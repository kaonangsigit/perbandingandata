import streamlit as st
import pandas as pd
import io
import re
import os
import subprocess
import logging
from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

import threading as _threading
import json as _json
_insw_threads = {}

def _insw_state_path(sid):
    return f"/tmp/insw_{sid}.json"

def _write_insw_state(sid, state):
    path = _insw_state_path(sid)
    tmp_path = path + ".tmp"
    try:
        with open(tmp_path, 'w') as f:
            _json.dump(state, f)
        os.rename(tmp_path, path)
    except Exception as e:
        logger.error(f"[INSW] Failed to write state file: {e}")

def _read_insw_state(sid):
    path = _insw_state_path(sid)
    try:
        with open(path, 'r') as f:
            return _json.load(f)
    except (FileNotFoundError, _json.JSONDecodeError, OSError):
        return {}

def _cleanup_insw_state(sid):
    path = _insw_state_path(sid)
    try:
        os.remove(path)
    except (FileNotFoundError, OSError):
        pass

def _setup_playwright_env():
    if os.environ.get('_PLAYWRIGHT_SETUP_DONE'):
        return

    if 'libgbm' not in os.environ.get("LD_LIBRARY_PATH", ""):
        gbm_lib_dir = None
        try:
            r = subprocess.run(["pkg-config", "--libs-only-L", "gbm"],
                             capture_output=True, text=True, timeout=5)
            if r.returncode == 0 and r.stdout.strip():
                gbm_lib_dir = r.stdout.strip().replace("-L", "")
        except Exception:
            pass

        if not gbm_lib_dir:
            try:
                r = subprocess.run(["nix-build", "<nixpkgs>", "-A", "libgbm", "--no-out-link"],
                                 capture_output=True, text=True, timeout=30)
                p = r.stdout.strip()
                if p and os.path.exists(p + "/lib/libgbm.so.1"):
                    gbm_lib_dir = p + "/lib"
            except Exception:
                pass

        if gbm_lib_dir and os.path.isdir(gbm_lib_dir):
            os.environ["LD_LIBRARY_PATH"] = gbm_lib_dir + ":" + os.environ.get("LD_LIBRARY_PATH", "")

    for bpath in [os.path.expanduser("~/.cache/ms-playwright"),
                  os.path.join(os.getcwd(), ".cache/ms-playwright"),
                  "/home/runner/workspace/.cache/ms-playwright"]:
        if os.path.exists(bpath) and os.listdir(bpath):
            os.environ["PLAYWRIGHT_BROWSERS_PATH"] = bpath
            break
    else:
        try:
            bp = os.path.expanduser("~/.cache/ms-playwright")
            os.environ["PLAYWRIGHT_BROWSERS_PATH"] = bp
            subprocess.run(["python3", "-m", "playwright", "install", "chromium"],
                         capture_output=True, timeout=120, env={**os.environ, "PLAYWRIGHT_BROWSERS_PATH": bp})
        except Exception:
            pass

    os.environ['_PLAYWRIGHT_SETUP_DONE'] = '1'

_setup_playwright_env()

st.set_page_config(
    page_title="Perbandingan Data Impor", 
    page_icon="📊", 
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        background: linear-gradient(90deg, #1e3a8a, #3b82f6);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        padding: 1rem 0;
    }
    .sub-header {
        font-size: 1.2rem;
        color: #64748b;
        text-align: center;
        margin-bottom: 2rem;
    }
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        padding: 10px 20px;
        background-color: #f1f5f9;
        border-radius: 10px;
        font-weight: 600;
    }
    .stTabs [aria-selected="true"] {
        background-color: #3b82f6;
        color: white;
    }
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 10px;
        color: white;
    }
    .upload-section {
        background-color: #f8fafc;
        padding: 1.5rem;
        border-radius: 15px;
        border: 2px dashed #cbd5e1;
        margin-bottom: 1rem;
    }
    .success-box {
        background-color: #dcfce7;
        border-left: 4px solid #22c55e;
        padding: 1rem;
        border-radius: 5px;
    }
    .warning-box {
        background-color: #fef3c7;
        border-left: 4px solid #f59e0b;
        padding: 1rem;
        border-radius: 5px;
    }
    .info-box {
        background-color: #e0f2fe;
        border-left: 4px solid #0ea5e9;
        padding: 1rem;
        border-radius: 5px;
    }
</style>
""", unsafe_allow_html=True)

st.markdown('<p class="main-header">📊 Perbandingan Data Realisasi Impor</p>', unsafe_allow_html=True)
st.markdown('<p class="sub-header">Aplikasi untuk membandingkan dan menganalisis data impor dengan mudah</p>', unsafe_allow_html=True)

tab_main, tab_hs, tab_analysis, tab_petugas, tab_absen, tab_importir, tab_merge, tab_notulen, tab_laporan_magang, tab_pdf = st.tabs(["📋 Perbandingan Data", "💊 Cek HS Code Obat", "📈 Analisis Data", "👤 Cek Petugas Loket S2", "📋 Cek Kehadiran", "🏢 Analisis Importir", "🔗 Gabung Data Excel", "📝 Notulen Rapat", "🎓 Laporan Magang BPOM", "📄 Edit PDF"])

def clean_value(value):
    if pd.isna(value):
        return ''
    val_str = str(value).strip()
    
    val_str = val_str.replace("'", "").replace('"', "").replace("'", "").replace("'", "")
    val_str = val_str.replace(";", "").replace(",", "")
    
    date_patterns = [
        r'\s*/\s*\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s+\d{4}',
        r'\s*/\s*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}',
        r'\s*/\s*\d{4}[-/]\d{1,2}[-/]\d{1,2}',
        r'\s*-\s*\d{1,2}\s+(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|Januari|Februari|Maret|April|Mei|Juni|Juli|Agustus|September|Oktober|November|Desember)\s+\d{4}',
    ]
    
    for pattern in date_patterns:
        val_str = re.sub(pattern, '', val_str, flags=re.IGNORECASE)
    
    val_str = re.sub(r'\s+', ' ', val_str)
    
    return val_str.strip()

def clean_number(value):
    if pd.isna(value):
        return ''
    val_str = str(value).strip()
    val_str = val_str.replace("'", "").replace('"', "").replace("'", "").replace("'", "")
    val_str = re.sub(r'[^\d]', '', val_str)
    return val_str

def get_invoice_list(value):
    if pd.isna(value):
        return []
    val_str = str(value).strip()
    val_str = val_str.replace("'", "").replace('"', "").replace("'", "").replace("'", "")
    
    if ';' in val_str or ',' in val_str:
        val_str = val_str.replace(';', ',')
        invoices = [inv.strip().strip(';').strip(',').strip() for inv in val_str.split(',')]
        invoices = [inv for inv in invoices if inv]
        return invoices
    
    val_str = val_str.strip(';').strip(',').strip()
    return [val_str] if val_str else []

def find_invoice_column(df):
    for col in df.columns:
        col_lower = str(col).lower().strip()
        if 'invoice' in col_lower and 'no' in col_lower:
            return col
        if col_lower == 'no. invoice' or col_lower == 'no.invoice' or col_lower == 'noinvoice':
            return col
    for col in df.columns:
        if 'invoice' in str(col).lower():
            return col
    return None

def load_invoice_set(file_invoice, label):
    invoice_set = set()
    if file_invoice:
        df_invoice = pd.read_excel(file_invoice)
        invoice_col = find_invoice_column(df_invoice)
        if invoice_col:
            for inv_value in df_invoice[invoice_col].dropna():
                inv_list = get_invoice_list(inv_value)
                invoice_set.update(inv_list)
            st.success(f"✅ **{label}**: {len(invoice_set)} NO. INVOICE unik ditemukan")
        else:
            st.warning(f"⚠️ Kolom NO. INVOICE tidak ditemukan di {label}")
    return invoice_set

def is_numeric_column(col_name):
    col_lower = str(col_name).lower()
    numeric_keywords = ['pib', 'pengajuan']
    return any(keyword in col_lower for keyword in numeric_keywords)

with tab_main:
    st.markdown("### 📁 Upload File")
    
    with st.expander("📖 Petunjuk Penggunaan", expanded=False):
        st.markdown("""
        1. Upload **File Tarikan** (bisa multiple file, akan digabung otomatis)
        2. Upload **File Data Anda** (data yang ingin dibandingkan)
        3. **Pilih kolom** yang ingin digunakan untuk perbandingan
        4. Upload **File Invoice** (opsional) untuk cek NO. INVOICE
        5. Klik **Bandingkan Data**
        6. Download hasil: Data SAMA = **Kuning**, Data berbeda = **Putih**
        """)

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 📥 File Tarikan (bisa multiple)")
        files_tarikan = st.file_uploader("Data hasil tarikan sistem", type=['xlsx', 'xls'], key="tarikan", help="Upload file Excel dari sistem (bisa pilih banyak file)", accept_multiple_files=True)

    with col2:
        st.markdown("#### 📤 File Data Anda")
        file_upload = st.file_uploader("Data Anda untuk dibandingkan", type=['xlsx', 'xls'], key="upload", help="Upload file Excel Anda")

    st.markdown("---")
    st.markdown("### 📑 File Invoice (Opsional)")

    col1, col2 = st.columns(2)

    with col1:
        st.markdown("#### 💊 Bahan Tambahan Obat")
        file_invoice_obat = st.file_uploader("File Invoice Bahan Tambahan Obat", type=['xlsx', 'xls'], key="invoice_obat")

    with col2:
        st.markdown("#### 🧪 Bahan Kimia")
        file_invoice_kimia = st.file_uploader("File Invoice Bahan Kimia", type=['xlsx', 'xls'], key="invoice_kimia")

    if files_tarikan and file_upload:
        try:
            tarikan_files_data = []
            dfs_all = []
            for f in files_tarikan:
                df_temp = pd.read_excel(f)
                file_name = f.name.replace('.xlsx', '').replace('.xls', '')[:31]
                tarikan_files_data.append({'name': file_name, 'df': df_temp})
                dfs_all.append(df_temp)
            
            df_tarikan = pd.concat(dfs_all, ignore_index=True)
            st.success(f"✅ {len(files_tarikan)} file tarikan dimuat: {len(df_tarikan)} baris total")
            
            df_upload = pd.read_excel(file_upload)
            
            st.markdown("---")
            st.markdown("### ⚙️ Konfigurasi Perbandingan")
            
            col_tarikan_list = df_tarikan.columns.tolist()
            col_upload_list = df_upload.columns.tolist()
            
            common_cols = [col for col in col_tarikan_list if col in col_upload_list]
            
            col1, col2 = st.columns(2)
            
            with col1:
                selected_col_tarikan = st.selectbox(
                    "📌 Kolom File Tarikan",
                    options=col_tarikan_list,
                    index=0,
                    key="col_tarikan"
                )
            
            with col2:
                default_index = col_upload_list.index(selected_col_tarikan) if selected_col_tarikan in col_upload_list else 0
                selected_col_upload = st.selectbox(
                    "📌 Kolom File Data Anda",
                    options=col_upload_list,
                    index=default_index,
                    key="col_upload"
                )
            
            use_numeric_cleaning = st.checkbox(
                "🔢 Bersihkan numerik saja (HANYA untuk kolom angka murni seperti NO. PIB)",
                value=is_numeric_column(selected_col_tarikan),
                help="⚠️ JANGAN centang jika data mengandung huruf seperti ST.03.04.35.352A..."
            )
            
            st.markdown("---")
            st.markdown("### 📋 Pilihan Jenis Output Download")
            output_option = st.radio(
                "Pilih jenis output yang diinginkan:",
                options=[
                    "❌ Download HANYA data yang TIDAK ADA di file lain (Output Lama)",
                    "📊 Download SEMUA data dengan highlight kuning untuk yang SAMA (Output Baru)"
                ],
                index=0,
                help="Pilih jenis output: Output Lama = hanya data tidak cocok, Output Baru = semua data dengan warna"
            )
            
            if common_cols:
                st.info(f"💡 Kolom yang sama di kedua file: **{', '.join(common_cols)}**")
            
            invoice_col_tarikan = find_invoice_column(df_tarikan)
            
            invoice_set_obat = set()
            invoice_set_kimia = set()
            
            st.markdown("---")
            st.markdown("### 📋 Status File Invoice")
            
            col1, col2 = st.columns(2)
            with col1:
                if file_invoice_obat:
                    invoice_set_obat = load_invoice_set(file_invoice_obat, "Bahan Tambahan Obat")
                else:
                    st.info("📭 File Invoice Bahan Tambahan Obat belum diupload")
            
            with col2:
                if file_invoice_kimia:
                    invoice_set_kimia = load_invoice_set(file_invoice_kimia, "Bahan Kimia")
                else:
                    st.info("📭 File Invoice Bahan Kimia belum diupload")
            
            st.markdown("---")
            st.markdown("### 👀 Preview Data")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown("#### 📊 Data Tarikan")
                st.caption(f"📝 {len(df_tarikan)} baris | Kolom: **{selected_col_tarikan}**")
                st.dataframe(df_tarikan.head(5), use_container_width=True, height=200)
            
            with col2:
                st.markdown("#### 📊 Data Anda")
                st.caption(f"📝 {len(df_upload)} baris | Kolom: **{selected_col_upload}**")
                st.dataframe(df_upload.head(5), use_container_width=True, height=200)
            
            st.markdown("---")
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                compare_btn = st.button("🔍 Bandingkan Data", type="primary", use_container_width=True)
            
            if compare_btn:
                st.markdown("---")
                st.markdown(f"### 📊 Hasil Perbandingan: {selected_col_tarikan}")
                
                if use_numeric_cleaning:
                    df_tarikan['_clean_key'] = df_tarikan[selected_col_tarikan].apply(clean_number)
                    df_upload['_clean_key'] = df_upload[selected_col_upload].apply(clean_number)
                else:
                    df_tarikan['_clean_key'] = df_tarikan[selected_col_tarikan].apply(clean_value)
                    df_upload['_clean_key'] = df_upload[selected_col_upload].apply(clean_value)
                
                with st.expander("🔎 Preview Hasil Pembersihan Data (klik untuk lihat)", expanded=False):
                    st.markdown("**File Tarikan - Sample Data Sebelum & Sesudah Pembersihan:**")
                    preview_tarikan = df_tarikan[[selected_col_tarikan, '_clean_key']].head(5).copy()
                    preview_tarikan.columns = ['Data Asli', 'Setelah Dibersihkan']
                    st.dataframe(preview_tarikan, use_container_width=True)
                    
                    st.markdown("**File Anda - Sample Data Sebelum & Sesudah Pembersihan:**")
                    preview_upload = df_upload[[selected_col_upload, '_clean_key']].head(5).copy()
                    preview_upload.columns = ['Data Asli', 'Setelah Dibersihkan']
                    st.dataframe(preview_upload, use_container_width=True)
                
                tarikan_keys = set(df_tarikan['_clean_key'].dropna())
                tarikan_keys = {k for k in tarikan_keys if k != ''}
                
                upload_keys = set(df_upload['_clean_key'].dropna())
                upload_keys = {k for k in upload_keys if k != ''}
                
                matching_keys = tarikan_keys & upload_keys
                missing_in_upload = tarikan_keys - upload_keys
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("📥 Data Tarikan", len(tarikan_keys), help="Jumlah data unik di file tarikan")
                with col2:
                    st.metric("📤 Data Anda", len(upload_keys), help="Jumlah data unik di file Anda")
                with col3:
                    st.metric("✅ Data SAMA", len(matching_keys), delta=f"+{len(matching_keys)}" if matching_keys else None, delta_color="normal", help="Data yang ada di KEDUA file")
                with col4:
                    st.metric("❌ Tidak Ada", len(missing_in_upload), delta=f"-{len(missing_in_upload)}" if missing_in_upload else None, delta_color="inverse", help="Data tarikan yang tidak ada di file Anda")
                
                df_tarikan_display = df_tarikan.copy()
                df_tarikan_display['Status'] = df_tarikan_display['_clean_key'].apply(
                    lambda x: '✅ Sama' if x in matching_keys else '❌ Tidak Sama'
                )
                df_tarikan_display = df_tarikan_display.drop(columns=['_clean_key'])
                
                jumlah_sama = len(df_tarikan_display[df_tarikan_display['Status'] == '✅ Sama'])
                jumlah_tidak_sama = len(df_tarikan_display[df_tarikan_display['Status'] == '❌ Tidak Sama'])
                
                if missing_in_upload:
                    st.markdown(f"### 🔴 Data Tarikan yang Tidak Ada di File Anda")
                    st.warning(f"Ditemukan **{len(missing_in_upload)}** data unik dari tarikan yang tidak ada di file Anda.")
                    
                    df_missing = df_tarikan[df_tarikan['_clean_key'].isin(missing_in_upload)].copy()
                    df_missing = df_missing.drop(columns=['_clean_key'])
                    
                    st.dataframe(df_missing, use_container_width=True, height=300)
                
                st.markdown("---")
                
                show_only_missing = "HANYA" in output_option
                
                if show_only_missing:
                    st.markdown("### 📥 Download Data yang TIDAK ADA (Output Lama)")
                    st.markdown("File Excel berisi **hanya data yang tidak ada** di file lain")
                    
                    if missing_in_upload:
                        df_missing = df_tarikan[df_tarikan['_clean_key'].isin(missing_in_upload)].copy()
                        df_missing = df_missing.drop(columns=['_clean_key'])
                        
                        output_missing = io.BytesIO()
                        with pd.ExcelWriter(output_missing, engine='openpyxl') as writer:
                            header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                            header_font = Font(bold=True, color='FFFFFF')
                            thin_border = Border(
                                left=Side(style='thin'),
                                right=Side(style='thin'),
                                top=Side(style='thin'),
                                bottom=Side(style='thin')
                            )
                            
                            for file_data in tarikan_files_data:
                                df_file = file_data['df'].copy()
                                if use_numeric_cleaning:
                                    df_file['_clean_key'] = df_file[selected_col_tarikan].apply(clean_number)
                                else:
                                    df_file['_clean_key'] = df_file[selected_col_tarikan].apply(clean_value)
                                
                                df_file_missing = df_file[df_file['_clean_key'].isin(missing_in_upload)].copy()
                                df_file_missing = df_file_missing.drop(columns=['_clean_key'])
                                
                                if len(df_file_missing) > 0:
                                    sheet_name = file_data['name'][:31]
                                    df_file_missing.to_excel(writer, index=False, sheet_name=sheet_name)
                                    
                                    worksheet = writer.sheets[sheet_name]
                                    for col_idx, col in enumerate(df_file_missing.columns, 1):
                                        cell = worksheet.cell(row=1, column=col_idx)
                                        cell.fill = header_fill
                                        cell.font = header_font
                                        cell.alignment = Alignment(horizontal='center')
                                        cell.border = thin_border
                                    
                                    for row_idx in range(2, len(df_file_missing) + 2):
                                        for col_idx in range(1, len(df_file_missing.columns) + 1):
                                            cell = worksheet.cell(row=row_idx, column=col_idx)
                                            cell.border = thin_border
                                    
                                    for col_idx, col in enumerate(df_file_missing.columns, 1):
                                        max_len = max(df_file_missing[col].astype(str).apply(len).max(), len(str(col))) + 2
                                        worksheet.column_dimensions[worksheet.cell(row=1, column=col_idx).column_letter].width = min(max_len, 50)
                        
                        output_missing.seek(0)
                        
                        st.metric("❌ Total Data Tidak Ada", len(missing_in_upload))
                        
                        st.download_button(
                            label="📥 Download Data yang Tidak Ada",
                            data=output_missing,
                            file_name="data_tidak_ada.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                    else:
                        st.success("✅ Semua data tarikan sudah ada di file Anda!")
                
                else:
                    st.markdown("### 📊 Download Data Lengkap dengan Warna (Output Baru)")
                    st.markdown(f"File Excel akan memiliki **{len(tarikan_files_data) + 1} sheet/laman**:")
                    for i, file_data in enumerate(tarikan_files_data, 1):
                        st.markdown(f"- 📥 **Sheet {i}**: {file_data['name']}")
                    st.markdown(f"- 📤 **Sheet Terakhir**: Data Anda")
                    st.markdown("- 🟡 **Warna Kuning**: Data yang **SAMA** di kedua file")
                    st.markdown("- ⬜ **Putih**: Data yang **TIDAK ADA** di file lain")
                    
                    col1, col2 = st.columns(2)
                    with col1:
                        st.metric("🟡 Data Kuning (Sama)", jumlah_sama)
                    with col2:
                        st.metric("⬜ Data Putih (Tidak Sama)", jumlah_tidak_sama)
                    
                    df_upload_display = df_upload.copy()
                    df_upload_display['Status'] = df_upload_display['_clean_key'].apply(
                        lambda x: '✅ Sama' if x in matching_keys else '❌ Tidak Sama'
                    )
                    df_upload_display = df_upload_display.drop(columns=['_clean_key'])
                    
                    output_colored = io.BytesIO()
                    with pd.ExcelWriter(output_colored, engine='openpyxl') as writer:
                        yellow_fill = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                        header_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                        header_font = Font(bold=True, color='FFFFFF')
                        thin_border = Border(
                            left=Side(style='thin'),
                            right=Side(style='thin'),
                            top=Side(style='thin'),
                            bottom=Side(style='thin')
                        )
                        
                        for file_data in tarikan_files_data:
                            df_file = file_data['df'].copy()
                            if use_numeric_cleaning:
                                df_file['_clean_key'] = df_file[selected_col_tarikan].apply(clean_number)
                            else:
                                df_file['_clean_key'] = df_file[selected_col_tarikan].apply(clean_value)
                            df_file['Status'] = df_file['_clean_key'].apply(
                                lambda x: '✅ Sama' if x in matching_keys else '❌ Tidak Sama'
                            )
                            df_file = df_file.drop(columns=['_clean_key'])
                            
                            sheet_name = file_data['name'][:31]
                            df_file.to_excel(writer, index=False, sheet_name=sheet_name)
                            
                            worksheet = writer.sheets[sheet_name]
                            for col_idx, col in enumerate(df_file.columns, 1):
                                cell = worksheet.cell(row=1, column=col_idx)
                                cell.fill = header_fill
                                cell.font = header_font
                                cell.alignment = Alignment(horizontal='center')
                                cell.border = thin_border
                            
                            status_col_idx = df_file.columns.get_loc('Status') + 1
                            for row_idx in range(2, len(df_file) + 2):
                                status_cell = worksheet.cell(row=row_idx, column=status_col_idx)
                                if '✅' in str(status_cell.value):
                                    for col_idx in range(1, len(df_file.columns) + 1):
                                        cell = worksheet.cell(row=row_idx, column=col_idx)
                                        cell.fill = yellow_fill
                                        cell.border = thin_border
                                else:
                                    for col_idx in range(1, len(df_file.columns) + 1):
                                        cell = worksheet.cell(row=row_idx, column=col_idx)
                                        cell.border = thin_border
                            
                            for col_idx, col in enumerate(df_file.columns, 1):
                                max_len = max(df_file[col].astype(str).apply(len).max(), len(str(col))) + 2
                                worksheet.column_dimensions[worksheet.cell(row=1, column=col_idx).column_letter].width = min(max_len, 50)
                        
                        df_upload_display.to_excel(writer, index=False, sheet_name='Data Anda')
                        worksheet_upload = writer.sheets['Data Anda']
                        
                        for col_idx, col in enumerate(df_upload_display.columns, 1):
                            cell = worksheet_upload.cell(row=1, column=col_idx)
                            cell.fill = header_fill
                            cell.font = header_font
                            cell.alignment = Alignment(horizontal='center')
                            cell.border = thin_border
                        
                        status_col_idx = df_upload_display.columns.get_loc('Status') + 1
                        for row_idx in range(2, len(df_upload_display) + 2):
                            status_cell = worksheet_upload.cell(row=row_idx, column=status_col_idx)
                            if '✅' in str(status_cell.value):
                                for col_idx in range(1, len(df_upload_display.columns) + 1):
                                    cell = worksheet_upload.cell(row=row_idx, column=col_idx)
                                    cell.fill = yellow_fill
                                    cell.border = thin_border
                            else:
                                for col_idx in range(1, len(df_upload_display.columns) + 1):
                                    cell = worksheet_upload.cell(row=row_idx, column=col_idx)
                                    cell.border = thin_border
                        
                        for col_idx, col in enumerate(df_upload_display.columns, 1):
                            max_len = max(df_upload_display[col].astype(str).apply(len).max(), len(str(col))) + 2
                            worksheet_upload.column_dimensions[worksheet_upload.cell(row=1, column=col_idx).column_letter].width = min(max_len, 50)
                        
                    output_colored.seek(0)
                    
                    st.download_button(
                        label="📥 Download Excel dengan Warna",
                        data=output_colored,
                        file_name="hasil_perbandingan_berwarna.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                if missing_in_upload:
                    
                    if invoice_col_tarikan and (invoice_set_obat or invoice_set_kimia):
                        st.markdown("---")
                        st.markdown("### 📋 Cek NO. INVOICE")
                        
                        def check_invoice_obat(inv_value):
                            if not invoice_set_obat:
                                return '-'
                            inv_list = get_invoice_list(inv_value)
                            if not inv_list:
                                return '❌ Tidak Ada'
                            found = sum(1 for inv in inv_list if inv in invoice_set_obat)
                            if found == len(inv_list):
                                return '✅ Ada'
                            elif found > 0:
                                return f'⚠️ Sebagian ({found}/{len(inv_list)})'
                            else:
                                return '❌ Tidak Ada'
                        
                        def check_invoice_kimia(inv_value):
                            if not invoice_set_kimia:
                                return '-'
                            inv_list = get_invoice_list(inv_value)
                            if not inv_list:
                                return '❌ Tidak Ada'
                            found = sum(1 for inv in inv_list if inv in invoice_set_kimia)
                            if found == len(inv_list):
                                return '✅ Ada'
                            elif found > 0:
                                return f'⚠️ Sebagian ({found}/{len(inv_list)})'
                            else:
                                return '❌ Tidak Ada'
                        
                        df_invoice_check = df_missing.copy()
                        
                        if invoice_set_obat:
                            st.markdown("#### 💊 Cek di Bahan Tambahan Obat")
                            df_invoice_check['Cek Bahan Obat'] = df_invoice_check[invoice_col_tarikan].apply(check_invoice_obat)
                            
                            ada_obat = df_invoice_check[df_invoice_check['Cek Bahan Obat'] == '✅ Ada']
                            sebagian_obat = df_invoice_check[df_invoice_check['Cek Bahan Obat'].str.contains('Sebagian', na=False)]
                            tidak_obat = df_invoice_check[df_invoice_check['Cek Bahan Obat'] == '❌ Tidak Ada']
                            
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("✅ Ada", len(ada_obat))
                            with col2:
                                st.metric("⚠️ Sebagian", len(sebagian_obat))
                            with col3:
                                st.metric("❌ Tidak Ada", len(tidak_obat))
                        
                        if invoice_set_kimia:
                            st.markdown("#### 🧪 Cek di Bahan Kimia")
                            df_invoice_check['Cek Bahan Kimia'] = df_invoice_check[invoice_col_tarikan].apply(check_invoice_kimia)
                            
                            ada_kimia = df_invoice_check[df_invoice_check['Cek Bahan Kimia'] == '✅ Ada']
                            sebagian_kimia = df_invoice_check[df_invoice_check['Cek Bahan Kimia'].str.contains('Sebagian', na=False)]
                            tidak_kimia = df_invoice_check[df_invoice_check['Cek Bahan Kimia'] == '❌ Tidak Ada']
                            
                            col1, col2, col3 = st.columns(3)
                            with col1:
                                st.metric("✅ Ada", len(ada_kimia))
                            with col2:
                                st.metric("⚠️ Sebagian", len(sebagian_kimia))
                            with col3:
                                st.metric("❌ Tidak Ada", len(tidak_kimia))
                        
                        st.markdown("#### 📊 Data Lengkap dengan Status Invoice")
                        st.dataframe(df_invoice_check, use_container_width=True, height=300)
                        
                        if invoice_set_obat:
                            st.markdown("---")
                            st.markdown("##### 💊 Filter Bahan Tambahan Obat")
                            tab1, tab2, tab3 = st.tabs(["✅ Ada", "⚠️ Sebagian", "❌ Tidak Ada"])
                            
                            with tab1:
                                if len(ada_obat) > 0:
                                    st.dataframe(ada_obat, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                            
                            with tab2:
                                if len(sebagian_obat) > 0:
                                    st.dataframe(sebagian_obat, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                            
                            with tab3:
                                if len(tidak_obat) > 0:
                                    st.dataframe(tidak_obat, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                        
                        if invoice_set_kimia:
                            st.markdown("---")
                            st.markdown("##### 🧪 Filter Bahan Kimia")
                            tab1, tab2, tab3 = st.tabs(["✅ Ada ", "⚠️ Sebagian ", "❌ Tidak Ada "])
                            
                            with tab1:
                                if len(ada_kimia) > 0:
                                    st.dataframe(ada_kimia, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                            
                            with tab2:
                                if len(sebagian_kimia) > 0:
                                    st.dataframe(sebagian_kimia, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                            
                            with tab3:
                                if len(tidak_kimia) > 0:
                                    st.dataframe(tidak_kimia, use_container_width=True)
                                else:
                                    st.info("Tidak ada data")
                        
                        output_invoice = io.BytesIO()
                        with pd.ExcelWriter(output_invoice, engine='openpyxl') as writer:
                            df_invoice_check.to_excel(writer, index=False, sheet_name='Hasil Cek Invoice')
                        output_invoice.seek(0)
                        
                        st.download_button(
                            label="📥 Download Hasil Cek Invoice",
                            data=output_invoice,
                            file_name="hasil_cek_invoice.xlsx",
                            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                            use_container_width=True
                        )
                        
                else:
                    st.success("✅ Semua data dari tarikan sudah tersedia di file Anda!")
                
        except Exception as e:
            st.error(f"❌ Terjadi kesalahan: {str(e)}")
            st.info("💡 Pastikan file Excel dalam format yang benar (.xlsx atau .xls)")

    else:
        st.info("👆 Silakan upload **File Tarikan** dan **File Data Anda** untuk memulai perbandingan.")

with tab_hs:
    st.markdown("### 🌐 Cek INSW Otomatis (Indonesia National Single Window)")
    st.markdown("Upload file data dari BPS, pilih chapter yang ingin dicek, lalu sistem akan mengecek otomatis di website INSW untuk mengetahui **regulasi impor**, **regulasi ekspor**, dan klasifikasi **obat/farmasi**.")

    with st.expander("📖 Petunjuk Penggunaan", expanded=False):
        st.markdown("""
        1. Upload file data dari **BPS** (format .xlsx/.xls)
        2. Sistem mendeteksi semua **chapter** HS Code dalam file
        3. **Pilih chapter** yang ingin dicek (bebas pilih berapa pun)
        4. Klik **Mulai Cek INSW Otomatis**
        5. Lihat hasil: regulasi **impor**, **ekspor**, **BPOM**, dan klasifikasi **obat**
        6. Download hasil dalam format Excel
        """)

    file_hs = st.file_uploader("📁 Upload file data BPS", type=['xlsx', 'xls'], key="hs_check")

    if file_hs:
        try:
            xls = pd.ExcelFile(file_hs)
            all_sheet_names = xls.sheet_names

            selected_sheet = all_sheet_names[0]
            if len(all_sheet_names) > 1:
                selected_sheet = st.selectbox(
                    "📄 Pilih Sheet:",
                    options=all_sheet_names,
                    key="sheet_select"
                )

            df_hs_raw = pd.read_excel(xls, sheet_name=selected_sheet, header=None, dtype=str)
            for col in df_hs_raw.columns:
                df_hs_raw[col] = df_hs_raw[col].astype(str).replace('nan', '')

            header_row = None
            for i in range(min(10, len(df_hs_raw))):
                val = str(df_hs_raw.iloc[i, 0]).strip().lower()
                if 'kode hs' in val or 'hs code' in val:
                    header_row = i
                    break

            if header_row is None:
                header_row = 3

            data_start = header_row + 1

            hs_items = []
            for idx in range(data_start, len(df_hs_raw)):
                val = str(df_hs_raw.iloc[idx, 0]).strip()
                match = re.match(r'\[(\d+)\]\s*(.*)', val)
                if match:
                    code = match.group(1)
                    desc = match.group(2).strip()
                    hs_items.append({
                        'row_idx': idx,
                        'raw_value': val,
                        'hs_code': code,
                        'description': desc,
                        'prefix': code[:2]
                    })

            st.success(f"Sheet **{selected_sheet}**: Total **{len(hs_items)}** HS Code ditemukan")

            all_prefixes = sorted(list(set(h['prefix'] for h in hs_items)))
            prefix_counts = {}
            for p in all_prefixes:
                prefix_counts[p] = len([h for h in hs_items if h['prefix'] == p])

            st.markdown("---")
            st.markdown("### 📊 Pilih Chapter untuk Dicek")

            chapter_labels = []
            for p in all_prefixes:
                chapter_labels.append(f"{p} ({prefix_counts[p]} HS Code)")

            default_chapters = [lbl for lbl in chapter_labels if lbl.startswith(('28 ', '29 ', '30 ', '31 '))]

            selected_chapter_labels = st.multiselect(
                "🔎 Pilih Chapter (awalan HS Code):",
                options=chapter_labels,
                default=default_chapters,
                key="chapter_select",
                help="Pilih chapter yang ingin dicek di INSW. Bisa pilih berapa pun, bebas kombinasi."
            )

            selected_prefixes = [lbl.split(' ')[0] for lbl in selected_chapter_labels]

            hs_filtered = [h for h in hs_items if h['prefix'] in selected_prefixes]

            if selected_prefixes:
                n_cols = min(len(selected_prefixes) + 1, 6)
                cols = st.columns(n_cols)
                with cols[0]:
                    st.metric("Total Terpilih", len(hs_filtered))
                for i, p in enumerate(selected_prefixes[:n_cols-1]):
                    with cols[i + 1]:
                        st.metric(f"Chapter {p}", prefix_counts.get(p, 0))

            all_hs_desc_map = {h['hs_code']: h['description'] for h in hs_items}
            codes_to_check = list(dict.fromkeys([h['hs_code'] for h in hs_filtered]))

            st.markdown("---")

            if 'playwright_available' not in st.session_state:
                try:
                    from playwright.sync_api import sync_playwright as _pw_check
                    with _pw_check() as _pw_test:
                        _test_browser = _pw_test.chromium.launch(
                            headless=True,
                            args=['--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu', '--single-process']
                        )
                        _test_browser.close()
                        st.session_state['playwright_available'] = True
                except Exception as _pw_err:
                    st.session_state['playwright_available'] = False
                    st.session_state['playwright_error'] = str(_pw_err)
            playwright_available = st.session_state['playwright_available']

            btn_insw = False
            if not playwright_available:
                st.warning("Browser otomatis (Playwright/Chromium) tidak tersedia. Gunakan link berikut untuk cek manual.")
                st.markdown(f"[Buka INSW Detail Komoditas](https://insw.go.id/intr/detail-komoditas)")
                insw_manual_data = [{'No': i+1, 'HS Code': h['hs_code'], 'Deskripsi': h['description']} for i, h in enumerate(hs_filtered[:200])]
                if insw_manual_data:
                    st.dataframe(pd.DataFrame(insw_manual_data), use_container_width=True, height=300)
                if len(hs_filtered) > 200:
                    st.caption(f"Menampilkan 200 dari {len(hs_filtered)} HS Code")
            else:
                if len(codes_to_check) == 0:
                    st.warning("Tidak ada HS Code yang dipilih. Pilih minimal 1 chapter di atas.")
                else:
                    est_seconds = len(codes_to_check) * 5
                    est_minutes = max(1, est_seconds // 60)

                    st.info(f"Akan mengecek **{len(codes_to_check)}** HS Code unik dari chapter **{', '.join(selected_prefixes)}**. Estimasi waktu: **~{est_minutes} menit**.")

                    if len(codes_to_check) > 0:
                        preview_data = [{'No': i+1, 'HS Code': c, 'Deskripsi': all_hs_desc_map.get(c, '')} for i, c in enumerate(codes_to_check[:10])]
                        st.dataframe(pd.DataFrame(preview_data), use_container_width=True, height=200)
                        if len(codes_to_check) > 10:
                            st.caption(f"... dan {len(codes_to_check) - 10} HS Code lainnya")

                    col_insw1, col_insw2, col_insw3 = st.columns([1, 2, 1])
                    with col_insw2:
                        btn_insw = st.button("🔍 Mulai Cek INSW Otomatis", type="primary", use_container_width=True, key="btn_insw")

            if playwright_available and btn_insw and len(codes_to_check) > 0 and not st.session_state.get('insw_running', False):
                st.session_state['insw_running'] = True
                st.session_state['insw_complete'] = False
                st.session_state.pop('insw_error', None)
                st.session_state['insw_results'] = []
                st.session_state['insw_progress_current'] = 0
                st.session_state['insw_progress_total'] = len(codes_to_check)
                st.session_state['insw_progress_hs'] = ''
                st.session_state['insw_progress_desc'] = ''
                st.session_state['insw_checked_prefixes'] = selected_prefixes
                st.session_state['insw_codes_to_check'] = list(codes_to_check)
                st.session_state['insw_desc_map'] = dict(all_hs_desc_map)

                import threading

                def _run_insw_scraping(codes, desc_map, session_id):
                    import time as _time

                    _file_state = {
                        'results': [],
                        'complete': False,
                        'current': 0,
                        'total': len(codes),
                        'current_hs': '',
                        'current_desc': '',
                        'error_count': 0,
                        'error_msg': '',
                        'heartbeat': _time.time(),
                        'status': 'running',
                    }
                    _write_insw_state(session_id, _file_state)

                    def _update_shared(key, value):
                        _file_state[key] = value
                        _file_state['heartbeat'] = _time.time()
                        _write_insw_state(session_id, _file_state)

                    def _update_shared_multi(updates):
                        _file_state.update(updates)
                        _file_state['heartbeat'] = _time.time()
                        _write_insw_state(session_id, _file_state)

                    INSW_URL = "https://insw.go.id/intr/detail-komoditas"
                    OBAT_KEYWORDS = ['obat', 'farmasi', 'pharmaceutical', 'medicine', 'drug',
                                    'suplemen kesehatan', 'bahan baku obat', 'kosmetik',
                                    'vaksin', 'vitamin', 'narkotik', 'psikotropik',
                                    'kuasi', 'prekursor', 'narkotika', 'psikotropika']
                    BROWSER_ARGS = [
                        '--no-sandbox', '--disable-dev-shm-usage', '--disable-gpu',
                        '--single-process', '--disable-extensions',
                        '--disable-background-networking',
                        '--disable-software-rasterizer',
                        '--disable-translate',
                        '--no-first-run',
                        '--no-zygote',
                    ]
                    max_retries = 3

                    def format_hs_dotted(code):
                        if len(code) == 8:
                            return f"{code[:4]}.{code[4:6]}.{code[6:8]}"
                        return code

                    def search_and_click_detail(pw_page, hs_code):
                        search_queries = [hs_code, format_hs_dotted(hs_code)]
                        for attempt, query in enumerate(search_queries):
                            try:
                                logger.info(f"[INSW] Searching {hs_code} with query '{query}' (attempt {attempt+1})")
                                pw_page.goto(INSW_URL, timeout=30000, wait_until='domcontentloaded')
                                pw_page.wait_for_timeout(2000)
                                search_input = pw_page.wait_for_selector("input[placeholder='Cari kode HS / Uraian HS']", timeout=20000)
                                search_input.fill(query)
                                search_input.press("Enter")
                                try:
                                    pw_page.wait_for_selector("button:has-text('Detail')", timeout=20000)
                                except Exception:
                                    logger.info(f"[INSW] No Detail button found for query '{query}'")
                                    continue
                                pw_page.wait_for_timeout(1500)
                                body_text = pw_page.inner_text("body")
                                if hs_code not in body_text:
                                    logger.info(f"[INSW] HS code {hs_code} not in search results for query '{query}'")
                                    continue
                                rows = pw_page.query_selector_all("tr")
                                for row in rows:
                                    row_text = row.inner_text()
                                    if hs_code in row_text:
                                        detail_btn = row.query_selector("button:has-text('Detail')")
                                        if detail_btn:
                                            detail_btn.click()
                                            pw_page.wait_for_timeout(3000)
                                            logger.info(f"[INSW] Clicked Detail for {hs_code}")
                                            return True
                                detail_btns = pw_page.query_selector_all("button:has-text('Detail')")
                                if detail_btns:
                                    detail_btns[0].click()
                                    pw_page.wait_for_timeout(3000)
                                    return True
                            except Exception as e:
                                logger.error(f"[INSW] Error searching {hs_code} with query '{query}': {str(e)[:100]}")
                                continue
                        return False

                    def extract_insw_detail(pw_page, hs_code, desc_text=''):
                        entry = {
                            'HS Code': hs_code, 'Deskripsi': desc_text, 'Jenis': '-',
                            'Ada Regulasi Impor': 'Tidak', 'Lartas Border': 'Tidak',
                            'Tata Niaga Post Border': 'Tidak', 'Ada Regulasi Ekspor': 'Tidak',
                            'Lartas Ekspor': 'Tidak', 'Komoditi INSW': '-',
                            'Terkait Obat (INSW)': 'Tidak', 'Ada BPOM': 'Tidak',
                            'Keterangan Impor': '-', 'Keterangan Ekspor': '-',
                        }
                        found = search_and_click_detail(pw_page, hs_code)
                        if not found:
                            entry['Jenis'] = 'Tidak ditemukan'
                            entry['Keterangan Impor'] = 'Tidak ditemukan di INSW'
                            entry['Keterangan Ekspor'] = 'Tidak ditemukan di INSW'
                            return entry
                        pw_page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
                        body = pw_page.inner_text("body")
                        has_lartas_border = "Regulasi Impor (Lartas Border)" in body
                        has_tata_niaga = "Regulasi Impor (Tata Niaga Post Border)" in body
                        has_import = has_lartas_border or has_tata_niaga or "Regulasi Impor" in body
                        has_lartas_ekspor = "Regulasi Ekspor (Lartas Ekspor)" in body or "Lartas Ekspor" in body
                        has_export = has_lartas_ekspor or "Regulasi Ekspor" in body
                        entry['Ada Regulasi Impor'] = 'YA' if has_import else 'Tidak'
                        entry['Lartas Border'] = 'YA' if has_lartas_border else 'Tidak'
                        entry['Tata Niaga Post Border'] = 'YA' if has_tata_niaga else 'Tidak'
                        entry['Ada Regulasi Ekspor'] = 'YA' if has_export else 'Tidak'
                        entry['Lartas Ekspor'] = 'YA' if has_lartas_ekspor else 'Tidak'
                        komoditi_list = []
                        is_obat = False
                        ket_impor_parts = []
                        ket_ekspor_parts = []
                        lines = body.split('\n')
                        for li, line in enumerate(lines):
                            stripped = line.strip()
                            if stripped == 'Komoditi':
                                for offset in range(1, 6):
                                    if li + offset < len(lines):
                                        next_line = lines[li + offset].strip()
                                        if next_line.startswith('[') and next_line.endswith(']'):
                                            komoditi_val = next_line[1:-1]
                                            if komoditi_val and komoditi_val not in komoditi_list:
                                                komoditi_list.append(komoditi_val)
                                            break
                                        elif next_line == ':':
                                            continue
                                        elif next_line and next_line not in ('Regulasi', 'Deskripsi', ''):
                                            break
                        if komoditi_list:
                            entry['Komoditi INSW'] = '; '.join(komoditi_list)
                            for k_val in komoditi_list:
                                k_lower = k_val.lower()
                                for ok in OBAT_KEYWORDS:
                                    if ok in k_lower:
                                        is_obat = True
                                        break
                        body_lower = body.lower()
                        if 'bahan obat' in body_lower or 'bahan baku obat' in body_lower:
                            is_obat = True
                        has_bpom = 'BPOM' in body
                        entry['Ada BPOM'] = 'YA' if has_bpom else 'Tidak'
                        if has_lartas_border:
                            ket_impor_parts.append('Lartas Border')
                        if has_tata_niaga:
                            ket_impor_parts.append('Tata Niaga Post Border')
                        if has_bpom:
                            ket_impor_parts.append('BPOM')
                        if is_obat:
                            ket_impor_parts.append('Terkait Obat/Farmasi')
                        if has_lartas_ekspor:
                            ket_ekspor_parts.append('Lartas Ekspor')
                        entry['Keterangan Impor'] = '; '.join(ket_impor_parts) if ket_impor_parts else 'Tidak ada regulasi impor'
                        entry['Keterangan Ekspor'] = '; '.join(ket_ekspor_parts) if ket_ekspor_parts else 'Tidak ada regulasi ekspor'
                        if has_import and has_export:
                            entry['Jenis'] = 'IMPOR & EKSPOR'
                        elif has_import:
                            entry['Jenis'] = 'IMPOR'
                        elif has_export:
                            entry['Jenis'] = 'EKSPOR'
                        else:
                            entry['Jenis'] = 'Tidak ada lartas'
                        entry['Terkait Obat (INSW)'] = 'YA' if is_obat else 'Tidak'
                        return entry

                    results = []
                    error_count = 0
                    pw_browser = None

                    try:
                        from playwright.sync_api import sync_playwright
                        with sync_playwright() as pw:
                            logger.info("[INSW-Thread] Launching Chromium browser...")
                            pw_browser = pw.chromium.launch(headless=True, args=BROWSER_ARGS)
                            pw_page = pw_browser.new_page()
                            pw_page.set_default_timeout(60000)
                            logger.info("[INSW-Thread] Browser launched successfully")

                            for idx_hs, hs_code in enumerate(codes):
                                _update_shared_multi({
                                    'current': idx_hs + 1,
                                    'current_hs': hs_code,
                                    'current_desc': desc_map.get(hs_code, '')[:60],
                                })

                                last_error_msg = ''
                                result_entry = None
                                for retry in range(max_retries + 1):
                                    try:
                                        result_entry = extract_insw_detail(pw_page, hs_code, desc_map.get(hs_code, ''))
                                        break
                                    except Exception as e_hs:
                                        last_error_msg = str(e_hs)[:120]
                                        logger.error(f"[INSW-Thread] Error on {hs_code} retry {retry}: {last_error_msg}")
                                        if retry < max_retries:
                                            try:
                                                pw_page.close()
                                            except Exception:
                                                pass
                                            try:
                                                pw_browser.close()
                                            except Exception:
                                                pass
                                            _time.sleep(2)
                                            try:
                                                pw_browser = pw.chromium.launch(headless=True, args=BROWSER_ARGS)
                                                pw_page = pw_browser.new_page()
                                                pw_page.set_default_timeout(60000)
                                                logger.info(f"[INSW-Thread] Browser restarted for retry {retry+1}")
                                            except Exception as e_launch:
                                                last_error_msg = f'Browser restart error: {str(e_launch)[:80]}'
                                                logger.error(f"[INSW-Thread] {last_error_msg}")
                                                break

                                if result_entry is None:
                                    error_count += 1
                                    result_entry = {
                                        'HS Code': hs_code, 'Deskripsi': desc_map.get(hs_code, ''),
                                        'Jenis': 'Error',
                                        'Ada Regulasi Impor': '-', 'Lartas Border': '-',
                                        'Tata Niaga Post Border': '-', 'Ada Regulasi Ekspor': '-',
                                        'Lartas Ekspor': '-', 'Komoditi INSW': '-',
                                        'Terkait Obat (INSW)': '-', 'Ada BPOM': '-',
                                        'Keterangan Impor': f'Error: {last_error_msg}',
                                        'Keterangan Ekspor': '-',
                                    }

                                results.append(result_entry)
                                _update_shared('results', list(results))

                            try:
                                pw_browser.close()
                            except Exception:
                                pass

                        _update_shared_multi({
                            'results': results,
                            'complete': True,
                            'error_count': error_count,
                            'status': 'completed',
                        })
                        logger.info(f"[INSW-Thread] Completed. {len(results)}/{len(codes)} checked, {error_count} errors")

                    except Exception as e_insw:
                        error_detail = str(e_insw)
                        logger.error(f"[INSW-Thread] Fatal error: {error_detail}")
                        _update_shared_multi({
                            'results': results,
                            'complete': True,
                            'error_msg': error_detail[:200],
                            'error_count': error_count,
                            'status': 'error',
                        })
                        try:
                            if pw_browser:
                                pw_browser.close()
                        except Exception:
                            pass

                import uuid
                sid = str(uuid.uuid4())[:8]
                st.session_state['insw_session_id'] = sid

                st.session_state['insw_thread_started'] = True

                t = threading.Thread(
                    target=_run_insw_scraping,
                    args=(list(codes_to_check), dict(all_hs_desc_map), sid),
                    daemon=True
                )
                t.start()
                _insw_threads[sid] = t
                st.rerun()

            if st.session_state.get('insw_running', False) and st.session_state.get('insw_thread_started', False):
                import time as _time
                sid = st.session_state.get('insw_session_id', '')

                shared = _read_insw_state(sid)

                thread = _insw_threads.get(sid)
                thread_alive = thread is not None and thread.is_alive()

                file_status = shared.get('status', '')
                heartbeat = shared.get('heartbeat', 0)
                heartbeat_stale = (heartbeat > 0 and (_time.time() - heartbeat) > 120)

                if not shared and not thread_alive:
                    st.session_state['insw_running'] = False
                    st.session_state.pop('insw_thread_started', None)
                    st.session_state['insw_error'] = "Proses terganggu (koneksi terputus). Silakan klik tombol 'Mulai Cek INSW Otomatis' lagi."
                    st.rerun()

                total = shared.get('total', st.session_state.get('insw_progress_total', 0))
                current = shared.get('current', 0)
                current_hs = shared.get('current_hs', '')
                current_desc = shared.get('current_desc', '')
                is_complete = shared.get('complete', False) or file_status in ('completed', 'error')
                partial_results = shared.get('results', [])
                error_msg = shared.get('error_msg', '')
                error_count = shared.get('error_count', 0)

                if not is_complete and heartbeat_stale:
                    is_complete = True
                    partial_results = shared.get('results', [])
                    error_msg = error_msg or "Proses scraping berhenti (tidak ada update selama 60 detik)"
                    logger.warning(f"[INSW] Heartbeat stale for {sid}, marking complete with {len(partial_results)} partial results")

                if not is_complete:
                    progress_val = current / total if total > 0 else 0
                    st.progress(progress_val, text=f"Mengecek HS Code {current_hs} ({current}/{total})...")
                    st.info(f"Sedang memproses: **{current_hs}** - {current_desc}")

                    if partial_results:
                        st.caption(f"{len(partial_results)} HS Code sudah dicek...")

                    _time.sleep(3)
                    st.rerun()
                else:
                    st.session_state['insw_results'] = partial_results
                    st.session_state['insw_running'] = False
                    st.session_state['insw_complete'] = True
                    st.session_state.pop('insw_thread_started', None)

                    if error_msg:
                        st.session_state['insw_error'] = f"Proses selesai dengan error. {len(partial_results)}/{total} HS Code dicek. Error: {error_msg}"
                    elif error_count > 0:
                        st.session_state['insw_error'] = f"Selesai! {len(partial_results)}/{total} HS Code dicek ({error_count} error)"

                    _cleanup_insw_state(sid)
                    _insw_threads.pop(sid, None)

                    st.rerun()

            if st.session_state.get('insw_error'):
                err_msg = st.session_state.pop('insw_error')
                if st.session_state.get('insw_results'):
                    st.warning(err_msg + " Hasil parsial ditampilkan di bawah.")
                else:
                    st.error(err_msg)

            insw_results_stored = st.session_state.get('insw_results', [])
            if insw_results_stored:
                df_insw_results = pd.DataFrame(insw_results_stored)

                st.markdown("---")
                st.markdown("### 📊 Hasil Pengecekan INSW")

                insw_impor_count = len(df_insw_results[df_insw_results['Ada Regulasi Impor'] == 'YA'])
                insw_ekspor_count = len(df_insw_results[df_insw_results['Ada Regulasi Ekspor'] == 'YA'])
                insw_obat_count = len(df_insw_results[df_insw_results['Terkait Obat (INSW)'] == 'YA'])
                insw_bpom_count = len(df_insw_results[df_insw_results['Ada BPOM'] == 'YA'])
                insw_both_count = len(df_insw_results[
                    (df_insw_results['Ada Regulasi Impor'] == 'YA') &
                    (df_insw_results['Ada Regulasi Ekspor'] == 'YA')
                ])
                insw_no_lartas = len(df_insw_results[
                    (df_insw_results['Ada Regulasi Impor'] == 'Tidak') &
                    (df_insw_results['Ada Regulasi Ekspor'] == 'Tidak')
                ])

                col_r1, col_r2, col_r3, col_r4, col_r5, col_r6 = st.columns(6)
                with col_r1:
                    st.metric("Total Dicek", len(insw_results_stored))
                with col_r2:
                    st.metric("Regulasi Impor", insw_impor_count)
                with col_r3:
                    st.metric("Regulasi Ekspor", insw_ekspor_count)
                with col_r4:
                    st.metric("Terkait Obat", insw_obat_count)
                with col_r5:
                    st.metric("Ada BPOM", insw_bpom_count)
                with col_r6:
                    st.metric("Tidak Ada Lartas", insw_no_lartas)

                st.markdown("---")

                result_prefixes = sorted(list(set(str(r.get('HS Code', ''))[:2] for r in insw_results_stored if r.get('HS Code', ''))))

                col_f1, col_f2 = st.columns(2)
                with col_f1:
                    filter_hs_prefix = st.multiselect(
                        "🔎 Filter per Chapter:",
                        options=result_prefixes,
                        default=[],
                        key="insw_filter_prefix",
                        help="Filter berdasarkan chapter. Kosongkan untuk menampilkan semua."
                    )
                with col_f2:
                    filter_insw_type = st.multiselect(
                        "🔎 Filter berdasarkan hasil:",
                        options=["Ada Regulasi Impor", "Ada Regulasi Ekspor", "Impor & Ekspor", "Terkait Obat", "Ada BPOM", "Tidak Ada Lartas"],
                        default=[],
                        key="insw_filter_type",
                        help="Filter berdasarkan jenis regulasi. Kosongkan untuk menampilkan semua."
                    )

                df_insw_display = df_insw_results

                if filter_hs_prefix:
                    df_insw_display = df_insw_display[df_insw_display['HS Code'].astype(str).str[:2].isin(filter_hs_prefix)]

                if filter_insw_type:
                    mask = pd.Series([False] * len(df_insw_display), index=df_insw_display.index)
                    if "Ada Regulasi Impor" in filter_insw_type:
                        mask = mask | (df_insw_display['Ada Regulasi Impor'] == 'YA')
                    if "Ada Regulasi Ekspor" in filter_insw_type:
                        mask = mask | (df_insw_display['Ada Regulasi Ekspor'] == 'YA')
                    if "Impor & Ekspor" in filter_insw_type:
                        mask = mask | ((df_insw_display['Ada Regulasi Impor'] == 'YA') & (df_insw_display['Ada Regulasi Ekspor'] == 'YA'))
                    if "Terkait Obat" in filter_insw_type:
                        mask = mask | (df_insw_display['Terkait Obat (INSW)'] == 'YA')
                    if "Ada BPOM" in filter_insw_type:
                        mask = mask | (df_insw_display['Ada BPOM'] == 'YA')
                    if "Tidak Ada Lartas" in filter_insw_type:
                        mask = mask | ((df_insw_display['Ada Regulasi Impor'] == 'Tidak') & (df_insw_display['Ada Regulasi Ekspor'] == 'Tidak'))
                    df_insw_display = df_insw_display[mask]

                if filter_hs_prefix or filter_insw_type:
                    st.caption(f"Menampilkan {len(df_insw_display)} dari {len(df_insw_results)} HS Code")

                tab_insw_all, tab_insw_impor, tab_insw_ekspor, tab_insw_obat, tab_insw_bpom = st.tabs(
                    ["📋 Semua", "📦 Regulasi Impor", "🚢 Regulasi Ekspor", "💊 Terkait Obat", "🏥 Ada BPOM"]
                )

                def highlight_insw(row):
                    jenis = str(row.get('Jenis', ''))
                    if row.get('Terkait Obat (INSW)') == 'YA':
                        return ['background-color: #dcfce7'] * len(row)
                    elif jenis == 'IMPOR & EKSPOR':
                        return ['background-color: #fce7f3'] * len(row)
                    elif row.get('Ada Regulasi Ekspor') == 'YA':
                        return ['background-color: #fef3c7'] * len(row)
                    elif row.get('Ada Regulasi Impor') == 'YA':
                        return ['background-color: #dbeafe'] * len(row)
                    elif 'Error' in str(row.get('Keterangan Impor', '')):
                        return ['background-color: #fef2f2'] * len(row)
                    return [''] * len(row)

                with tab_insw_all:
                    if len(df_insw_display) > 0:
                        st.markdown("**Legenda warna:** 🟢 Terkait Obat | 🩷 Impor & Ekspor | 🔵 Impor | 🟡 Ekspor | ⬜ Tidak ada lartas")
                        styled_insw = df_insw_display.style.apply(highlight_insw, axis=1)
                        st.dataframe(styled_insw, use_container_width=True, height=400)
                    else:
                        st.info("Tidak ada HS Code yang cocok dengan filter yang dipilih.")

                with tab_insw_impor:
                    df_insw_imp = df_insw_display[df_insw_display['Ada Regulasi Impor'] == 'YA']
                    if len(df_insw_imp) > 0:
                        st.success(f"**{len(df_insw_imp)}** HS Code memiliki regulasi impor")
                        impor_cols = ['HS Code', 'Deskripsi', 'Jenis', 'Lartas Border', 'Tata Niaga Post Border', 'Komoditi INSW', 'Ada BPOM', 'Terkait Obat (INSW)', 'Keterangan Impor']
                        display_cols = [c for c in impor_cols if c in df_insw_imp.columns]
                        st.dataframe(df_insw_imp[display_cols], use_container_width=True, height=400)
                    else:
                        st.info("Tidak ada HS Code yang memiliki regulasi impor")

                with tab_insw_ekspor:
                    df_insw_eks = df_insw_display[df_insw_display['Ada Regulasi Ekspor'] == 'YA']
                    if len(df_insw_eks) > 0:
                        st.success(f"**{len(df_insw_eks)}** HS Code memiliki regulasi ekspor")
                        ekspor_cols = ['HS Code', 'Deskripsi', 'Jenis', 'Lartas Ekspor', 'Komoditi INSW', 'Keterangan Ekspor']
                        display_cols = [c for c in ekspor_cols if c in df_insw_eks.columns]
                        st.dataframe(df_insw_eks[display_cols], use_container_width=True, height=400)
                    else:
                        st.info("Tidak ada HS Code yang memiliki regulasi ekspor")

                with tab_insw_obat:
                    df_insw_obat_data = df_insw_display[df_insw_display['Terkait Obat (INSW)'] == 'YA']
                    if len(df_insw_obat_data) > 0:
                        st.success(f"**{len(df_insw_obat_data)}** HS Code terkait obat/farmasi")
                        obat_cols = ['HS Code', 'Deskripsi', 'Jenis', 'Komoditi INSW', 'Ada BPOM', 'Keterangan Impor', 'Keterangan Ekspor']
                        display_cols = [c for c in obat_cols if c in df_insw_obat_data.columns]
                        st.dataframe(df_insw_obat_data[display_cols], use_container_width=True, height=400)
                    else:
                        st.info("Tidak ada HS Code yang terkait obat menurut INSW")

                with tab_insw_bpom:
                    df_insw_bpom_data = df_insw_display[df_insw_display['Ada BPOM'] == 'YA']
                    if len(df_insw_bpom_data) > 0:
                        st.success(f"**{len(df_insw_bpom_data)}** HS Code memiliki regulasi BPOM")
                        bpom_cols = ['HS Code', 'Deskripsi', 'Jenis', 'Komoditi INSW', 'Ada BPOM', 'Keterangan Impor']
                        display_cols = [c for c in bpom_cols if c in df_insw_bpom_data.columns]
                        st.dataframe(df_insw_bpom_data[display_cols], use_container_width=True, height=400)
                    else:
                        st.info("Tidak ada HS Code yang memiliki regulasi BPOM")

                st.markdown("---")
                st.markdown("### 📥 Download Hasil INSW")

                output_insw = io.BytesIO()
                with pd.ExcelWriter(output_insw, engine='openpyxl') as writer:
                    header_fill_insw = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                    header_font_insw = Font(bold=True, color='FFFFFF')
                    green_fill_insw = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
                    yellow_fill_insw = PatternFill(start_color='FFFF00', end_color='FFFF00', fill_type='solid')
                    blue_fill_insw = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
                    pink_fill_insw = PatternFill(start_color='FCE4EC', end_color='FCE4EC', fill_type='solid')
                    thin_border_insw = Border(
                        left=Side(style='thin'), right=Side(style='thin'),
                        top=Side(style='thin'), bottom=Side(style='thin')
                    )

                    df_insw_results.to_excel(writer, index=False, sheet_name='Semua Hasil INSW')
                    ws_insw = writer.sheets['Semua Hasil INSW']

                    for col_idx in range(1, len(df_insw_results.columns) + 1):
                        cell = ws_insw.cell(row=1, column=col_idx)
                        cell.fill = header_fill_insw
                        cell.font = header_font_insw
                        cell.alignment = Alignment(horizontal='center')
                        cell.border = thin_border_insw

                    obat_col = list(df_insw_results.columns).index('Terkait Obat (INSW)') + 1
                    jenis_col = list(df_insw_results.columns).index('Jenis') + 1
                    ekspor_col = list(df_insw_results.columns).index('Ada Regulasi Ekspor') + 1
                    impor_col = list(df_insw_results.columns).index('Ada Regulasi Impor') + 1

                    for row_idx in range(2, len(df_insw_results) + 2):
                        obat_val = ws_insw.cell(row=row_idx, column=obat_col).value
                        jenis_val = ws_insw.cell(row=row_idx, column=jenis_col).value
                        ekspor_val = ws_insw.cell(row=row_idx, column=ekspor_col).value
                        impor_val = ws_insw.cell(row=row_idx, column=impor_col).value
                        for col_idx in range(1, len(df_insw_results.columns) + 1):
                            cell = ws_insw.cell(row=row_idx, column=col_idx)
                            cell.border = thin_border_insw
                            if obat_val == 'YA':
                                cell.fill = green_fill_insw
                            elif jenis_val == 'IMPOR & EKSPOR':
                                cell.fill = pink_fill_insw
                            elif ekspor_val == 'YA':
                                cell.fill = yellow_fill_insw
                            elif impor_val == 'YA':
                                cell.fill = blue_fill_insw

                    for col_idx, col in enumerate(df_insw_results.columns, 1):
                        max_len = max(df_insw_results[col].astype(str).apply(len).max(), len(str(col))) + 2
                        ws_insw.column_dimensions[ws_insw.cell(row=1, column=col_idx).column_letter].width = min(max_len, 60)

                    df_impor_only = df_insw_results[df_insw_results['Ada Regulasi Impor'] == 'YA'].copy()
                    if len(df_impor_only) > 0:
                        df_impor_only.to_excel(writer, index=False, sheet_name='Regulasi Impor')
                        ws_imp = writer.sheets['Regulasi Impor']
                        for col_idx in range(1, len(df_impor_only.columns) + 1):
                            cell = ws_imp.cell(row=1, column=col_idx)
                            cell.fill = header_fill_insw
                            cell.font = header_font_insw
                            cell.alignment = Alignment(horizontal='center')
                            cell.border = thin_border_insw
                        for row_idx in range(2, len(df_impor_only) + 2):
                            for col_idx in range(1, len(df_impor_only.columns) + 1):
                                ws_imp.cell(row=row_idx, column=col_idx).border = thin_border_insw
                                ws_imp.cell(row=row_idx, column=col_idx).fill = blue_fill_insw

                    df_ekspor_only = df_insw_results[df_insw_results['Ada Regulasi Ekspor'] == 'YA'].copy()
                    if len(df_ekspor_only) > 0:
                        df_ekspor_only.to_excel(writer, index=False, sheet_name='Regulasi Ekspor')
                        ws_eks = writer.sheets['Regulasi Ekspor']
                        for col_idx in range(1, len(df_ekspor_only.columns) + 1):
                            cell = ws_eks.cell(row=1, column=col_idx)
                            cell.fill = header_fill_insw
                            cell.font = header_font_insw
                            cell.alignment = Alignment(horizontal='center')
                            cell.border = thin_border_insw
                        for row_idx in range(2, len(df_ekspor_only) + 2):
                            for col_idx in range(1, len(df_ekspor_only.columns) + 1):
                                ws_eks.cell(row=row_idx, column=col_idx).border = thin_border_insw
                                ws_eks.cell(row=row_idx, column=col_idx).fill = yellow_fill_insw

                    df_obat_only = df_insw_results[df_insw_results['Terkait Obat (INSW)'] == 'YA'].copy()
                    if len(df_obat_only) > 0:
                        df_obat_only.to_excel(writer, index=False, sheet_name='Terkait Obat')
                        ws_obat = writer.sheets['Terkait Obat']
                        for col_idx in range(1, len(df_obat_only.columns) + 1):
                            cell = ws_obat.cell(row=1, column=col_idx)
                            cell.fill = header_fill_insw
                            cell.font = header_font_insw
                            cell.alignment = Alignment(horizontal='center')
                            cell.border = thin_border_insw
                        for row_idx in range(2, len(df_obat_only) + 2):
                            for col_idx in range(1, len(df_obat_only.columns) + 1):
                                ws_obat.cell(row=row_idx, column=col_idx).border = thin_border_insw
                                ws_obat.cell(row=row_idx, column=col_idx).fill = green_fill_insw

                output_insw.seek(0)
                st.download_button(
                    label="📥 Download Hasil Cek INSW (Excel)",
                    data=output_insw,
                    file_name="hasil_cek_insw.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

        except Exception as e:
            st.error(f"Terjadi kesalahan: {str(e)}")
            st.info("Pastikan file dalam format yang benar (.xlsx atau .xls)")
    else:
        st.info("Silakan upload file data BPS untuk memulai pengecekan INSW.")

with tab_analysis:
    st.markdown("### 📈 Analisis Data")
    st.markdown("Upload file Excel untuk menganalisis dan memvisualisasikan data Anda.")
    
    file_analysis = st.file_uploader("📁 Upload file untuk analisis", type=['xlsx', 'xls'], key="analysis")
    
    if file_analysis:
        try:
            df_analysis = pd.read_excel(file_analysis)
            
            st.markdown("---")
            st.markdown("### 📋 Preview Data")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📝 Jumlah Baris", len(df_analysis))
            with col2:
                st.metric("📊 Jumlah Kolom", len(df_analysis.columns))
            with col3:
                st.metric("📁 Ukuran Data", f"{df_analysis.memory_usage(deep=True).sum() / 1024:.1f} KB")
            
            st.dataframe(df_analysis.head(10), use_container_width=True, height=250)
            
            st.markdown("---")
            st.markdown("### ⚙️ Konfigurasi Analisis")
            
            col_list = df_analysis.columns.tolist()
            
            col1, col2 = st.columns(2)
            
            with col1:
                selected_analysis_col = st.selectbox(
                    "📌 Pilih kolom untuk dianalisis",
                    options=col_list,
                    key="analysis_col",
                    help="Pilih kolom yang ingin Anda analisis (misalnya: Negara, Jenis Obat, dll)"
                )
            
            with col2:
                top_n = st.slider("🔢 Tampilkan Top N data", min_value=5, max_value=50, value=10, key="top_n")
            
            col1, col2, col3 = st.columns([1, 2, 1])
            with col2:
                analysis_btn = st.button("🔍 Analisis Data", type="primary", use_container_width=True, key="btn_analysis")
            
            if analysis_btn:
                st.markdown("---")
                
                value_counts = df_analysis[selected_analysis_col].value_counts().head(top_n)
                
                st.markdown(f"### 📊 Top {top_n} {selected_analysis_col}")
                
                total_data = len(df_analysis)
                unique_values = df_analysis[selected_analysis_col].nunique()
                top_value = value_counts.index[0] if len(value_counts) > 0 else '-'
                top_count = value_counts.values[0] if len(value_counts) > 0 else 0
                
                col1, col2, col3, col4 = st.columns(4)
                
                with col1:
                    st.metric("📝 Total Data", total_data)
                with col2:
                    st.metric("🔢 Nilai Unik", unique_values)
                with col3:
                    st.metric("🏆 Terbanyak", str(top_value)[:20])
                with col4:
                    st.metric("📊 Jumlah", top_count)
                
                st.markdown("---")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    st.markdown("#### 📋 Tabel Data")
                    df_counts = value_counts.reset_index()
                    df_counts.columns = [selected_analysis_col, 'Jumlah']
                    df_counts['Persentase'] = (df_counts['Jumlah'] / df_counts['Jumlah'].sum() * 100).round(2).astype(str) + '%'
                    df_counts.index = range(1, len(df_counts) + 1)
                    st.dataframe(df_counts, use_container_width=True, height=400)
                
                with col2:
                    st.markdown("#### 📊 Grafik Bar")
                    st.bar_chart(value_counts, use_container_width=True, height=400)
                
                st.markdown("---")
                st.markdown("#### 🥧 Grafik Pie")
                
                import matplotlib.pyplot as plt
                
                fig, ax = plt.subplots(figsize=(12, 8))
                colors = plt.cm.Set3(range(len(value_counts)))
                
                wedges, texts, autotexts = ax.pie(
                    value_counts.values, 
                    labels=None,
                    autopct='%1.1f%%',
                    colors=colors,
                    startangle=90,
                    explode=[0.02] * len(value_counts)
                )
                
                for autotext in autotexts:
                    autotext.set_fontsize(9)
                    autotext.set_fontweight('bold')
                
                ax.legend(
                    wedges, 
                    [f"{str(label)[:30]} ({count:,})" for label, count in zip(value_counts.index, value_counts.values)],
                    title=selected_analysis_col,
                    loc="center left",
                    bbox_to_anchor=(1, 0, 0.5, 1),
                    fontsize=9
                )
                
                ax.set_title(f"Distribusi {selected_analysis_col}", fontsize=14, fontweight='bold', pad=20)
                plt.tight_layout()
                
                st.pyplot(fig)
                
                st.markdown("---")
                st.markdown("### 📥 Download Hasil Analisis")
                
                col1, col2 = st.columns(2)
                
                with col1:
                    output_analysis = io.BytesIO()
                    with pd.ExcelWriter(output_analysis, engine='openpyxl') as writer:
                        df_counts.to_excel(writer, index=True, sheet_name='Hasil Analisis')
                    output_analysis.seek(0)
                    
                    st.download_button(
                        label="📥 Download Data (Excel)",
                        data=output_analysis,
                        file_name=f"analisis_{selected_analysis_col}.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )
                
                with col2:
                    img_buffer = io.BytesIO()
                    fig.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', facecolor='white')
                    img_buffer.seek(0)
                    
                    st.download_button(
                        label="📥 Download Grafik (PNG)",
                        data=img_buffer,
                        file_name=f"grafik_{selected_analysis_col}.png",
                        mime="image/png",
                        use_container_width=True
                    )
                
        except Exception as e:
            st.error(f"❌ Terjadi kesalahan: {str(e)}")
            st.info("💡 Pastikan file Excel dalam format yang benar (.xlsx atau .xls)")
    else:
        st.info("👆 Silakan upload file Excel untuk memulai analisis data.")

with tab_petugas:
    st.markdown("### 👤 Cek & Lengkapi Nama Petugas Loket S2")
    st.markdown("Otomatis melengkapi **Nama Petugas** yang kosong di data Loket S2 berdasarkan data **Form Konsultasi**, serta mengisi **Skor** berdasarkan tingkat kepuasan.")

    col_up1, col_up2 = st.columns(2)
    with col_up1:
        st.markdown("**📁 File Loket S2**")
        file_loket = st.file_uploader("Upload file Loket S2 (.xlsx)", type=["xlsx", "xls"], key="loket_s2_file")
    with col_up2:
        st.markdown("**📁 File Form Konsultasi (bisa lebih dari 1)**")
        files_form = st.file_uploader("Upload file Form Konsultasi (.xlsx)", type=["xlsx", "xls"], key="form_konsul_files", accept_multiple_files=True)

    if 'petugas_result' not in st.session_state:
        st.session_state.petugas_result = None
    if 'petugas_excel' not in st.session_state:
        st.session_state.petugas_excel = None

    if file_loket and files_form:
        if st.button("🔍 Proses & Lengkapi Data Petugas", key="btn_cek_petugas"):
            with st.spinner("Memproses data..."):
                try:
                    df_loket_raw = pd.read_excel(file_loket)

                    skor_map = {'Sangat Puas': 2, 'Puas': 1, 'Tidak Puas': 0}

                    loket_records = []
                    current_date = None
                    current_satisfaction = None
                    idx_loket = 0
                    while idx_loket < len(df_loket_raw):
                        val0 = str(df_loket_raw.iloc[idx_loket, 0]).strip() if pd.notna(df_loket_raw.iloc[idx_loket, 0]) else ''
                        val1 = str(df_loket_raw.iloc[idx_loket, 1]).strip() if pd.notna(df_loket_raw.iloc[idx_loket, 1]) else ''

                        if val0 in ['', 'Row Labels', 'Grand Total'] or 'Nama petugas' in val1:
                            idx_loket += 1
                            continue

                        if 'Sangat Puas' in val0:
                            current_satisfaction = 'Sangat Puas'
                            idx_loket += 1
                            continue
                        elif 'Tidak Puas' in val0:
                            current_satisfaction = 'Tidak Puas'
                            idx_loket += 1
                            continue
                        elif 'Puas' in val0 and 'Sangat' not in val0 and 'Tidak' not in val0:
                            current_satisfaction = 'Puas'
                            idx_loket += 1
                            continue

                        try:
                            date_val = pd.to_datetime(val0)
                            if date_val.year >= 2025:
                                current_date = date_val
                                idx_loket += 1
                                continue
                        except (ValueError, TypeError):
                            pass

                        if '@' in val0:
                            idx_loket += 1
                            continue

                        if idx_loket + 1 < len(df_loket_raw):
                            next_val = str(df_loket_raw.iloc[idx_loket + 1, 0]).strip() if pd.notna(df_loket_raw.iloc[idx_loket + 1, 0]) else ''
                            if '@' in next_val:
                                nama = val0
                                petugas_loket = val1 if val1 else ''
                                email = next_val.lower()
                                skor_otomatis = skor_map.get(current_satisfaction, '')
                                loket_records.append({
                                    'Tanggal': current_date,
                                    'Nama': nama,
                                    'Email': email,
                                    'Petugas_Loket': petugas_loket,
                                    'Kepuasan': current_satisfaction if current_satisfaction else '',
                                    'Skor': skor_otomatis
                                })
                                idx_loket += 2
                                continue

                        idx_loket += 1

                    df_loket = pd.DataFrame(loket_records)

                    form_records = []
                    for ff in files_form:
                        df_f = pd.read_excel(ff)
                        has_loket_col = 'Pilihan Loket Layanan' in df_f.columns
                        if has_loket_col:
                            df_f = df_f[df_f['Pilihan Loket Layanan'].astype(str).str.contains('S2', case=False, na=False)]

                        col_nama = 'Nama' if 'Nama' in df_f.columns else None
                        col_email = 'Email Address' if 'Email Address' in df_f.columns else None
                        col_tanggal = 'Tanggal Konsultasi' if 'Tanggal Konsultasi' in df_f.columns else None
                        col_petugas = 'Nama Petugas' if 'Nama Petugas' in df_f.columns else None

                        if not all([col_nama, col_email, col_petugas]):
                            continue

                        for _, row in df_f.iterrows():
                            f_nama = str(row[col_nama]).strip() if pd.notna(row[col_nama]) else ''
                            f_email = str(row[col_email]).strip().lower() if pd.notna(row[col_email]) else ''
                            f_tanggal = None
                            if col_tanggal and pd.notna(row[col_tanggal]):
                                try:
                                    f_tanggal = pd.to_datetime(row[col_tanggal])
                                except (ValueError, TypeError):
                                    pass
                            f_petugas = str(row[col_petugas]).strip() if pd.notna(row[col_petugas]) else ''

                            if f_nama and f_email:
                                form_records.append({
                                    'Nama_Form': f_nama,
                                    'Email_Form': f_email,
                                    'Tanggal_Form': f_tanggal,
                                    'Petugas_Form': f_petugas,
                                    'Sumber': ff.name
                                })

                    df_forms = pd.DataFrame(form_records)

                    if df_loket.empty:
                        st.session_state.petugas_result = None
                        st.error("❌ Tidak ada data yang berhasil diparsing dari file Loket S2.")
                    elif df_forms.empty:
                        st.session_state.petugas_result = None
                        st.error("❌ Tidak ada data Form Konsultasi yang ditemukan.")
                    else:
                        def normalize_name(name):
                            if not name:
                                return ''
                            return re.sub(r'\s+', ' ', str(name).strip().lower())

                        def find_form_petugas(email, tanggal, nama):
                            candidates = df_forms[df_forms['Email_Form'] == email]
                            if not candidates.empty and tanggal is not None:
                                for _, fr in candidates.iterrows():
                                    if fr['Tanggal_Form'] is not None:
                                        try:
                                            if pd.to_datetime(tanggal).date() == pd.to_datetime(fr['Tanggal_Form']).date():
                                                return fr['Petugas_Form'], fr['Sumber']
                                        except (ValueError, TypeError):
                                            pass
                            if not candidates.empty:
                                first = candidates.iloc[0]
                                if first['Petugas_Form']:
                                    return first['Petugas_Form'], first['Sumber']
                            norm_nama = normalize_name(nama)
                            if norm_nama:
                                for _, fr in df_forms.iterrows():
                                    if normalize_name(fr['Nama_Form']) == norm_nama:
                                        if tanggal is not None and fr['Tanggal_Form'] is not None:
                                            try:
                                                if pd.to_datetime(tanggal).date() == pd.to_datetime(fr['Tanggal_Form']).date():
                                                    return fr['Petugas_Form'], fr['Sumber']
                                            except (ValueError, TypeError):
                                                pass
                            return '', ''

                        def match_short_to_full(short_name, full_name):
                            sn = normalize_name(short_name)
                            fn = normalize_name(full_name)
                            if not sn or not fn:
                                return False
                            if sn == fn:
                                return True
                            if sn in fn or fn in sn:
                                return True
                            sn_parts = sn.split()
                            fn_parts = fn.split()
                            for p in sn_parts:
                                if len(p) > 2 and p in fn_parts:
                                    return True
                            return False

                        results = []
                        for _, lr in df_loket.iterrows():
                            petugas_loket = lr['Petugas_Loket']
                            form_petugas, sumber = find_form_petugas(lr['Email'], lr['Tanggal'], lr['Nama'])

                            if petugas_loket and form_petugas:
                                if match_short_to_full(petugas_loket, form_petugas):
                                    status = 'Cocok'
                                    petugas_final = form_petugas
                                else:
                                    status = 'Tidak Cocok'
                                    petugas_final = petugas_loket
                            elif petugas_loket and not form_petugas:
                                status = 'Tidak Ada di Form'
                                petugas_final = petugas_loket
                            elif not petugas_loket and form_petugas:
                                status = 'Otomatis Terisi'
                                petugas_final = form_petugas
                            else:
                                status = 'Kosong'
                                petugas_final = ''

                            tanggal_str = ''
                            if lr['Tanggal'] is not None:
                                try:
                                    tanggal_str = pd.to_datetime(lr['Tanggal']).strftime('%d-%m-%Y')
                                except Exception:
                                    tanggal_str = str(lr['Tanggal'])

                            results.append({
                                'Tanggal': tanggal_str,
                                'Nama': lr['Nama'],
                                'Email': lr['Email'],
                                'Kepuasan': lr['Kepuasan'],
                                'Skor': lr['Skor'],
                                'Petugas (Loket S2)': petugas_loket if petugas_loket else '-',
                                'Petugas (Form)': form_petugas if form_petugas else '-',
                                'Petugas Final': petugas_final if petugas_final else '-',
                                'Status': status,
                                'Sumber File': sumber if sumber else '-'
                            })

                        df_result = pd.DataFrame(results)
                        st.session_state.petugas_result = df_result

                        out_buf = io.BytesIO()
                        with pd.ExcelWriter(out_buf, engine='openpyxl') as writer:
                            df_result.to_excel(writer, index=False, sheet_name='Hasil Lengkap')

                            wb = writer.book
                            ws = wb['Hasil Lengkap']

                            green_fill = PatternFill(start_color='D4EDDA', end_color='D4EDDA', fill_type='solid')
                            blue_fill = PatternFill(start_color='CCE5FF', end_color='CCE5FF', fill_type='solid')
                            red_fill = PatternFill(start_color='F8D7DA', end_color='F8D7DA', fill_type='solid')
                            yellow_fill = PatternFill(start_color='FFF3CD', end_color='FFF3CD', fill_type='solid')
                            gray_fill = PatternFill(start_color='E2E3E5', end_color='E2E3E5', fill_type='solid')
                            hdr_fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                            hdr_font = Font(bold=True, color='FFFFFF')
                            border_thin = Border(
                                left=Side(style='thin'), right=Side(style='thin'),
                                top=Side(style='thin'), bottom=Side(style='thin')
                            )

                            for cell in ws[1]:
                                cell.fill = hdr_fill
                                cell.font = hdr_font
                                cell.alignment = Alignment(horizontal='center', vertical='center')
                                cell.border = border_thin

                            status_ci = list(df_result.columns).index('Status') + 1
                            fill_map = {
                                'Cocok': green_fill, 'Otomatis Terisi': blue_fill,
                                'Tidak Cocok': red_fill, 'Kosong': yellow_fill,
                                'Tidak Ada di Form': gray_fill
                            }
                            for ri in range(2, ws.max_row + 1):
                                sv = ws.cell(row=ri, column=status_ci).value or ''
                                fl = fill_map.get(sv)
                                for ci_x in range(1, ws.max_column + 1):
                                    c = ws.cell(row=ri, column=ci_x)
                                    c.border = border_thin
                                    if fl:
                                        c.fill = fl

                            for ci_x in range(1, ws.max_column + 1):
                                ml = 0
                                for ri in range(1, ws.max_row + 1):
                                    cv = ws.cell(row=ri, column=ci_x).value
                                    if cv:
                                        ml = max(ml, len(str(cv)))
                                ws.column_dimensions[ws.cell(row=1, column=ci_x).column_letter].width = min(ml + 3, 40)

                            for sheet_status, sheet_name in [('Cocok', 'Cocok'), ('Otomatis Terisi', 'Otomatis Terisi'), ('Tidak Cocok', 'Tidak Cocok'), ('Kosong', 'Petugas Kosong'), ('Tidak Ada di Form', 'Tidak Ada di Form')]:
                                df_sheet = df_result[df_result['Status'] == sheet_status]
                                if not df_sheet.empty:
                                    df_sheet.to_excel(writer, index=False, sheet_name=sheet_name)

                        out_buf.seek(0)
                        st.session_state.petugas_excel = out_buf.getvalue()
                        st.rerun()

                except Exception as e:
                    st.error(f"❌ Terjadi kesalahan: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())

    if st.session_state.petugas_result is not None:
        df_result = st.session_state.petugas_result

        total = len(df_result)
        cocok = len(df_result[df_result['Status'] == 'Cocok'])
        otomatis = len(df_result[df_result['Status'] == 'Otomatis Terisi'])
        tidak_cocok = len(df_result[df_result['Status'] == 'Tidak Cocok'])
        kosong = len(df_result[df_result['Status'] == 'Kosong'])
        no_form = len(df_result[df_result['Status'] == 'Tidak Ada di Form'])

        st.success(f"✅ Berhasil memproses {total} data Loket S2")

        c1, c2, c3, c4, c5, c6 = st.columns(6)
        c1.metric("Total", total)
        c2.metric("Cocok", cocok)
        c3.metric("Otomatis Terisi", otomatis)
        c4.metric("Tidak Cocok", tidak_cocok)
        c5.metric("Kosong", kosong)
        c6.metric("Tidak Ada di Form", no_form)

        filter_st = st.selectbox("Filter Status:", ["Semua", "Cocok", "Otomatis Terisi", "Tidak Cocok", "Kosong", "Tidak Ada di Form"], key="filter_petugas")
        df_show = df_result if filter_st == "Semua" else df_result[df_result['Status'] == filter_st]

        def color_row(row):
            s = row['Status']
            if s == 'Cocok':
                return ['background-color: #d4edda'] * len(row)
            elif s == 'Otomatis Terisi':
                return ['background-color: #cce5ff'] * len(row)
            elif s == 'Tidak Cocok':
                return ['background-color: #f8d7da'] * len(row)
            elif s == 'Kosong':
                return ['background-color: #fff3cd'] * len(row)
            elif s == 'Tidak Ada di Form':
                return ['background-color: #e2e3e5'] * len(row)
            return [''] * len(row)

        st.dataframe(df_show.style.apply(color_row, axis=1), height=500)

        if st.session_state.petugas_excel is not None:
            st.download_button(
                label="📥 Download Hasil Pengecekan Petugas (Excel)",
                data=st.session_state.petugas_excel,
                file_name="hasil_cek_petugas_loket_s2.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    elif not file_loket or not files_form:
        st.info("👆 Silakan upload file **Loket S2** dan **Form Konsultasi** untuk memulai pengecekan petugas.")

with tab_absen:
    st.markdown("### 📋 Cek Kehadiran Pegawai")
    st.markdown("Bandingkan data **Pegawai** dengan **Daftar Hadir** untuk mengetahui siapa saja yang **tidak hadir**.")

    col_ab1, col_ab2 = st.columns(2)
    with col_ab1:
        st.markdown("**📁 File Data Pegawai**")
        file_pegawai = st.file_uploader("Upload file Pegawai (.xlsx)", type=["xlsx", "xls"], key="file_pegawai")
    with col_ab2:
        st.markdown("**📁 File Daftar Hadir**")
        file_hadir = st.file_uploader("Upload file Daftar Hadir (.xlsx)", type=["xlsx", "xls"], key="file_hadir")

    if 'absen_result' not in st.session_state:
        st.session_state.absen_result = None
    if 'absen_excel' not in st.session_state:
        st.session_state.absen_excel = None

    if file_pegawai and file_hadir:
        if st.button("🔍 Cek Kehadiran", key="btn_cek_absen"):
            with st.spinner("Memproses data..."):
                try:
                    df_peg_raw = pd.read_excel(file_pegawai)
                    df_hadir_raw = pd.read_excel(file_hadir)

                    pegawai_list = []
                    nama_col_peg = None
                    jabatan_col_peg = None
                    for ci in range(min(10, df_peg_raw.shape[1])):
                        for ri in range(min(10, len(df_peg_raw))):
                            val = str(df_peg_raw.iloc[ri, ci]).strip().upper() if pd.notna(df_peg_raw.iloc[ri, ci]) else ''
                            if val == 'NAMA':
                                nama_col_peg = ci
                                start_row_peg = ri + 1
                                if ci + 1 < df_peg_raw.shape[1]:
                                    jabatan_col_peg = ci + 1
                                break
                        if nama_col_peg is not None:
                            break

                    if nama_col_peg is None:
                        for ci in range(min(10, df_peg_raw.shape[1])):
                            for ri in range(min(10, len(df_peg_raw))):
                                val = str(df_peg_raw.iloc[ri, ci]).strip().lower() if pd.notna(df_peg_raw.iloc[ri, ci]) else ''
                                if 'nama' in val:
                                    nama_col_peg = ci
                                    start_row_peg = ri + 1
                                    break
                            if nama_col_peg is not None:
                                break

                    if nama_col_peg is None:
                        nama_col_peg = 2
                        jabatan_col_peg = 3
                        start_row_peg = 4

                    for ri in range(start_row_peg, len(df_peg_raw)):
                        nama = str(df_peg_raw.iloc[ri, nama_col_peg]).strip() if pd.notna(df_peg_raw.iloc[ri, nama_col_peg]) else ''
                        jabatan = ''
                        if jabatan_col_peg is not None and jabatan_col_peg < df_peg_raw.shape[1]:
                            jabatan = str(df_peg_raw.iloc[ri, jabatan_col_peg]).strip() if pd.notna(df_peg_raw.iloc[ri, jabatan_col_peg]) else ''
                        if nama and nama != 'nan':
                            pegawai_list.append({'Nama': nama, 'Jabatan': jabatan if jabatan != 'nan' else ''})

                    hadir_list = []
                    nama_col_h = None
                    kehadiran_col_h = None
                    waktu_col_h = None
                    for ci in range(min(10, df_hadir_raw.shape[1])):
                        for ri in range(min(10, len(df_hadir_raw))):
                            val = str(df_hadir_raw.iloc[ri, ci]).strip().lower() if pd.notna(df_hadir_raw.iloc[ri, ci]) else ''
                            if 'nama' in val:
                                nama_col_h = ci
                                start_row_h = ri + 1
                                break
                        if nama_col_h is not None:
                            break

                    if nama_col_h is None:
                        nama_col_h = 0
                        start_row_h = 6

                    for ci in range(min(10, df_hadir_raw.shape[1])):
                        for ri in range(min(10, len(df_hadir_raw))):
                            val = str(df_hadir_raw.iloc[ri, ci]).strip().lower() if pd.notna(df_hadir_raw.iloc[ri, ci]) else ''
                            if 'kehadiran' in val:
                                kehadiran_col_h = ci
                            if 'waktu' in val:
                                waktu_col_h = ci

                    for ri in range(start_row_h, len(df_hadir_raw)):
                        nama = str(df_hadir_raw.iloc[ri, nama_col_h]).strip() if pd.notna(df_hadir_raw.iloc[ri, nama_col_h]) else ''
                        kehadiran = ''
                        waktu = ''
                        if kehadiran_col_h is not None:
                            kehadiran = str(df_hadir_raw.iloc[ri, kehadiran_col_h]).strip() if pd.notna(df_hadir_raw.iloc[ri, kehadiran_col_h]) else ''
                        if waktu_col_h is not None:
                            waktu = str(df_hadir_raw.iloc[ri, waktu_col_h]).strip() if pd.notna(df_hadir_raw.iloc[ri, waktu_col_h]) else ''
                        if nama and nama != 'nan':
                            hadir_list.append({'Nama_Hadir': nama, 'Kehadiran': kehadiran if kehadiran != 'nan' else '', 'Waktu': waktu if waktu != 'nan' else ''})

                    def clean_name_absen(name):
                        cleaned = re.sub(r',?\s*(S\.Si|S\.Farm|S\.E|S\.Kom|S\.IP|S\.Ak|S\.Sos|S\.K\.M|SKM|A\.Md|Apt|apt|M\.Si|M\.S|M\.Sc|M\.Farm|M\.Med\.Sc|M\.Epid|M\.K\.M|MKM|M\.T|Dra\.|Drs\.|Dr\.|drg\.|Rr\.)\.*', '', name, flags=re.IGNORECASE)
                        cleaned = re.sub(r'[,.]', ' ', cleaned)
                        cleaned = re.sub(r'\s+', ' ', cleaned).strip()
                        return cleaned.lower()

                    hadir_clean_map = {}
                    for h in hadir_list:
                        hadir_clean_map[clean_name_absen(h['Nama_Hadir'])] = h

                    results_absen = []
                    for p in pegawai_list:
                        p_clean = clean_name_absen(p['Nama'])
                        matched_hadir = None
                        for h_clean, h_data in hadir_clean_map.items():
                            if p_clean == h_clean:
                                matched_hadir = h_data
                                break
                            if p_clean in h_clean or h_clean in p_clean:
                                matched_hadir = h_data
                                break
                            p_parts = p_clean.split()
                            h_parts = h_clean.split()
                            if len(p_parts) >= 2 and len(h_parts) >= 2:
                                if p_parts[0] == h_parts[0] and p_parts[-1] == h_parts[-1]:
                                    matched_hadir = h_data
                                    break

                        if matched_hadir:
                            results_absen.append({
                                'Nama Pegawai': p['Nama'],
                                'Jabatan': p['Jabatan'],
                                'Status': 'Hadir',
                                'Kehadiran': matched_hadir.get('Kehadiran', ''),
                                'Waktu': matched_hadir.get('Waktu', '')
                            })
                        else:
                            results_absen.append({
                                'Nama Pegawai': p['Nama'],
                                'Jabatan': p['Jabatan'],
                                'Status': 'Tidak Hadir',
                                'Kehadiran': '-',
                                'Waktu': '-'
                            })

                    df_absen = pd.DataFrame(results_absen)
                    st.session_state.absen_result = df_absen

                    out_absen = io.BytesIO()
                    with pd.ExcelWriter(out_absen, engine='openpyxl') as writer:
                        df_absen.to_excel(writer, index=False, sheet_name='Semua Pegawai')

                        wb = writer.book
                        ws = wb['Semua Pegawai']

                        green_f = PatternFill(start_color='D4EDDA', end_color='D4EDDA', fill_type='solid')
                        red_f = PatternFill(start_color='F8D7DA', end_color='F8D7DA', fill_type='solid')
                        hdr_f = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                        hdr_fn = Font(bold=True, color='FFFFFF')
                        bdr = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

                        for cell in ws[1]:
                            cell.fill = hdr_f
                            cell.font = hdr_fn
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            cell.border = bdr

                        sc = list(df_absen.columns).index('Status') + 1
                        for ri in range(2, ws.max_row + 1):
                            sv = ws.cell(row=ri, column=sc).value or ''
                            fl = green_f if sv == 'Hadir' else red_f if sv == 'Tidak Hadir' else None
                            for ci_x in range(1, ws.max_column + 1):
                                cell = ws.cell(row=ri, column=ci_x)
                                cell.border = bdr
                                if fl:
                                    cell.fill = fl

                        for ci_x in range(1, ws.max_column + 1):
                            ml = 0
                            for ri in range(1, ws.max_row + 1):
                                cv = ws.cell(row=ri, column=ci_x).value
                                if cv:
                                    ml = max(ml, len(str(cv)))
                            ws.column_dimensions[ws.cell(row=1, column=ci_x).column_letter].width = min(ml + 3, 50)

                        df_tidak = df_absen[df_absen['Status'] == 'Tidak Hadir']
                        if not df_tidak.empty:
                            df_tidak.to_excel(writer, index=False, sheet_name='Tidak Hadir')

                    out_absen.seek(0)
                    st.session_state.absen_excel = out_absen.getvalue()
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Terjadi kesalahan: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())

    if st.session_state.absen_result is not None:
        df_absen = st.session_state.absen_result

        total_peg = len(df_absen)
        total_hadir = len(df_absen[df_absen['Status'] == 'Hadir'])
        total_tidak = len(df_absen[df_absen['Status'] == 'Tidak Hadir'])

        st.success(f"✅ Total Pegawai: {total_peg} | Hadir: {total_hadir} | Tidak Hadir: {total_tidak}")

        c_a1, c_a2, c_a3 = st.columns(3)
        c_a1.metric("Total Pegawai", total_peg)
        c_a2.metric("Hadir", total_hadir)
        c_a3.metric("Tidak Hadir", total_tidak)

        filter_absen = st.selectbox("Filter:", ["Semua", "Hadir", "Tidak Hadir"], key="filter_absen")
        df_show_absen = df_absen if filter_absen == "Semua" else df_absen[df_absen['Status'] == filter_absen]

        def color_absen(row):
            if row['Status'] == 'Hadir':
                return ['background-color: #d4edda'] * len(row)
            elif row['Status'] == 'Tidak Hadir':
                return ['background-color: #f8d7da'] * len(row)
            return [''] * len(row)

        st.dataframe(df_show_absen.style.apply(color_absen, axis=1), height=500)

        if st.session_state.absen_excel is not None:
            st.download_button(
                label="📥 Download Hasil Cek Kehadiran (Excel)",
                data=st.session_state.absen_excel,
                file_name="hasil_cek_kehadiran.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    elif not file_pegawai or not file_hadir:
        st.info("👆 Silakan upload file **Data Pegawai** dan **Daftar Hadir** untuk memulai pengecekan.")

with tab_importir:
    st.markdown("### 🏢 Analisis Bidang Usaha Importir")
    st.markdown("Upload file Excel, pilih kolom **Nama Importir**, lalu sistem akan **otomatis menganalisis bidang usaha** setiap importir menggunakan AI dan menentukan apakah termasuk **CEK** (obat/kosmetik/OT/food) atau **NOM** (bukan komoditas BPOM).")

    file_importir = st.file_uploader("Upload file Excel (.xlsx)", type=["xlsx", "xls"], key="file_importir_upload")

    if 'importir_df_raw' not in st.session_state:
        st.session_state.importir_df_raw = None
    if 'importir_headers' not in st.session_state:
        st.session_state.importir_headers = None
    if 'importir_header_row' not in st.session_state:
        st.session_state.importir_header_row = None
    if 'importir_result' not in st.session_state:
        st.session_state.importir_result = None
    if 'importir_excel' not in st.session_state:
        st.session_state.importir_excel = None
    if 'importir_progress' not in st.session_state:
        st.session_state.importir_progress = None
    if 'importir_file_id' not in st.session_state:
        st.session_state.importir_file_id = None

    if file_importir:
        current_file_id = f"{file_importir.name}_{file_importir.size}"
        if st.session_state.importir_file_id != current_file_id:
            st.session_state.importir_result = None
            st.session_state.importir_excel = None
            st.session_state.importir_file_id = current_file_id
    elif not file_importir and st.session_state.importir_file_id is not None:
        st.session_state.importir_result = None
        st.session_state.importir_excel = None
        st.session_state.importir_file_id = None

    if file_importir:
        try:
            df_raw_imp = pd.read_excel(file_importir, header=None)
            header_row_imp = 0
            for ri in range(min(10, len(df_raw_imp))):
                row_vals = [str(df_raw_imp.iloc[ri, ci]).strip().upper() if pd.notna(df_raw_imp.iloc[ri, ci]) else '' for ci in range(min(10, df_raw_imp.shape[1]))]
                if any('NAMA' in v for v in row_vals):
                    header_row_imp = ri
                    break

            df_with_header = pd.read_excel(file_importir, header=header_row_imp)
            st.session_state.importir_df_raw = df_with_header
            st.session_state.importir_headers = list(df_with_header.columns)
            st.session_state.importir_header_row = header_row_imp

            st.success(f"✅ File berhasil dibaca! **{len(df_with_header)}** baris data, **{len(df_with_header.columns)}** kolom.")

            with st.expander("📋 Preview Data (10 baris pertama)", expanded=True):
                st.dataframe(df_with_header.head(10), height=300)

            with st.expander("📊 Struktur Kolom"):
                col_info = []
                for i, col in enumerate(df_with_header.columns):
                    col_letter = chr(65 + i) if i < 26 else chr(65 + (i // 26 - 1)) + chr(65 + (i % 26))
                    non_null = df_with_header[col].notna().sum()
                    col_info.append({
                        'Kolom Excel': col_letter,
                        'Nama Kolom': str(col)[:50],
                        'Jumlah Data': non_null,
                        'Kosong': len(df_with_header) - non_null
                    })
                st.dataframe(pd.DataFrame(col_info), height=400)

            st.markdown("---")
            st.markdown("#### ⚙️ Pengaturan Analisis")

            col_names = [str(c) for c in df_with_header.columns]

            nama_col_default = 0
            for i, cn in enumerate(col_names):
                if 'NAMA_IMPORTIR' in cn.upper() or 'NAMA IMPORTIR' in cn.upper():
                    nama_col_default = i
                    break

            selected_nama_col = st.selectbox(
                "Pilih kolom **Nama Importir** yang akan dianalisis:",
                options=col_names,
                index=nama_col_default,
                key="sel_nama_importir"
            )

            keterangan_col_options = ["(Buat kolom baru)"] + col_names
            ket_default = 0
            for i, cn in enumerate(col_names):
                cn_up = cn.upper().strip()
                if 'PENJELASAN' in cn_up or 'KETERANGAN' in cn_up:
                    ket_default = i + 1
                    break
            selected_ket_col = st.selectbox(
                "Pilih kolom untuk **Keterangan (NOM/CEK)**:",
                options=keterangan_col_options,
                index=ket_default,
                key="sel_ket_col"
            )

            bidang_col_options = ["(Buat kolom baru)"] + col_names
            bid_default = 0
            for i, cn in enumerate(col_names):
                cn_up = cn.upper().strip()
                if 'HASIL' in cn_up and 'ANALISIS' in cn_up:
                    bid_default = i + 1
                    break
            selected_bidang_col = st.selectbox(
                "Pilih kolom untuk **Bidang Usaha / Hasil Analisis**:",
                options=bidang_col_options,
                index=bid_default,
                key="sel_bidang_col"
            )

            only_empty = st.checkbox("Hanya analisis baris yang kolom Keterangan-nya masih kosong", value=True, key="only_empty_importir")

            product_col_candidates = {}
            for cn_s in col_names:
                cn_up = cn_s.upper()
                if 'BRGURAI' in cn_up:
                    product_col_candidates['brgurai'] = cn_s
                elif 'NOHS' in cn_up or cn_up == 'NOHS':
                    product_col_candidates['nohs'] = cn_s
                elif 'URAIAN_HS' in cn_up or 'URAIAN HS' in cn_up:
                    product_col_candidates['uraian_hs'] = cn_s
                elif 'ALAMAT' in cn_up:
                    product_col_candidates['alamat'] = cn_s

            unique_importers = df_with_header[selected_nama_col].dropna().unique()
            if only_empty and selected_ket_col != "(Buat kolom baru)":
                mask_empty = df_with_header[selected_ket_col].isna() | (df_with_header[selected_ket_col].astype(str).str.strip() == '')
                unique_importers = df_with_header.loc[mask_empty, selected_nama_col].dropna().unique()

            unique_importers = [str(n).strip() for n in unique_importers if str(n).strip() and str(n).strip().lower() != 'nan']
            unique_importers = list(dict.fromkeys(unique_importers))

            importir_context = {}
            for imp_name in unique_importers:
                rows_imp = df_with_header[df_with_header[selected_nama_col].astype(str).str.strip() == imp_name]
                products = []
                alamat = ''
                for _, row_imp in rows_imp.head(5).iterrows():
                    prod_info = {}
                    if 'brgurai' in product_col_candidates:
                        v = str(row_imp.get(product_col_candidates['brgurai'], '')).strip()
                        if v and v != 'nan':
                            prod_info['barang'] = v[:80]
                    if 'nohs' in product_col_candidates:
                        v = str(row_imp.get(product_col_candidates['nohs'], '')).strip()
                        if v and v != 'nan':
                            prod_info['hs'] = v
                    if 'uraian_hs' in product_col_candidates:
                        v = str(row_imp.get(product_col_candidates['uraian_hs'], '')).strip()
                        if v and v != 'nan':
                            prod_info['uraian'] = v[:80]
                    if prod_info:
                        products.append(prod_info)
                    if not alamat and 'alamat' in product_col_candidates:
                        v = str(row_imp.get(product_col_candidates['alamat'], '')).strip()
                        if v and v != 'nan':
                            alamat = v[:100]
                importir_context[imp_name] = {'products': products, 'alamat': alamat}

            st.info(f"📊 Ditemukan **{len(unique_importers)}** importir unik yang perlu dianalisis.")

            if st.button("🤖 Mulai Analisis Otomatis dengan AI", key="btn_analisis_importir"):
                from openai import OpenAI
                from concurrent.futures import ThreadPoolExecutor, as_completed
                from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception
                import time

                ai_key = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY")
                ai_url = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL")

                if not ai_key or not ai_url:
                    st.error("❌ AI Integration belum dikonfigurasi. Pastikan OpenAI AI Integration sudah terinstall.")
                else:
                    client = OpenAI(api_key=ai_key, base_url=ai_url)

                    def is_rate_limit_error(exception):
                        error_msg = str(exception)
                        return ("429" in error_msg or "RATELIMIT" in error_msg.upper()
                                or "quota" in error_msg.lower() or "rate limit" in error_msg.lower()
                                or (hasattr(exception, "status_code") and exception.status_code == 429))

                    @retry(stop=stop_after_attempt(5), wait=wait_exponential(multiplier=1, min=2, max=60), retry=retry_if_exception(is_rate_limit_error), reraise=True)
                    def classify_batch(names_batch, context_map):
                        entries = []
                        for i, n in enumerate(names_batch):
                            ctx = context_map.get(n, {})
                            entry = f"{i+1}. Importir: {n}"
                            if ctx.get('alamat'):
                                entry += f"\n   Alamat: {ctx['alamat']}"
                            if ctx.get('products'):
                                prods = ctx['products'][:3]
                                prod_strs = []
                                for p in prods:
                                    ps = ""
                                    if p.get('barang'):
                                        ps += p['barang']
                                    if p.get('hs'):
                                        ps += f" (HS: {p['hs']})"
                                    if p.get('uraian'):
                                        ps += f" - {p['uraian']}"
                                    if ps:
                                        prod_strs.append(ps)
                                if prod_strs:
                                    entry += "\n   Produk yang diimpor: " + "; ".join(prod_strs)
                            entries.append(entry)
                        names_text = "\n".join(entries)

                        prompt = f"""Kamu adalah analis perdagangan Indonesia yang sangat ahli dalam mengidentifikasi bidang usaha perusahaan importir dan menganalisis komoditas impor terkait regulasi BPOM.

Untuk setiap importir di bawah ini, berikan:
1. "bidang": Bidang usaha utama perusahaan (singkat, 2-5 kata)
2. "kelas": "CEK" jika bidang usahanya terkait obat/farmasi, kosmetik, obat tradisional/herbal, makanan/minuman/food/pangan, suplemen kesehatan, atau bahan baku untuk produk-produk tersebut. "NOM" jika BUKAN terkait hal-hal tersebut.
3. "alasan": Penjelasan detail (2-3 kalimat) yang mencakup:
   - Deskripsi produk yang diimpor (berdasarkan data barang/HS Code jika tersedia)
   - Kegunaan produk tersebut
   - Mengapa diklasifikasikan NOM atau CEK
   - Informasi tentang importir dan bidang usahanya

Contoh alasan yang baik:
"Thermal grease dengan CAS Number 63148-62-9 merupakan bahan berbasis Polydimethylsiloxane (silicone oil) yang digunakan sebagai pasta penghantar panas (heat transfer compound), dan tidak termasuk bahan obat maupun makanan. Importir produk ini adalah PT Jaya Refrigeration Equipment yaitu perusahaan yang bergerak di bidang perdagangan dan penyediaan peralatan sistem pendingin (refrigeration) serta komponen pendukungnya."

PENTING:
- Perusahaan bahan kimia industri/specialty chemicals yang TIDAK spesifik untuk farmasi/food → NOM
- Perusahaan yang jelas bergerak di farmasi/pharmaceutical → CEK
- Perusahaan food ingredients/flavor/fragrance → CEK
- Jika ragu, klasifikasikan sebagai CEK
- Alasan harus ditulis dalam Bahasa Indonesia yang formal dan jelas
- Gunakan informasi produk (nama barang, HS Code, uraian) untuk memberikan alasan yang spesifik

Jawab HANYA dalam format JSON object dengan key "results" berisi array, contoh:
{{"results": [{{"nama": "PT ABC", "bidang": "Farmasi", "kelas": "CEK", "alasan": "PT ABC merupakan perusahaan farmasi yang mengimpor bahan baku obat..."}}, {{"nama": "PT XYZ", "bidang": "Peralatan Pendingin", "kelas": "NOM", "alasan": "PT XYZ mengimpor komponen refrigerasi yang bukan merupakan komoditas BPOM..."}}]}}

Daftar importir:
{names_text}"""

                        response = client.chat.completions.create(
                            model="gpt-5-mini",
                            messages=[{"role": "user", "content": prompt}],
                            response_format={"type": "json_object"},
                            max_completion_tokens=8192
                        )
                        content = response.choices[0].message.content or "[]"
                        try:
                            parsed = _json.loads(content)
                            if isinstance(parsed, dict):
                                for key in parsed:
                                    if isinstance(parsed[key], list):
                                        return parsed[key]
                                return []
                            return parsed
                        except:
                            return []

                    progress_bar = st.progress(0, text="Memulai analisis...")
                    status_text = st.empty()

                    batch_size = 15
                    batches = [unique_importers[i:i+batch_size] for i in range(0, len(unique_importers), batch_size)]
                    all_results = {}
                    total_batches = len(batches)
                    errors_count = 0

                    for bi, batch in enumerate(batches):
                        progress = (bi + 1) / total_batches
                        progress_bar.progress(progress, text=f"Menganalisis batch {bi+1}/{total_batches} ({len(all_results)}/{len(unique_importers)} importir)...")
                        status_text.text(f"🔄 Sedang memproses: {batch[0][:30]}... s/d {batch[-1][:30]}...")

                        try:
                            batch_ctx = {n: importir_context.get(n, {}) for n in batch}
                            results = classify_batch(batch, batch_ctx)
                            matched_in_batch = set()
                            if isinstance(results, list):
                                for r in results:
                                    if isinstance(r, dict) and 'nama' in r:
                                        rname = r['nama'].strip().upper()
                                        all_results[rname] = {
                                            'bidang': r.get('bidang', ''),
                                            'kelas': r.get('kelas', 'CEK'),
                                            'alasan': r.get('alasan', '')
                                        }
                                        matched_in_batch.add(rname)
                            for name in batch:
                                if name.strip().upper() not in matched_in_batch:
                                    norm_name = re.sub(r'[^A-Z0-9\s]', '', name.strip().upper()).strip()
                                    found = False
                                    for mk in matched_in_batch:
                                        mk_norm = re.sub(r'[^A-Z0-9\s]', '', mk).strip()
                                        if norm_name == mk_norm or (len(norm_name) > 5 and (norm_name in mk_norm or mk_norm in norm_name)):
                                            found = True
                                            break
                                    if not found:
                                        all_results[name.strip().upper()] = {'bidang': 'Perlu cek manual', 'kelas': 'CEK', 'alasan': 'Data importir tidak dapat dianalisis secara otomatis, perlu pengecekan manual.'}
                        except Exception as e:
                            errors_count += 1
                            st.warning(f"⚠️ Error batch {bi+1}: {str(e)[:100]}")
                            for name in batch:
                                all_results[name.strip().upper()] = {'bidang': 'Error - perlu cek manual', 'kelas': 'CEK', 'alasan': 'Terjadi error saat analisis, perlu pengecekan manual.'}

                        if bi < total_batches - 1:
                            time.sleep(0.5)

                    progress_bar.progress(1.0, text="✅ Analisis selesai!")
                    status_text.text(f"✅ Selesai! {len(all_results)} importir dianalisis" + (f" ({errors_count} batch error)" if errors_count else ""))

                    df_result = df_with_header.copy()

                    if selected_ket_col == "(Buat kolom baru)":
                        ket_col_name = "Keterangan_AI"
                        df_result[ket_col_name] = ""
                    else:
                        ket_col_name = selected_ket_col

                    if selected_bidang_col == "(Buat kolom baru)":
                        bidang_col_name = "Bidang_Usaha_AI"
                        df_result[bidang_col_name] = ""
                    else:
                        bidang_col_name = selected_bidang_col

                    alasan_col_name = "Alasan_Analisis"
                    df_result[alasan_col_name] = ""

                    def normalize_imp_name(n):
                        return re.sub(r'[^A-Z0-9\s]', '', n).strip()

                    norm_map = {normalize_imp_name(k): v for k, v in all_results.items()}

                    filled_count = 0
                    for idx in range(len(df_result)):
                        nama_val = str(df_result.at[idx, selected_nama_col]).strip().upper() if pd.notna(df_result.at[idx, selected_nama_col]) else ''
                        if not nama_val or nama_val == 'NAN':
                            continue

                        if only_empty and selected_ket_col != "(Buat kolom baru)":
                            existing = str(df_result.at[idx, ket_col_name]).strip() if pd.notna(df_result.at[idx, ket_col_name]) else ''
                            if existing:
                                continue

                        nama_norm = normalize_imp_name(nama_val)
                        matched = all_results.get(nama_val)
                        if not matched:
                            matched = norm_map.get(nama_norm)
                        if not matched:
                            nama_words = set(nama_val.split())
                            best_score = 0
                            for key, val in all_results.items():
                                key_words = set(key.split())
                                if len(nama_words) >= 2 and len(key_words) >= 2:
                                    common = len(nama_words & key_words)
                                    score = common / max(len(nama_words), len(key_words))
                                    if score > best_score and score >= 0.6:
                                        best_score = score
                                        matched = val

                        if matched:
                            df_result.at[idx, ket_col_name] = matched['kelas']
                            df_result.at[idx, bidang_col_name] = matched['bidang']
                            df_result.at[idx, alasan_col_name] = matched.get('alasan', '')
                            filled_count += 1
                        else:
                            df_result.at[idx, ket_col_name] = 'CEK'
                            df_result.at[idx, bidang_col_name] = 'Perlu cek manual'
                            df_result.at[idx, alasan_col_name] = 'Data importir tidak dapat dianalisis secara otomatis, perlu pengecekan manual.'
                            filled_count += 1

                    st.session_state.importir_result = df_result
                    st.session_state.importir_ket_col = ket_col_name
                    st.session_state.importir_bidang_col = bidang_col_name
                    st.session_state.importir_alasan_col = alasan_col_name

                    out_imp = io.BytesIO()
                    with pd.ExcelWriter(out_imp, engine='openpyxl') as writer:
                        df_result.to_excel(writer, index=False, sheet_name='Data Lengkap')

                        wb = writer.book
                        ws = wb['Data Lengkap']

                        hdr_f_imp = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
                        hdr_fn_imp = Font(bold=True, color='FFFFFF')
                        green_fi = PatternFill(start_color='D4EDDA', end_color='D4EDDA', fill_type='solid')
                        yellow_fi = PatternFill(start_color='FFF3CD', end_color='FFF3CD', fill_type='solid')
                        bdr_i = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))

                        for cell in ws[1]:
                            cell.fill = hdr_f_imp
                            cell.font = hdr_fn_imp
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            cell.border = bdr_i

                        col_names_result = list(df_result.columns)
                        ket_ci = col_names_result.index(ket_col_name) + 1 if ket_col_name in col_names_result else None
                        bid_ci = col_names_result.index(bidang_col_name) + 1 if bidang_col_name in col_names_result else None

                        if ket_ci:
                            for ri in range(2, ws.max_row + 1):
                                sv = str(ws.cell(row=ri, column=ket_ci).value or '').strip().upper()
                                if sv == 'NOM':
                                    ws.cell(row=ri, column=ket_ci).fill = green_fi
                                    if bid_ci:
                                        ws.cell(row=ri, column=bid_ci).fill = green_fi
                                elif sv == 'CEK':
                                    ws.cell(row=ri, column=ket_ci).fill = yellow_fi
                                    if bid_ci:
                                        ws.cell(row=ri, column=bid_ci).fill = yellow_fi
                                ws.cell(row=ri, column=ket_ci).border = bdr_i
                                if bid_ci:
                                    ws.cell(row=ri, column=bid_ci).border = bdr_i

                        summary_data = []
                        for name_upper, info in all_results.items():
                            summary_data.append({
                                'Nama Importir': name_upper,
                                'Bidang Usaha': info['bidang'],
                                'Klasifikasi': info['kelas'],
                                'Alasan': info.get('alasan', '')
                            })
                        df_summary = pd.DataFrame(summary_data)
                        df_summary.to_excel(writer, index=False, sheet_name='Ringkasan Importir')

                        ws2 = wb['Ringkasan Importir']
                        for cell in ws2[1]:
                            cell.fill = hdr_f_imp
                            cell.font = hdr_fn_imp
                            cell.alignment = Alignment(horizontal='center', vertical='center')
                            cell.border = bdr_i

                        kls_ci = 3
                        for ri in range(2, ws2.max_row + 1):
                            sv = str(ws2.cell(row=ri, column=kls_ci).value or '').strip().upper()
                            for ci_x in range(1, ws2.max_column + 1):
                                ws2.cell(row=ri, column=ci_x).border = bdr_i
                            if sv == 'NOM':
                                for ci_x in range(1, ws2.max_column + 1):
                                    ws2.cell(row=ri, column=ci_x).fill = green_fi
                            elif sv == 'CEK':
                                for ci_x in range(1, ws2.max_column + 1):
                                    ws2.cell(row=ri, column=ci_x).fill = yellow_fi

                        for ws_x in [ws, ws2]:
                            for ci_x in range(1, ws_x.max_column + 1):
                                ml = 0
                                for ri in range(1, min(ws_x.max_row + 1, 100)):
                                    cv = ws_x.cell(row=ri, column=ci_x).value
                                    if cv:
                                        ml = max(ml, len(str(cv)))
                                ws_x.column_dimensions[ws_x.cell(row=1, column=ci_x).column_letter].width = min(ml + 3, 50)

                    out_imp.seek(0)
                    st.session_state.importir_excel = out_imp.getvalue()
                    st.rerun()

        except Exception as e:
            st.error(f"❌ Gagal membaca file: {str(e)}")
            import traceback
            st.code(traceback.format_exc())

    if st.session_state.importir_result is not None:
        df_res = st.session_state.importir_result
        ket_col_name = st.session_state.get('importir_ket_col', 'Keterangan_AI')
        bidang_col_name = st.session_state.get('importir_bidang_col', 'Bidang_Usaha_AI')
        alasan_col_name = st.session_state.get('importir_alasan_col', 'Alasan_Analisis')

        if ket_col_name in df_res.columns:
            total_nom = len(df_res[df_res[ket_col_name].astype(str).str.strip().str.upper() == 'NOM'])
            total_cek = len(df_res[df_res[ket_col_name].astype(str).str.strip().str.upper() == 'CEK'])
            total_kosong = len(df_res[df_res[ket_col_name].isna() | (df_res[ket_col_name].astype(str).str.strip() == '')])

            st.success(f"✅ Analisis selesai! NOM: {total_nom} | CEK: {total_cek} | Belum diisi: {total_kosong}")

            c_i1, c_i2, c_i3 = st.columns(3)
            c_i1.metric("NOM (Bukan BPOM)", total_nom)
            c_i2.metric("CEK (Perlu Dicek)", total_cek)
            c_i3.metric("Belum Diisi", total_kosong)

        filter_imp = st.selectbox("Filter Klasifikasi:", ["Semua", "CEK", "NOM", "Belum Diisi"], key="filter_importir")
        if filter_imp == "CEK":
            df_show_imp = df_res[df_res[ket_col_name].astype(str).str.strip().str.upper() == 'CEK']
        elif filter_imp == "NOM":
            df_show_imp = df_res[df_res[ket_col_name].astype(str).str.strip().str.upper() == 'NOM']
        elif filter_imp == "Belum Diisi":
            df_show_imp = df_res[df_res[ket_col_name].isna() | (df_res[ket_col_name].astype(str).str.strip() == '')]
        else:
            df_show_imp = df_res

        show_cols = []
        for cn_s in df_res.columns:
            cn_upper = str(cn_s).upper()
            if 'NAMA_IMPORTIR' in cn_upper or 'NAMA IMPORTIR' in cn_upper or cn_s == ket_col_name or cn_s == bidang_col_name or cn_s == alasan_col_name or 'BRGURAI' in cn_upper or 'NOHS' in cn_upper or 'STATUS' in cn_upper:
                show_cols.append(cn_s)

        if not show_cols:
            show_cols = list(df_res.columns)

        def color_imp_row(row):
            kls = str(row.get(ket_col_name, '')).strip().upper()
            if kls == 'NOM':
                return ['background-color: #d4edda'] * len(row)
            elif kls == 'CEK':
                return ['background-color: #fff3cd'] * len(row)
            return [''] * len(row)

        st.dataframe(df_show_imp[show_cols].style.apply(color_imp_row, axis=1), height=500)

        if st.session_state.importir_excel is not None:
            st.download_button(
                label="📥 Download Hasil Analisis Importir (Excel)",
                data=st.session_state.importir_excel,
                file_name="hasil_analisis_importir.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    elif not file_importir:
        st.info("👆 Silakan upload file **Excel** untuk memulai analisis importir.")

with tab_merge:
    st.markdown("### 🔗 Gabung Data dari 2 File Excel")
    st.markdown("""Upload **2 file Excel** dengan struktur yang sama. Sistem akan:
- **Mempertahankan format asli** File Utama (filter, warna, font, lebar kolom, dll)
- **Hanya mengisi sel kosong** di File Utama dengan data dari File Pelengkap
- **Tidak mengubah** data yang sudah ada""")

    col_mg1, col_mg2 = st.columns(2)
    with col_mg1:
        st.markdown("**📁 File Utama** (yang ada sel kosong)")
        file_merge_main = st.file_uploader("Upload File Utama (.xlsx)", type=["xlsx"], key="file_merge_main")
    with col_mg2:
        st.markdown("**📁 File Pelengkap** (yang datanya lebih lengkap)")
        file_merge_source = st.file_uploader("Upload File Pelengkap (.xlsx)", type=["xlsx"], key="file_merge_source")

    if 'merge_excel' not in st.session_state:
        st.session_state.merge_excel = None
    if 'merge_stats' not in st.session_state:
        st.session_state.merge_stats = None
    if 'merge_file_id' not in st.session_state:
        st.session_state.merge_file_id = None
    if 'merge_filename' not in st.session_state:
        st.session_state.merge_filename = None

    if file_merge_main and file_merge_source:
        current_merge_id = f"{file_merge_main.name}_{file_merge_main.size}_{file_merge_source.name}_{file_merge_source.size}"
        if st.session_state.merge_file_id != current_merge_id:
            st.session_state.merge_excel = None
            st.session_state.merge_stats = None
            st.session_state.merge_file_id = current_merge_id

        try:
            from openpyxl import load_workbook

            df_main_peek = pd.read_excel(file_merge_main, header=None, nrows=5)
            file_merge_main.seek(0)
            df_src_peek = pd.read_excel(file_merge_source, header=None, nrows=5)
            file_merge_source.seek(0)

            df_main_info = pd.read_excel(file_merge_main, header=None, nrows=0)
            file_merge_main.seek(0)
            df_src_info = pd.read_excel(file_merge_source, header=None, nrows=0)
            file_merge_source.seek(0)

            main_cols = df_main_peek.shape[1]
            src_cols = df_src_peek.shape[1]

            main_rows_est = None
            src_rows_est = None

            st.success(f"✅ File berhasil dibaca!")

            col_info_mg1, col_info_mg2 = st.columns(2)
            with col_info_mg1:
                st.info(f"**File Utama**: {main_cols} kolom")
            with col_info_mg2:
                st.info(f"**File Pelengkap**: {src_cols} kolom")

            col_headers = []
            for ci in range(main_cols):
                found_header = None
                for ri in range(min(5, len(df_main_peek))):
                    v = df_main_peek.iloc[ri, ci]
                    if pd.notna(v) and str(v).strip() and str(v).strip().upper() != 'NAN':
                        vs = str(v).strip().upper()
                        if any(kw in vs for kw in ['NAMA', 'NO', 'TANGGAL', 'KODE', 'STATUS', 'KANTOR', 'ALAMAT', 'NPWP', 'SERIAL', 'SATUAN', 'KEMASAN', 'NEGARA', 'PELABUHAN', 'PENJELASAN', 'HASIL', 'ESTIMASI', 'CENTANG', 'NOMOR', 'ALASAN']):
                            found_header = str(v).strip()
                            break
                if not found_header:
                    for ri in range(min(5, len(df_main_peek))):
                        v = df_main_peek.iloc[ri, ci]
                        if pd.notna(v) and str(v).strip() and str(v).strip().upper() != 'NAN' and len(str(v).strip()) > 2:
                            found_header = str(v).strip()
                            break
                if not found_header:
                    found_header = f'Kolom_{ci}'

                col_idx_1 = ci + 1
                col_letter = chr(64 + col_idx_1) if col_idx_1 <= 26 else chr(64 + ((col_idx_1 - 1) // 26)) + chr(65 + ((col_idx_1 - 1) % 26))
                col_headers.append(f"{col_letter}: {found_header[:40]}")

            st.markdown("---")
            st.markdown("#### ⚙️ Pengaturan Penggabungan")

            mode_merge = st.radio(
                "Mode penggabungan:",
                ["Isi semua sel kosong di File Utama dari File Pelengkap", "Pilih kolom tertentu saja"],
                key="mode_merge"
            )

            selected_cols_merge = list(range(1, main_cols + 1))
            if mode_merge == "Pilih kolom tertentu saja":
                selected_headers = st.multiselect(
                    "Pilih kolom yang ingin digabungkan:",
                    options=col_headers,
                    default=[],
                    key="sel_cols_merge"
                )
                selected_cols_merge = []
                for sh in selected_headers:
                    col_letter = sh.split(":")[0].strip()
                    if len(col_letter) == 1:
                        selected_cols_merge.append(ord(col_letter) - 64)
                    elif len(col_letter) == 2:
                        selected_cols_merge.append((ord(col_letter[0]) - 64) * 26 + (ord(col_letter[1]) - 64))

            start_row_mg = st.number_input("Mulai dari baris ke- (di Excel):", min_value=1, value=1, step=1, key="start_row_merge")
            end_row_mg = st.number_input("Sampai baris ke- (di Excel, 0 = sampai akhir):", min_value=0, value=0, step=1, key="end_row_merge")

            overwrite_mode = st.checkbox("Timpa data yang sudah ada (overwrite)", value=False, key="overwrite_merge")

            if st.button("🔄 Gabungkan Data", key="btn_merge"):
                progress_mg = st.progress(0, text="Membaca File Pelengkap (cepat via pandas)...")

                file_merge_source.seek(0)
                df_src_all = pd.read_excel(file_merge_source, header=None)
                src_total_rows = len(df_src_all)

                progress_mg.progress(15, text="Membuka File Utama (mempertahankan format)...")

                file_merge_main.seek(0)
                wb_main = load_workbook(file_merge_main)
                ws_main = wb_main.active

                main_total_rows = ws_main.max_row or 0
                max_r = min(main_total_rows, src_total_rows)
                max_c = min(ws_main.max_column or 0, df_src_all.shape[1])

                actual_start = max(start_row_mg, 1)
                actual_end = end_row_mg if end_row_mg > 0 else max_r

                if main_total_rows != src_total_rows:
                    st.warning(f"⚠️ Jumlah baris berbeda! File Utama: {main_total_rows}, File Pelengkap: {src_total_rows}. Diproses sampai baris terpendek.")

                progress_mg.progress(30, text="Mengisi sel kosong...")

                fill_stats = {}
                total_cols_to_process = len([c for c in selected_cols_merge if c <= max_c])
                processed_cols = 0

                for ci in selected_cols_merge:
                    if ci > max_c:
                        continue
                    col_label = col_headers[ci - 1] if ci - 1 < len(col_headers) else f'Kolom_{ci}'
                    count = 0
                    pandas_ci = ci - 1

                    for ri in range(actual_start, min(actual_end + 1, max_r + 1)):
                        v_main = ws_main.cell(row=ri, column=ci).value
                        main_empty = v_main is None or str(v_main).strip() == '' or str(v_main).strip().lower() == 'nan'

                        if main_empty or overwrite_mode:
                            src_ri = ri - 1
                            if src_ri < len(df_src_all):
                                v_src = df_src_all.iloc[src_ri, pandas_ci]
                                src_filled = pd.notna(v_src) and str(v_src).strip() != '' and str(v_src).strip().lower() != 'nan'
                                if src_filled:
                                    ws_main.cell(row=ri, column=ci).value = v_src if not isinstance(v_src, float) or not v_src != v_src else None
                                    count += 1

                    if count > 0:
                        fill_stats[col_label] = count

                    processed_cols += 1
                    pct = 30 + int(50 * processed_cols / max(total_cols_to_process, 1))
                    progress_mg.progress(pct, text=f"Mengisi kolom {col_label}... ({processed_cols}/{total_cols_to_process})")

                progress_mg.progress(85, text="Menyimpan file (mempertahankan format asli)...")

                out_merge = io.BytesIO()
                wb_main.save(out_merge)
                wb_main.close()
                out_merge.seek(0)

                st.session_state.merge_excel = out_merge.getvalue()
                st.session_state.merge_stats = fill_stats

                base_name = file_merge_main.name
                if base_name.endswith('.xlsx'):
                    base_name = base_name[:-5]
                st.session_state.merge_filename = f"{base_name}_LENGKAP.xlsx"

                progress_mg.progress(100, text="Selesai!")
                st.rerun()

        except Exception as e:
            st.error(f"❌ Gagal membaca file: {str(e)}")
            import traceback
            st.code(traceback.format_exc())

    if st.session_state.merge_stats is not None:
        fill_stats = st.session_state.merge_stats
        total_filled = sum(fill_stats.values())

        st.success(f"✅ Penggabungan selesai! Total **{total_filled}** sel berhasil diisi.")
        st.markdown("📌 **Format asli file dipertahankan** (filter, warna, font, lebar kolom, dll.)")

        if fill_stats:
            st.markdown("**📊 Rincian per kolom:**")
            stats_df = pd.DataFrame([{'Kolom': k, 'Jumlah Sel Terisi': v} for k, v in fill_stats.items()])
            st.dataframe(stats_df, height=min(len(stats_df) * 40 + 50, 400))
        else:
            st.info("Tidak ada sel yang perlu diisi. Data di File Utama sudah lengkap atau tidak ada data pelengkap yang cocok.")

        if st.session_state.merge_excel is not None:
            dl_name = st.session_state.merge_filename or "file_gabungan.xlsx"
            st.download_button(
                label="📥 Download File Gabungan (Excel)",
                data=st.session_state.merge_excel,
                file_name=dl_name,
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

    elif not file_merge_main or not file_merge_source:
        st.info("👆 Silakan upload **File Utama** dan **File Pelengkap** untuk memulai penggabungan data.")

with tab_notulen:
    st.markdown("### 📝 Generator Notulen Rapat")
    st.markdown("""Buat dokumen notulen rapat resmi dalam format **Word (.docx)** sesuai template standar.
- **Upload PDF undangan** → sistem otomatis baca dan isi informasi rapat
- **Paste ringkasan/summary** → AI otomatis susun jadi pendahuluan, pembahasan, dan kesimpulan
- Semua field bisa diedit sebelum generate""")

    if 'notulen_docx' not in st.session_state:
        st.session_state.notulen_docx = None
    if 'notulen_filename' not in st.session_state:
        st.session_state.notulen_filename = None
    if 'nt_parsed_undangan' not in st.session_state:
        st.session_state.nt_parsed_undangan = None
    if 'nt_parsed_summary' not in st.session_state:
        st.session_state.nt_parsed_summary = None

    if 'nt_pending_fill' in st.session_state and st.session_state.nt_pending_fill:
        pf = st.session_state.nt_pending_fill
        new_pb_count = pf.get("nt_pembahasan_count", None)
        if new_pb_count is not None:
            old_count = st.session_state.get("nt_pembahasan_count", 1)
            for oi in range(new_pb_count, max(old_count, new_pb_count) + 5):
                st.session_state.pop(f"nt_speaker_{oi}", None)
                st.session_state.pop(f"nt_content_{oi}", None)
        for k, v in pf.items():
            if v is not None:
                st.session_state[k] = v
        st.session_state.nt_pending_fill = None

    with st.expander("⚙️ Pengaturan AI — Klik di sini jika fitur AI tidak berjalan", expanded=False):
        st.info("💡 **ChatGPT Plus tidak bisa dihubungkan langsung ke aplikasi.** ChatGPT Plus adalah layanan chat (chat.openai.com), sedangkan yang dibutuhkan aplikasi ini adalah **API Key** — layanan terpisah. Gunakan salah satu opsi gratis di bawah ini:")

        st.markdown("---")
        st.markdown("### 🆓 Opsi 1 — Groq (Direkomendasikan, Benar-benar GRATIS)")
        st.markdown("""
**Langkah mendapatkan Groq API Key:**
1. Buka [console.groq.com](https://console.groq.com) di browser
2. Klik **Sign Up** → daftar dengan Google atau email
3. Setelah login, klik menu **API Keys** di sebelah kiri
4. Klik **Create API Key** → beri nama → klik **Submit**
5. **Copy key-nya** (dimulai dengan `gsk_...`)
6. Paste di kolom di bawah ini

> Groq menyediakan akses AI gratis (model Llama 3) yang cukup cepat dan akurat untuk membuat notulen.
        """)

        st.markdown("### 💰 Opsi 2 — OpenAI API (Berbayar, kualitas tertinggi)")
        st.markdown("""
**Perlu diketahui:** Langganan **ChatGPT Plus tidak termasuk akses API**. API OpenAI ditagih terpisah berdasarkan penggunaan.

Jika ingin tetap pakai OpenAI:
1. Buka [platform.openai.com/api-keys](https://platform.openai.com/api-keys)
2. Login → klik **Create new secret key**
3. Isi saldo minimal $5 di menu Billing
4. Copy key (dimulai `sk-...`) dan paste di bawah
        """)

        st.markdown("---")
        nt_manual_key = st.text_input("Paste API Key di sini:", type="password", key="nt_openai_key_input",
                                       placeholder="gsk_... (Groq gratis) atau sk-... (OpenAI berbayar)")
        if nt_manual_key:
            st.session_state.nt_openai_key = nt_manual_key
            if nt_manual_key.startswith("gsk_"):
                st.success("✅ Groq API Key tersimpan! Fitur AI notulen siap digunakan untuk sesi ini.")
            else:
                st.success("✅ OpenAI API Key tersimpan untuk sesi ini.")
        elif "nt_openai_key" not in st.session_state:
            st.session_state.nt_openai_key = ""
        if st.session_state.get("nt_openai_key"):
            active = "Groq" if st.session_state.nt_openai_key.startswith("gsk_") else "OpenAI"
            st.caption(f"Key aktif saat ini: {active} (...{st.session_state.nt_openai_key[-6:]})")

    def _nt_call_ai(prompt, max_tokens=4000):
        import openai as _openai_mod
        ai_base = os.environ.get("AI_INTEGRATIONS_OPENAI_BASE_URL", "")
        ai_key_int = os.environ.get("AI_INTEGRATIONS_OPENAI_API_KEY", "")
        user_key = st.session_state.get("nt_openai_key", "").strip()
        last_error = None
        clients_to_try = []
        if ai_base and ai_key_int:
            clients_to_try.append(("integration", _openai_mod.OpenAI(base_url=ai_base, api_key=ai_key_int), "gpt-4o-mini"))
        if user_key:
            if user_key.startswith("gsk_"):
                groq_client = _openai_mod.OpenAI(api_key=user_key, base_url="https://api.groq.com/openai/v1")
                clients_to_try.append(("groq", groq_client, "llama-3.1-8b-instant"))
            else:
                clients_to_try.append(("openai", _openai_mod.OpenAI(api_key=user_key), "gpt-4o-mini"))
        if not clients_to_try:
            raise Exception("❌ AI tidak tersedia. Masukkan API Key di bagian '⚙️ Pengaturan API' di atas.\n\nGunakan Groq (gratis) dengan key gsk_... dari console.groq.com")
        for source, client, model in clients_to_try:
            try:
                prompt_to_use = prompt
                tokens_out = max_tokens
                if source == "groq":
                    tokens_out = min(max_tokens, 1800)
                    cutoff = prompt_to_use.find('"""', 500)
                    end_cutoff = prompt_to_use.find('"""', cutoff + 3) if cutoff > 0 else -1
                    if end_cutoff > 0:
                        content_part = prompt_to_use[cutoff+3:end_cutoff][:2200]
                        prompt_to_use = prompt_to_use[:cutoff+3] + content_part + prompt_to_use[end_cutoff:]
                    elif len(prompt_to_use) > 3000:
                        prompt_to_use = prompt_to_use[:3000]
                    if "json" not in prompt_to_use.lower():
                        prompt_to_use += "\n\nKembalikan HANYA dalam format JSON valid."
                kwargs = dict(
                    model=model,
                    messages=[{"role": "user", "content": prompt_to_use}],
                    response_format={"type": "json_object"}
                )
                if source == "integration":
                    kwargs["max_completion_tokens"] = tokens_out
                else:
                    kwargs["max_tokens"] = tokens_out
                resp = client.chat.completions.create(**kwargs)
                return resp.choices[0].message.content
            except Exception as e:
                last_error = e
                err_str = str(e)
                if any(x in err_str for x in ["401", "ApiKey", "Unauthorized", "ApiKeyNotApproved"]):
                    continue
                if source == "groq" and ("413" in err_str or "too large" in err_str.lower() or "rate_limit" in err_str or "tokens" in err_str):
                    try:
                        short_prompt = prompt[:2000] + '\n"""\n\nKembalikan HANYA dalam format JSON valid.'
                        resp2 = client.chat.completions.create(
                            model=model,
                            messages=[{"role": "user", "content": short_prompt}],
                            response_format={"type": "json_object"},
                            max_tokens=1500
                        )
                        return resp2.choices[0].message.content
                    except Exception:
                        pass
                raise e
        if last_error:
            if any(x in str(last_error) for x in ["401", "ApiKey", "Unauthorized"]):
                if not user_key:
                    raise Exception("❌ AI Integration tidak aktif. Masukkan API Key di bagian '⚙️ Pengaturan API'.\n\nGunakan Groq (gratis) dari console.groq.com")
            raise last_error

    st.markdown("---")
    st.markdown("#### 📎 Upload Otomatis (Opsional)")
    st.markdown("Upload file undangan dan/atau paste ringkasan rapat. Sistem akan otomatis mengisi form di bawah.")

    col_auto1, col_auto2 = st.columns(2)
    with col_auto1:
        nt_file_undangan = st.file_uploader("Upload Surat Undangan (PDF/Word):", type=["pdf", "docx"], key="nt_file_undangan")
        if nt_file_undangan and st.button("🔍 Baca & Isi dari Undangan", key="btn_parse_undangan"):
            with st.spinner("Membaca file undangan..."):
                try:
                    undangan_text = ""
                    if nt_file_undangan.name.lower().endswith('.pdf'):
                        import pdfplumber
                        nt_file_undangan.seek(0)
                        with pdfplumber.open(nt_file_undangan) as pdf_doc:
                            for page in pdf_doc.pages:
                                page_text = page.extract_text()
                                if page_text:
                                    undangan_text += page_text + "\n"
                    elif nt_file_undangan.name.lower().endswith('.docx'):
                        from docx import Document as DocxRead
                        nt_file_undangan.seek(0)
                        doc_read = DocxRead(nt_file_undangan)
                        for p in doc_read.paragraphs:
                            if p.text.strip():
                                undangan_text += p.text.strip() + "\n"
                        for tbl in doc_read.tables:
                            for row in tbl.rows:
                                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                                if row_text:
                                    undangan_text += row_text + "\n"

                    if not undangan_text.strip():
                        st.error("❌ Tidak dapat membaca teks dari file.")
                    else:
                        prompt_undangan = f"""Baca teks surat undangan rapat berikut dan ekstrak informasi dalam format JSON.

Teks undangan:
\"\"\"
{undangan_text[:4000]}
\"\"\"

Ekstrak dan kembalikan JSON dengan field berikut (isi string kosong "" jika tidak ditemukan):
{{
  "judul": "judul/topik/perihal rapat",
  "tanggal": "hari/tanggal rapat (format: Senin, 1 Januari 2026)",
  "waktu": "waktu rapat (format: 09.00 - 12.00 WIB)",
  "tempat": "tempat/platform rapat",
  "meeting_id": "meeting ID jika ada",
  "password": "password meeting jika ada",
  "pimpinan": "pimpinan/penyelenggara rapat",
  "peserta": ["daftar peserta/instansi yang diundang"]
}}

PENTING: Kembalikan HANYA JSON, tanpa markdown atau penjelasan."""

                        raw_und = _nt_call_ai(prompt_undangan, max_tokens=2000)
                        result_und = _json.loads(raw_und)
                        st.session_state.nt_parsed_undangan = result_und
                        pf = {}
                        if result_und.get("judul"):
                            pf["nt_judul"] = result_und["judul"]
                        tanggal_val = result_und.get("tanggal", "")
                        if result_und.get("waktu"):
                            tanggal_val = f"{tanggal_val}, {result_und['waktu']}" if tanggal_val else result_und['waktu']
                        if tanggal_val:
                            pf["nt_tanggal"] = tanggal_val
                        if result_und.get("meeting_id"):
                            pf["nt_meeting_id"] = result_und["meeting_id"]
                        if result_und.get("password"):
                            pf["nt_password"] = result_und["password"]
                        if result_und.get("tempat"):
                            pf["nt_tempat"] = result_und["tempat"]
                        if result_und.get("pimpinan"):
                            pf["nt_pimpinan"] = result_und["pimpinan"]
                        if result_und.get("peserta") and isinstance(result_und["peserta"], list):
                            pf["nt_peserta"] = "\n".join(result_und["peserta"])
                        st.session_state.nt_pending_fill = pf
                        st.rerun()
                except Exception as e:
                    st.error(f"❌ Gagal membaca undangan: {str(e)}")

    with col_auto2:
        nt_summary_input = st.text_area("Paste Ringkasan/Summary Rapat:", height=200, placeholder="Paste ringkasan hasil rapat di sini...\n\nBisa berupa catatan kasar, poin-poin, atau narasi bebas. AI akan menyusun menjadi format notulen.", key="nt_summary_input")
        if nt_summary_input.strip() and st.button("🤖 Olah Ringkasan dengan AI", key="btn_parse_summary"):
            with st.spinner("AI sedang menyusun ringkasan..."):
                try:
                    prompt_summary = f"""Kamu adalah notulis profesional senior instansi pemerintah Indonesia yang sangat berpengalaman dalam menyusun dokumen resmi.

INSTRUKSI WAJIB: Kembalikan HANYA JSON valid dengan PERSIS struktur berikut (jangan ubah nama field, jangan tambah field lain):
{{
  "judul": "judul/topik rapat",
  "tanggal": "tanggal rapat",
  "tempat": "tempat/platform rapat",
  "pimpinan": "pimpinan rapat",
  "peserta": ["peserta 1", "peserta 2"],
  "pendahuluan": "paragraf pembuka notulen dalam bahasa Indonesia resmi",
  "pembahasan": [
    {{"speaker": "Nama pembicara", "content": "Isi pembahasan dalam bahasa Indonesia resmi"}}
  ],
  "kesimpulan": ["poin kesimpulan 1", "poin kesimpulan 2"]
}}

Isi string kosong "" untuk field yang tidak ditemukan. Array peserta/kesimpulan boleh kosong [].

TUGAS: Baca catatan rapat berikut dan susun menjadi notulen formal yang komprehensif dan elaboratif:
\"\"\"
{nt_summary_input[:8000]}
\"\"\"

PANDUAN PENULISAN WAJIB:
- Gunakan bahasa Indonesia resmi, formal, dan baku sesuai standar dokumen pemerintah
- JANGAN meringkas secara berlebihan — uraikan setiap poin secara lengkap dan elaboratif dengan kalimat-kalimat yang panjang dan terstruktur
- Setiap topik pembahasan harus mencakup latar belakang, isi materi, data/angka yang disebutkan, permasalahan yang diangkat, dan upaya/rekomendasi yang disampaikan
- Gunakan kalimat majemuk bertingkat yang mencerminkan bahasa dokumen resmi pemerintah (contoh: "Dalam rangka..., maka dipandang perlu untuk..., sehingga dapat...")
- Pertahankan SEMUA nama, angka, regulasi, pasal, nomor peraturan, dan data teknis — jangan dihilangkan
- Setiap item pembahasan dalam array "pembahasan" boleh terdiri dari beberapa paragraf yang digabung dengan \\n\\n
- Kesimpulan ditulis sebagai kalimat lengkap, bukan hanya poin pendek
- Kembalikan HANYA JSON valid tanpa markdown atau penjelasan lain"""

                    raw_sum = _nt_call_ai(prompt_summary, max_tokens=4000)
                    result_sum = _json.loads(raw_sum)
                    st.session_state.nt_debug_raw = raw_sum[:3000]
                    st.session_state.nt_parsed_summary = result_sum
                    pf = {}
                    if result_sum.get("judul"):
                        pf["nt_judul"] = result_sum["judul"]
                    if result_sum.get("tanggal"):
                        pf["nt_tanggal"] = result_sum["tanggal"]
                    if result_sum.get("tempat"):
                        pf["nt_tempat"] = result_sum["tempat"]
                    if result_sum.get("pimpinan"):
                        pf["nt_pimpinan"] = result_sum["pimpinan"]
                    if result_sum.get("peserta") and isinstance(result_sum["peserta"], list):
                        pf["nt_peserta"] = "\n".join(result_sum["peserta"])
                    if result_sum.get("pendahuluan"):
                        pf["nt_pendahuluan"] = result_sum["pendahuluan"]
                    parsed_pb = result_sum.get("pembahasan", [])
                    if parsed_pb and isinstance(parsed_pb, list):
                        pf["nt_pembahasan_count"] = max(len(parsed_pb), 1)
                        for pi, pb in enumerate(parsed_pb):
                            pf[f"nt_speaker_{pi}"] = pb.get("speaker", "") if isinstance(pb, dict) else ""
                            pf[f"nt_content_{pi}"] = pb.get("content", "") if isinstance(pb, dict) else str(pb)
                    if result_sum.get("kesimpulan"):
                        kes = result_sum["kesimpulan"]
                        if isinstance(kes, list):
                            pf["nt_kesimpulan"] = "\n".join(kes)
                        elif isinstance(kes, str):
                            pf["nt_kesimpulan"] = kes
                    st.session_state.nt_pending_fill = pf
                    filled = [k for k, v in pf.items() if v]
                    st.session_state.nt_fill_summary = f"✅ AI berhasil mengisi {len(filled)} field: {', '.join(filled)}"
                    st.rerun()
                except Exception as e:
                    st.error(f"❌ Gagal mengolah ringkasan: {str(e)}")

    st.markdown("---")
    st.markdown("#### 📋 Upload Ringkasan tldv.io (Opsional)")
    st.markdown("Upload satu atau lebih file ringkasan dari **tldv.io** (format .txt). Cocok untuk rapat dengan beberapa sesi.")

    _tldv_json_path = "attached_assets/preloaded_notulen_7mei2026.json"
    import os as _os
    _has_preloaded = _os.path.exists(_tldv_json_path)
    if _has_preloaded:
        if st.button("📥 Muat Rapat 7 Mei 2026 (Langsung, Tanpa API)", key="btn_load_preloaded", type="primary"):
            try:
                with open(_tldv_json_path, "r", encoding="utf-8") as _fj:
                    _result_tldv = _json.load(_fj)

                pf = {}
                if _result_tldv.get("judul"):
                    pf["nt_judul"] = _result_tldv["judul"]
                if _result_tldv.get("tanggal"):
                    pf["nt_tanggal"] = _result_tldv["tanggal"]
                if _result_tldv.get("tempat"):
                    pf["nt_tempat"] = _result_tldv["tempat"]
                if _result_tldv.get("pimpinan"):
                    pf["nt_pimpinan"] = _result_tldv["pimpinan"]
                if _result_tldv.get("peserta") and isinstance(_result_tldv["peserta"], list):
                    pf["nt_peserta"] = "\n".join(_result_tldv["peserta"])
                if _result_tldv.get("pendahuluan"):
                    pf["nt_pendahuluan"] = _result_tldv["pendahuluan"]
                _parsed_pb = _result_tldv.get("pembahasan", [])
                if _parsed_pb and isinstance(_parsed_pb, list):
                    pf["nt_pembahasan_count"] = max(len(_parsed_pb), 1)
                    for _pi, _pb in enumerate(_parsed_pb):
                        pf[f"nt_speaker_{_pi}"] = _pb.get("speaker", "") if isinstance(_pb, dict) else ""
                        pf[f"nt_content_{_pi}"] = _pb.get("content", "") if isinstance(_pb, dict) else str(_pb)
                _kes = _result_tldv.get("kesimpulan", [])
                _tl = _result_tldv.get("tindak_lanjut", [])
                _all_kes = []
                if isinstance(_kes, list):
                    _all_kes.extend(_kes)
                elif isinstance(_kes, str) and _kes:
                    _all_kes.append(_kes)
                if isinstance(_tl, list) and _tl:
                    _all_kes.append("")
                    _all_kes.append("TINDAK LANJUT:")
                    _all_kes.extend(_tl)
                if _all_kes:
                    pf["nt_kesimpulan"] = "\n".join(_all_kes)

                st.session_state.nt_pending_fill = pf
                st.session_state.nt_fill_summary = f"✅ Rapat 7 Mei 2026 berhasil dimuat! {len(_parsed_pb)} topik pembahasan, {len(_kes)} poin kesimpulan, {len(_tl)} tindak lanjut — semuanya sudah terisi. Tinggal isi nama pimpinan, upload foto, dan klik Generate."
                st.rerun()
            except Exception as _e:
                st.error(f"❌ Gagal memuat: {str(_e)}")

    st.markdown("Atau upload file tldv.io Anda sendiri:")
    nt_tldv_files = st.file_uploader(
        "Upload file ringkasan tldv.io (.txt, bisa lebih dari satu untuk multi-sesi):",
        type=["txt"],
        accept_multiple_files=True,
        key="nt_tldv_files",
        help="File ringkasan dari tldv.io — format teks dengan topik dan bullet points"
    )

    if nt_tldv_files and st.button("📋 Olah Ringkasan tldv.io dengan AI", key="btn_parse_tldv", type="secondary"):
        with st.spinner("Memproses ringkasan tldv.io dengan AI..."):
            try:
                _texts = []
                for _i, _tf in enumerate(nt_tldv_files):
                    _raw = _tf.read()
                    try:
                        _texts.append(f"=== SESI {_i+1} ({_tf.name}) ===\n{_raw.decode('utf-8')}")
                    except UnicodeDecodeError:
                        _texts.append(f"=== SESI {_i+1} ({_tf.name}) ===\n{_raw.decode('latin-1')}")
                _combined_tldv = "\n\n".join(_texts)

                _prompt_tldv2 = f"""Kamu adalah notulis profesional senior instansi pemerintah Indonesia yang sangat berpengalaman dalam menyusun dokumen resmi pemerintah.

Berikut adalah ringkasan hasil rapat dari tldv.io:
\"\"\"
{_combined_tldv[:9000]}
\"\"\"

INSTRUKSI WAJIB: Kembalikan HANYA JSON valid dengan PERSIS struktur berikut:
{{
  "judul": "judul lengkap rapat berdasarkan topik yang dibahas",
  "tanggal": "tanggal rapat jika tersedia, atau string kosong",
  "tempat": "Zoom Meeting",
  "pimpinan": "pimpinan rapat jika teridentifikasi",
  "peserta": ["unit/instansi peserta 1", "unit/instansi peserta 2"],
  "pendahuluan": "paragraf pembuka notulen elaboratif dalam bahasa Indonesia resmi",
  "pembahasan": [
    {{"speaker": "Topik/Nama", "content": "Isi pembahasan elaboratif multi-paragraf"}}
  ],
  "kesimpulan": ["poin kesimpulan elaboratif 1", "poin kesimpulan elaboratif 2"],
  "tindak_lanjut": ["item tindak lanjut dari Item Tindakan"]
}}

PANDUAN PENULISAN WAJIB:
- Gunakan bahasa Indonesia resmi, formal, dan baku sesuai standar dokumen pemerintah
- Jika ada beberapa sesi, gabungkan menjadi satu notulen yang komprehensif dan kohesif
- JANGAN meringkas terlalu singkat — uraikan setiap topik secara lengkap dengan kalimat majemuk bertingkat gaya dokumen resmi
- Setiap item pembahasan harus mencakup konteks, isi diskusi, keputusan, dan tindak lanjut
- Pertahankan SEMUA angka, persentase, tanggal, nama regulasi, SLA, tarif, dan data teknis
- Kembalikan HANYA JSON valid tanpa markdown atau penjelasan lain"""

                _raw2 = _nt_call_ai(_prompt_tldv2, max_tokens=4000)
                _result2 = _json.loads(_raw2)

                pf2 = {}
                if _result2.get("judul"):
                    pf2["nt_judul"] = _result2["judul"]
                if _result2.get("tanggal"):
                    pf2["nt_tanggal"] = _result2["tanggal"]
                if _result2.get("tempat"):
                    pf2["nt_tempat"] = _result2["tempat"]
                if _result2.get("pimpinan"):
                    pf2["nt_pimpinan"] = _result2["pimpinan"]
                if _result2.get("peserta") and isinstance(_result2["peserta"], list):
                    pf2["nt_peserta"] = "\n".join(_result2["peserta"])
                if _result2.get("pendahuluan"):
                    pf2["nt_pendahuluan"] = _result2["pendahuluan"]
                _pb2 = _result2.get("pembahasan", [])
                if _pb2 and isinstance(_pb2, list):
                    pf2["nt_pembahasan_count"] = max(len(_pb2), 1)
                    for _pi2, _pb2i in enumerate(_pb2):
                        pf2[f"nt_speaker_{_pi2}"] = _pb2i.get("speaker", "") if isinstance(_pb2i, dict) else ""
                        pf2[f"nt_content_{_pi2}"] = _pb2i.get("content", "") if isinstance(_pb2i, dict) else str(_pb2i)
                _kes2 = _result2.get("kesimpulan", [])
                _tl2 = _result2.get("tindak_lanjut", [])
                _all2 = []
                if isinstance(_kes2, list):
                    _all2.extend(_kes2)
                elif isinstance(_kes2, str) and _kes2:
                    _all2.append(_kes2)
                if isinstance(_tl2, list) and _tl2:
                    _all2.append("")
                    _all2.append("TINDAK LANJUT:")
                    _all2.extend(_tl2)
                if _all2:
                    pf2["nt_kesimpulan"] = "\n".join(_all2)

                st.session_state.nt_pending_fill = pf2
                st.session_state.nt_fill_summary = f"✅ tldv.io ({len(nt_tldv_files)} file) berhasil diproses! {len(_pb2)} topik pembahasan diisi otomatis."
                st.rerun()
            except Exception as _e2:
                st.error(f"❌ Gagal memproses tldv.io: {str(_e2)}")
                import traceback as _tb2
                st.code(_tb2.format_exc())

    st.markdown("---")
    st.markdown("#### 🎬 Upload Transcript Zoom (Opsional)")
    with st.expander("ℹ️ Cara mendapatkan file transcript dari Zoom", expanded=False):
        st.markdown("""
**Langkah-langkah:**
1. Setelah meeting selesai, buka **zoom.us** → login
2. Klik **Recordings** di menu kiri
3. Klik nama meeting yang ingin dibuat notulennya
4. Di bagian bawah, klik **Audio Transcript** → **Download** (file .vtt)
5. Upload file .vtt tersebut di sini

> **Catatan:** Fitur transcript otomatis hanya tersedia di akun Zoom **Pro, Business, atau Enterprise** dengan Cloud Recording aktif.
> Jika recording lokal, Anda bisa copy-paste isi chat Zoom ke kolom "Ringkasan Rapat" di atas.
        """)

    nt_zoom_transcript = st.file_uploader(
        "Upload Transcript Zoom (.vtt atau .txt):",
        type=["vtt", "txt"],
        key="nt_zoom_transcript",
        help="File transcript dari Zoom Cloud Recording (format .vtt atau .txt)"
    )

    if nt_zoom_transcript and st.button("🎬 Olah Transcript Zoom dengan AI", key="btn_parse_zoom", type="secondary"):
        with st.spinner("Membaca dan memproses transcript Zoom..."):
            try:
                raw_bytes = nt_zoom_transcript.read()
                try:
                    transcript_text = raw_bytes.decode("utf-8")
                except UnicodeDecodeError:
                    transcript_text = raw_bytes.decode("latin-1")

                import re as _re_vtt

                def _parse_vtt(text):
                    lines = text.splitlines()
                    segments = []
                    current_speaker = ""
                    current_lines = []
                    i = 0
                    while i < len(lines):
                        line = lines[i].strip()
                        if '-->' in line:
                            i += 1
                            while i < len(lines) and lines[i].strip() and '-->' not in lines[i]:
                                seg_text = lines[i].strip()
                                if ':' in seg_text and not seg_text.startswith('['):
                                    parts = seg_text.split(':', 1)
                                    spk = parts[0].strip()
                                    if len(spk) < 60 and not any(c.isdigit() for c in spk[:3]):
                                        if spk != current_speaker:
                                            if current_speaker and current_lines:
                                                segments.append({"speaker": current_speaker, "text": " ".join(current_lines)})
                                            current_speaker = spk
                                            current_lines = [parts[1].strip()]
                                        else:
                                            current_lines.append(parts[1].strip())
                                    else:
                                        current_lines.append(seg_text)
                                else:
                                    current_lines.append(seg_text)
                                i += 1
                        else:
                            i += 1
                    if current_speaker and current_lines:
                        segments.append({"speaker": current_speaker, "text": " ".join(current_lines)})
                    return segments

                def _parse_zoom_txt(text):
                    lines = text.splitlines()
                    segments = []
                    current_speaker = ""
                    current_lines = []
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                        m = _re_vtt.match(r'^\[?(\d{1,2}:\d{2}(?::\d{2})?)\]?\s+(.+)', line)
                        if m:
                            rest = m.group(2)
                            if ':' in rest:
                                parts = rest.split(':', 1)
                                spk = parts[0].strip()
                                if len(spk) < 60:
                                    if spk != current_speaker:
                                        if current_speaker and current_lines:
                                            segments.append({"speaker": current_speaker, "text": " ".join(current_lines)})
                                        current_speaker = spk
                                        current_lines = [parts[1].strip()]
                                    else:
                                        current_lines.append(parts[1].strip())
                                    continue
                        if current_speaker:
                            current_lines.append(line)
                        elif ':' in line:
                            parts = line.split(':', 1)
                            spk = parts[0].strip()
                            if len(spk) < 60 and len(spk) > 1:
                                if spk != current_speaker:
                                    if current_speaker and current_lines:
                                        segments.append({"speaker": current_speaker, "text": " ".join(current_lines)})
                                    current_speaker = spk
                                    current_lines = [parts[1].strip()]
                    if current_speaker and current_lines:
                        segments.append({"speaker": current_speaker, "text": " ".join(current_lines)})
                    return segments

                fname_lower = nt_zoom_transcript.name.lower()
                if fname_lower.endswith('.vtt'):
                    segs = _parse_vtt(transcript_text)
                else:
                    segs = _parse_zoom_txt(transcript_text)

                if not segs:
                    segs_text = transcript_text[:8000]
                else:
                    segs_text = "\n".join([f"{s['speaker']}: {s['text']}" for s in segs])[:8000]

                unique_speakers = []
                seen = set()
                for s in segs:
                    if s['speaker'] not in seen and s['speaker']:
                        seen.add(s['speaker'])
                        unique_speakers.append(s['speaker'])

                prompt_zoom = f"""Kamu adalah notulis profesional instansi pemerintah Indonesia.

Berikut adalah transcript rekaman rapat Zoom:
\"\"\"
{segs_text}
\"\"\"

Pembicara yang teridentifikasi: {', '.join(unique_speakers[:20]) if unique_speakers else 'tidak diketahui'}

INSTRUKSI WAJIB: Kembalikan HANYA JSON valid dengan PERSIS struktur berikut:
{{
  "judul": "judul/topik rapat berdasarkan isi transcript",
  "tanggal": "",
  "tempat": "Zoom Meeting",
  "pimpinan": "nama pembicara yang paling sering memimpin atau membuka rapat",
  "peserta": ["nama peserta 1", "nama peserta 2"],
  "pendahuluan": "paragraf pembuka notulen dalam bahasa Indonesia resmi berdasarkan pembukaan rapat",
  "pembahasan": [
    {{"speaker": "Nama pembicara", "content": "Ringkasan isi pembicaraan dalam bahasa Indonesia resmi"}}
  ],
  "kesimpulan": ["poin kesimpulan 1", "poin kesimpulan 2"]
}}

PANDUAN PENULISAN WAJIB:
- Gunakan bahasa Indonesia resmi, formal, dan baku sesuai standar dokumen pemerintah
- Gabungkan pernyataan dari pembicara yang sama jika berurutan
- JANGAN meringkas terlalu singkat — uraikan setiap topik secara lengkap dan elaboratif dengan kalimat-kalimat yang panjang dan terstruktur
- Setiap topik pembahasan harus mencakup: latar belakang yang disampaikan, isi materi/paparan, data/angka yang disebutkan, permasalahan yang diangkat, serta upaya/solusi/rekomendasi yang disampaikan
- Gunakan kalimat majemuk bertingkat yang mencerminkan bahasa dokumen resmi pemerintah
- Pertahankan SEMUA nama, angka, regulasi, pasal, nomor peraturan, dan data teknis — jangan dihilangkan atau dipersingkat
- Setiap item pembahasan dalam array "pembahasan" boleh terdiri dari beberapa paragraf yang digabung dengan \\n\\n
- Kesimpulan ditulis sebagai kalimat lengkap yang komprehensif, bukan hanya poin singkat
- Kembalikan HANYA JSON valid, tanpa markdown atau penjelasan lain"""

                raw_zoom = _nt_call_ai(prompt_zoom, max_tokens=4000)
                result_zoom = _json.loads(raw_zoom)
                st.session_state.nt_parsed_summary = result_zoom

                pf = {}
                if result_zoom.get("judul"):
                    pf["nt_judul"] = result_zoom["judul"]
                if result_zoom.get("tanggal"):
                    pf["nt_tanggal"] = result_zoom["tanggal"]
                if result_zoom.get("tempat"):
                    pf["nt_tempat"] = result_zoom["tempat"]
                if result_zoom.get("pimpinan"):
                    pf["nt_pimpinan"] = result_zoom["pimpinan"]
                if result_zoom.get("peserta") and isinstance(result_zoom["peserta"], list):
                    pf["nt_peserta"] = "\n".join(result_zoom["peserta"])
                if result_zoom.get("pendahuluan"):
                    pf["nt_pendahuluan"] = result_zoom["pendahuluan"]
                parsed_pb = result_zoom.get("pembahasan", [])
                if parsed_pb and isinstance(parsed_pb, list):
                    pf["nt_pembahasan_count"] = max(len(parsed_pb), 1)
                    for pi, pb in enumerate(parsed_pb):
                        pf[f"nt_speaker_{pi}"] = pb.get("speaker", "") if isinstance(pb, dict) else ""
                        pf[f"nt_content_{pi}"] = pb.get("content", "") if isinstance(pb, dict) else str(pb)
                if result_zoom.get("kesimpulan"):
                    kes = result_zoom["kesimpulan"]
                    if isinstance(kes, list):
                        pf["nt_kesimpulan"] = "\n".join(kes)
                    elif isinstance(kes, str):
                        pf["nt_kesimpulan"] = kes

                n_spk = len(unique_speakers)
                n_seg = len(segs)
                st.session_state.nt_pending_fill = pf
                st.session_state.nt_fill_summary = f"✅ Transcript Zoom berhasil diproses! {n_spk} pembicara, {n_seg} segmen percakapan. Semua field sudah diisi otomatis."
                st.rerun()

            except Exception as e:
                st.error(f"❌ Gagal memproses transcript Zoom: {str(e)}")
                import traceback
                st.code(traceback.format_exc())

    if st.session_state.nt_parsed_undangan:
        st.success("✅ Informasi dari undangan berhasil diekstrak! Data sudah diisi di form bawah.")
    if st.session_state.nt_parsed_summary:
        fill_msg = st.session_state.get("nt_fill_summary", "")
        if fill_msg:
            st.success(fill_msg)
        else:
            st.success("✅ Ringkasan berhasil diolah! Pendahuluan, pembahasan, dan kesimpulan sudah diisi.")
        if st.session_state.get("nt_debug_raw"):
            with st.expander("🔍 Debug: Respons mentah dari AI", expanded=True):
                st.code(st.session_state.nt_debug_raw, language="json")

    st.markdown("---")
    st.markdown("#### A. Informasi Rapat")

    col_nt1, col_nt2 = st.columns(2)
    with col_nt1:
        nt_judul = st.text_input("Judul/Topik Rapat:", placeholder="Contoh: Rapat Penggunaan Satuan Terkecil pada Dokumen PIB", key="nt_judul")
        nt_tanggal = st.text_input("Hari/Tanggal:", placeholder="Contoh: Rabu, 4 Maret 2026", key="nt_tanggal")
        nt_meeting_id = st.text_input("Meeting ID (opsional):", placeholder="Contoh: 494 688 897 866 99", key="nt_meeting_id")
        nt_password = st.text_input("Password Meeting (opsional):", placeholder="Contoh: ET6DP6SP", key="nt_password")
    with col_nt2:
        nt_tempat = st.text_input("Tempat/Platform:", placeholder="Contoh: Microsoft Teams / Ruang Rapat Lt.3", key="nt_tempat")
        nt_pimpinan = st.text_input("Pimpinan Rapat:", placeholder="Contoh: Direktorat Jenderal Bea dan Cukai (DJBC)", key="nt_pimpinan")
        nt_notulis = st.text_input("Nama Notulis:", placeholder="Contoh: Dewi Putriayu P.", key="nt_notulis")

    nt_peserta = st.text_area("Peserta Rapat (satu per baris):", height=150, placeholder="Direktorat Pengawasan KMEI ONAPPZA\nBadan Pengawas Obat dan Makanan (BPOM)\nKementerian Perdagangan\n...", key="nt_peserta")

    st.markdown("---")
    st.markdown("#### B. Pendahuluan")
    nt_pendahuluan = st.text_area("Latar belakang / pendahuluan rapat:", height=200, placeholder="Tuliskan latar belakang dan tujuan rapat...", key="nt_pendahuluan")

    st.markdown("---")
    st.markdown("#### C. Pembahasan")
    st.markdown("Tambahkan poin-poin pembahasan. Setiap pembahasan berisi **pembicara** dan **isi pembicaraan**.")

    if 'nt_pembahasan_count' not in st.session_state:
        st.session_state.nt_pembahasan_count = 1

    col_add_rmv = st.columns([1, 1, 4])
    with col_add_rmv[0]:
        if st.button("➕ Tambah Pembahasan", key="nt_add_pembahasan"):
            st.session_state.nt_pembahasan_count += 1
            st.rerun()
    with col_add_rmv[1]:
        if st.session_state.nt_pembahasan_count > 1:
            if st.button("➖ Hapus Terakhir", key="nt_rmv_pembahasan"):
                st.session_state.nt_pembahasan_count -= 1
                st.rerun()

    nt_pembahasan_list = []
    for idx in range(st.session_state.nt_pembahasan_count):
        with st.expander(f"Pembahasan {idx + 1}", expanded=(idx == 0)):
            speaker = st.text_input(f"Pembicara/Instansi:", placeholder="Contoh: Direktorat Teknis Kepabeanan DJBS", key=f"nt_speaker_{idx}")
            content = st.text_area(f"Isi Pembahasan:", height=150, placeholder="Isi pembahasan dari pembicara...", key=f"nt_content_{idx}")
            nt_pembahasan_list.append({"speaker": speaker, "content": content})

    st.markdown("---")
    st.markdown("#### D. Dokumentasi")
    nt_fotos = st.file_uploader("Upload foto dokumentasi rapat (maks 35 foto, maks 5MB/foto):", type=["jpg", "jpeg", "png"], accept_multiple_files=True, key="nt_fotos")
    if nt_fotos and len(nt_fotos) > 35:
        st.warning("⚠️ Maksimal 35 foto. Hanya 35 foto pertama yang akan digunakan.")
        nt_fotos = nt_fotos[:35]
    if nt_fotos:
        oversized = [f.name for f in nt_fotos if f.size > 5 * 1024 * 1024]
        if oversized:
            st.warning(f"⚠️ Foto terlalu besar (>5MB): {', '.join(oversized)}. Foto ini akan dilewati.")
            nt_fotos = [f for f in nt_fotos if f.size <= 5 * 1024 * 1024]

    st.markdown("---")
    st.markdown("#### E. Kesimpulan")
    nt_kesimpulan = st.text_area("Kesimpulan rapat (setiap poin dipisahkan baris baru):", height=200, placeholder="Poin kesimpulan 1\nPoin kesimpulan 2\nPoin kesimpulan 3\n...", key="nt_kesimpulan")

    st.markdown("---")

    col_gen_reset = st.columns([1, 1, 4])
    with col_gen_reset[0]:
        gen_notulen_clicked = st.button("📄 Generate Notulen (.docx)", key="btn_gen_notulen", type="primary")
    with col_gen_reset[1]:
        if st.button("🔄 Reset Form", key="btn_reset_notulen"):
            st.session_state.nt_parsed_undangan = None
            st.session_state.nt_parsed_summary = None
            st.session_state.notulen_docx = None
            st.session_state.notulen_filename = None
            if 'nt_pembahasan_loaded' in st.session_state:
                del st.session_state.nt_pembahasan_loaded
            st.session_state.nt_pembahasan_count = 1
            st.rerun()

    if gen_notulen_clicked:
        if not nt_judul or not nt_tanggal:
            st.error("❌ Judul rapat dan tanggal wajib diisi!")
        else:
            with st.spinner("Membuat dokumen notulen..."):
                try:
                    from docx import Document as DocxDocument
                    from docx.shared import Pt, Cm, Inches, RGBColor, Emu
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from docx.enum.table import WD_TABLE_ALIGNMENT
                    from docx.oxml.ns import qn
                    from docx.oxml import OxmlElement

                    doc = DocxDocument()

                    section = doc.sections[0]
                    section.page_width = Cm(21)
                    section.page_height = Cm(29.7)
                    section.left_margin = Cm(2.54)
                    section.right_margin = Cm(2.54)
                    section.top_margin = Cm(2.54)
                    section.bottom_margin = Cm(2.54)

                    style = doc.styles['Normal']
                    font = style.font
                    font.name = 'Arial'
                    font.size = Pt(12)
                    style.paragraph_format.space_after = Pt(0)
                    style.paragraph_format.space_before = Pt(0)
                    style.paragraph_format.line_spacing = 1.15

                    BODY_INDENT = Cm(1.35)
                    SECTION_INDENT = Cm(1.42)

                    def _remove_table_borders(tbl):
                        tbl_pr = tbl._tbl.find(qn('w:tblPr'))
                        if tbl_pr is None:
                            tbl_pr = OxmlElement('w:tblPr')
                            tbl._tbl.insert(0, tbl_pr)
                        existing = tbl_pr.find(qn('w:tblBorders'))
                        if existing is not None:
                            tbl_pr.remove(existing)
                        borders = OxmlElement('w:tblBorders')
                        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                            el = OxmlElement(f'w:{edge}')
                            el.set(qn('w:val'), 'nil')
                            borders.append(el)
                        tbl_pr.append(borders)

                    def _add_run(paragraph, text, bold=False, size=Pt(12), font_name='Arial'):
                        run = paragraph.add_run(text)
                        run.bold = bold
                        run.font.name = font_name
                        run.font.size = size
                        return run

                    p_title = doc.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_title.paragraph_format.space_after = Pt(0)
                    _add_run(p_title, "NOTULEN", bold=True, size=Pt(12))

                    p_subtitle = doc.add_paragraph()
                    p_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_subtitle.paragraph_format.space_after = Pt(12)
                    _add_run(p_subtitle, nt_judul.upper(), bold=True, size=Pt(12))

                    doc.add_paragraph()

                    p_info_hdr = doc.add_paragraph()
                    p_info_hdr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_info_hdr.paragraph_format.left_indent = SECTION_INDENT
                    p_info_hdr.paragraph_format.first_line_indent = -SECTION_INDENT
                    p_info_hdr.paragraph_format.space_after = Pt(6)
                    _add_run(p_info_hdr, "Informasi Rapat", bold=True)

                    row0_labels = ["Hari/Tanggal"]
                    row0_values = [nt_tanggal]
                    if nt_meeting_id.strip():
                        row0_labels.append("Meeting ID")
                        row0_values.append(nt_meeting_id)
                    if nt_password.strip():
                        row0_labels.append("Password")
                        row0_values.append(nt_password)
                    if nt_tempat.strip():
                        row0_labels.append("Tempat")
                        row0_values.append(nt_tempat)

                    peserta_list = [p.strip() for p in nt_peserta.strip().split("\n") if p.strip()] if nt_peserta.strip() else []

                    table = doc.add_table(rows=3, cols=3)
                    table.alignment = WD_TABLE_ALIGNMENT.LEFT

                    c0 = table.cell(0, 0)
                    c1 = table.cell(0, 1)
                    c2 = table.cell(0, 2)
                    c0.text = ""
                    c1.text = ""
                    c2.text = ""
                    for li, lbl in enumerate(row0_labels):
                        if li == 0:
                            _add_run(c0.paragraphs[0], lbl)
                            _add_run(c1.paragraphs[0], ":")
                            _add_run(c2.paragraphs[0], row0_values[li])
                        else:
                            p_l = c0.add_paragraph()
                            _add_run(p_l, lbl)
                            p_c = c1.add_paragraph()
                            _add_run(p_c, ":")
                            p_v = c2.add_paragraph()
                            _add_run(p_v, row0_values[li])

                    c0_p = table.cell(1, 0)
                    c1_p = table.cell(1, 1)
                    c2_p = table.cell(1, 2)
                    c0_p.text = ""
                    c1_p.text = ""
                    c2_p.text = ""
                    _add_run(c0_p.paragraphs[0], "Pimpinan Rapat")
                    _add_run(c1_p.paragraphs[0], ":")
                    _add_run(c2_p.paragraphs[0], nt_pimpinan if nt_pimpinan.strip() else "-")

                    c0_ps = table.cell(2, 0)
                    c1_ps = table.cell(2, 1)
                    c2_ps = table.cell(2, 2)
                    c0_ps.text = ""
                    c1_ps.text = ""
                    c2_ps.text = ""
                    _add_run(c0_ps.paragraphs[0], "Peserta Rapat")
                    _add_run(c1_ps.paragraphs[0], ":")
                    if peserta_list:
                        _add_run(c2_ps.paragraphs[0], peserta_list[0])
                        for pi_idx in range(1, len(peserta_list)):
                            p_peserta = c2_ps.add_paragraph()
                            _add_run(p_peserta, peserta_list[pi_idx])
                    else:
                        _add_run(c2_ps.paragraphs[0], "-")

                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                paragraph.paragraph_format.space_after = Pt(0)
                                paragraph.paragraph_format.space_before = Pt(0)
                                paragraph.paragraph_format.line_spacing = 1.15

                    _remove_table_borders(table)

                    doc.add_paragraph()
                    doc.add_paragraph()

                    if nt_pendahuluan.strip():
                        p_pend_hdr = doc.add_paragraph()
                        p_pend_hdr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_pend_hdr.paragraph_format.left_indent = SECTION_INDENT
                        p_pend_hdr.paragraph_format.first_line_indent = -SECTION_INDENT
                        p_pend_hdr.paragraph_format.space_after = Pt(6)
                        _add_run(p_pend_hdr, "Pendahuluan", bold=True)

                        for para_text in nt_pendahuluan.strip().split("\n"):
                            if para_text.strip():
                                p_pend = doc.add_paragraph()
                                p_pend.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                p_pend.paragraph_format.left_indent = BODY_INDENT
                                p_pend.paragraph_format.space_after = Pt(6)
                                _add_run(p_pend, para_text.strip())

                        doc.add_paragraph()

                    has_pembahasan = any(pb["speaker"].strip() or pb["content"].strip() for pb in nt_pembahasan_list)
                    if has_pembahasan:
                        p_pemb_hdr = doc.add_paragraph()
                        p_pemb_hdr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_pemb_hdr.paragraph_format.left_indent = SECTION_INDENT
                        p_pemb_hdr.paragraph_format.first_line_indent = -SECTION_INDENT
                        p_pemb_hdr.paragraph_format.space_after = Pt(6)
                        _add_run(p_pemb_hdr, "Pembahasan", bold=True)

                        for pb in nt_pembahasan_list:
                            if pb["speaker"].strip() or pb["content"].strip():
                                if pb["speaker"].strip():
                                    p_spk = doc.add_paragraph()
                                    p_spk.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                    p_spk.paragraph_format.left_indent = BODY_INDENT
                                    p_spk.paragraph_format.space_before = Pt(6)
                                    _add_run(p_spk, pb["speaker"].strip(), bold=True)

                                if pb["content"].strip():
                                    for line in pb["content"].strip().split("\n"):
                                        if line.strip():
                                            p_ct = doc.add_paragraph()
                                            p_ct.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                                            p_ct.paragraph_format.left_indent = BODY_INDENT
                                            p_ct.paragraph_format.space_after = Pt(6)
                                            _add_run(p_ct, line.strip())

                                doc.add_paragraph()

                    if nt_fotos:
                        p_dok_hdr = doc.add_paragraph()
                        p_dok_hdr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_dok_hdr.paragraph_format.left_indent = SECTION_INDENT
                        p_dok_hdr.paragraph_format.first_line_indent = -SECTION_INDENT
                        p_dok_hdr.paragraph_format.space_after = Pt(6)
                        _add_run(p_dok_hdr, "D. Dokumentasi", bold=True)

                        doc.add_paragraph()

                        foto_pairs = []
                        for i in range(0, len(nt_fotos), 2):
                            pair = [nt_fotos[i]]
                            if i + 1 < len(nt_fotos):
                                pair.append(nt_fotos[i + 1])
                            foto_pairs.append(pair)

                        for pair in foto_pairs:
                            tbl_foto = doc.add_table(rows=1, cols=len(pair))
                            tbl_foto.alignment = WD_TABLE_ALIGNMENT.CENTER

                            for ci, foto in enumerate(pair):
                                cell = tbl_foto.cell(0, ci)
                                cell.text = ""
                                p_img = cell.paragraphs[0]
                                p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                foto.seek(0)
                                img_width = Cm(7) if len(pair) == 2 else Cm(12)
                                p_img.add_run().add_picture(foto, width=img_width)

                            _remove_table_borders(tbl_foto)
                            doc.add_paragraph()

                    if nt_kesimpulan.strip():
                        p_kes_hdr = doc.add_paragraph()
                        p_kes_hdr.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        p_kes_hdr.paragraph_format.left_indent = SECTION_INDENT
                        p_kes_hdr.paragraph_format.first_line_indent = -SECTION_INDENT
                        p_kes_hdr.paragraph_format.space_after = Pt(6)
                        _add_run(p_kes_hdr, "E. Kesimpulan", bold=True)

                        kesimpulan_lines = [k.strip() for k in nt_kesimpulan.strip().split("\n") if k.strip()]
                        for ki, kline in enumerate(kesimpulan_lines):
                            p_kes = doc.add_paragraph()
                            p_kes.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_kes.paragraph_format.left_indent = BODY_INDENT
                            p_kes.paragraph_format.space_after = Pt(6)
                            run_kes_text = _add_run(p_kes, f"{ki + 1}. {kline}")

                    doc.add_paragraph()
                    doc.add_paragraph()
                    doc.add_paragraph()

                    tbl_notulis = doc.add_table(rows=1, cols=1)
                    tbl_notulis.alignment = WD_TABLE_ALIGNMENT.CENTER
                    cell_nt = tbl_notulis.cell(0, 0)
                    cell_nt.text = ""
                    p_nt_label = cell_nt.paragraphs[0]
                    p_nt_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(p_nt_label, "Notulis")

                    cell_nt.add_paragraph()
                    cell_nt.add_paragraph()
                    cell_nt.add_paragraph()

                    p_nt_name = cell_nt.add_paragraph()
                    p_nt_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    _add_run(p_nt_name, nt_notulis if nt_notulis.strip() else "___________________")

                    _remove_table_borders(tbl_notulis)

                    out_notulen = io.BytesIO()
                    doc.save(out_notulen)
                    out_notulen.seek(0)

                    import re as _re
                    tanggal_clean = _re.sub(r'[^\w\s-]', '', nt_tanggal).strip().replace(" ", "_")
                    judul_clean = _re.sub(r'[^\w\s-]', '', nt_judul).strip().replace(" ", "_")[:50]
                    filename = f"Notulen_{tanggal_clean}_{judul_clean}.docx"

                    st.session_state.notulen_docx = out_notulen.getvalue()
                    st.session_state.notulen_filename = filename
                    st.rerun()

                except Exception as e:
                    st.error(f"❌ Gagal membuat notulen: {str(e)}")
                    import traceback
                    st.code(traceback.format_exc())

    if st.session_state.notulen_docx is not None:
        st.success("✅ Notulen berhasil dibuat!")
        st.download_button(
            label="📥 Download Notulen (.docx)",
            data=st.session_state.notulen_docx,
            file_name=st.session_state.notulen_filename or "notulen_rapat.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


with tab_laporan_magang:
    st.markdown("### 🎓 Generator Laporan Akhir Peserta Magang BPOM")
    st.markdown("Isi form di bawah, lalu klik **Generate** untuk mengunduh laporan dalam format **Word (.docx)** sesuai format resmi BPOM (Lampiran ND PPSDM POM No. HM.03.04.9.05.26.542).")

    if 'lm_docx' not in st.session_state:
        st.session_state.lm_docx = None
    if 'lm_filename' not in st.session_state:
        st.session_state.lm_filename = None
    if 'lm_iki_count' not in st.session_state:
        st.session_state.lm_iki_count = 3
    if 'lm_kendala_count' not in st.session_state:
        st.session_state.lm_kendala_count = 3
    if 'lm_output_count' not in st.session_state:
        st.session_state.lm_output_count = 3

    # ── A. Identitas ────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### A. Identitas Peserta")
    col_a1, col_a2 = st.columns(2)
    with col_a1:
        lm_nama       = st.text_input("1. Nama Peserta", key="lm_nama")
        lm_pt_prodi   = st.text_input("2. Perguruan Tinggi / Prodi", key="lm_pt_prodi")
        lm_tahun_lulus = st.text_input("3. Tahun Lulus", key="lm_tahun_lulus")
    with col_a2:
        lm_unit_kerja  = st.text_input("4. Unit Kerja Penempatan", key="lm_unit_kerja")
        lm_jabatan     = st.text_input("5. Jabatan / Posisi Magang", key="lm_jabatan")
        lm_mentor      = st.text_input("6. Nama Mentor", key="lm_mentor")

    # ── B. Durasi & Ringkasan ────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### B. Durasi dan Ringkasan Magang")
    col_b1, col_b2, col_b3 = st.columns([1, 2, 2])
    with col_b1:
        lm_durasi_bulan = st.number_input("Durasi (bulan)", min_value=1, max_value=12, value=6, key="lm_durasi_bulan")
    with col_b2:
        lm_tgl_mulai  = st.text_input("Tanggal Mulai (mis. 1 November 2025)", key="lm_tgl_mulai")
    with col_b3:
        lm_tgl_selesai = st.text_input("Tanggal Selesai (mis. 30 April 2026)", key="lm_tgl_selesai")
    lm_ringkasan = st.text_area("Ringkasan Pelaksanaan Magang (1 paragraf — fokus pekerjaan utama & kontribusi terpenting):",
                                 height=130, key="lm_ringkasan",
                                 placeholder="Selama periode magang, peserta berfokus pada … Kontribusi utama yang dihasilkan adalah …")

    # ── C1. IKI/SKI ─────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### C. Perencanaan Kinerja Awal Magang")
    st.markdown("**C1. Rencana Kinerja Magang (IKI/SKI)**")
    col_ciki_btns = st.columns([1, 1, 6])
    with col_ciki_btns[0]:
        if st.button("➕ Tambah IKI", key="lm_add_iki"):
            st.session_state.lm_iki_count += 1
            st.rerun()
    with col_ciki_btns[1]:
        if st.session_state.lm_iki_count > 1 and st.button("➖ Hapus", key="lm_rmv_iki"):
            st.session_state.lm_iki_count -= 1
            st.rerun()

    lm_iki_rows = []
    hdr_iki = st.columns([0.5, 3, 2, 2, 3])
    hdr_iki[0].markdown("**No**"); hdr_iki[1].markdown("**Sasaran/IKI (6 Bulan)**")
    hdr_iki[2].markdown("**Target Kuantitatif**"); hdr_iki[3].markdown("**Indikator Mutu**"); hdr_iki[4].markdown("**Output yang Direncanakan**")
    for i in range(st.session_state.lm_iki_count):
        ci = st.columns([0.5, 3, 2, 2, 3])
        ci[0].markdown(f"**{i+1}**")
        s = ci[1].text_input(" ", key=f"lm_iki_sasaran_{i}", label_visibility="collapsed", placeholder="Sasaran/IKI")
        t = ci[2].text_input(" ", key=f"lm_iki_target_{i}", label_visibility="collapsed", placeholder="Target")
        m = ci[3].text_input(" ", key=f"lm_iki_mutu_{i}", label_visibility="collapsed", placeholder="Mutu (opsional)")
        o = ci[4].text_input(" ", key=f"lm_iki_output_{i}", label_visibility="collapsed", placeholder="Output")
        lm_iki_rows.append({"no": i+1, "sasaran": s, "target": t, "mutu": m, "output": o})

    # ── C2. Target Bulanan ───────────────────────────────────────────────────
    st.markdown("**C2. Target Kinerja Bulanan**")
    lm_bulanan_rows = []
    hdr_b = st.columns([0.5, 0.7, 3, 3, 3])
    hdr_b[0].markdown("**No**"); hdr_b[1].markdown("**Bulan**")
    hdr_b[2].markdown("**Target Utama Bulan Ini**"); hdr_b[3].markdown("**Output Bulanan**"); hdr_b[4].markdown("**Catatan**")
    for i in range(int(lm_durasi_bulan)):
        cb = st.columns([0.5, 0.7, 3, 3, 3])
        cb[0].markdown(f"**{i+1}**"); cb[1].markdown(f"**{i+1}**")
        tm = cb[2].text_input(" ", key=f"lm_bln_target_{i}", label_visibility="collapsed", placeholder=f"Target bulan {i+1}")
        ob = cb[3].text_input(" ", key=f"lm_bln_output_{i}", label_visibility="collapsed", placeholder="Output")
        ct = cb[4].text_input(" ", key=f"lm_bln_catatan_{i}", label_visibility="collapsed", placeholder="Catatan")
        lm_bulanan_rows.append({"bulan": i+1, "target": tm, "output": ob, "catatan": ct})

    # ── D. Gambaran Proses ───────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### D. Gambaran Proses Pelaksanaan Magang")
    st.markdown("**D1. Aktivitas Utama per Bulan**")
    lm_aktivitas = []
    for i in range(int(lm_durasi_bulan)):
        ak = st.text_input(f"Bulan {i+1}:", key=f"lm_aktivitas_{i}",
                            placeholder=f"Narasi singkat aktivitas bulan {i+1}…")
        lm_aktivitas.append(ak)

    st.markdown("**D2. Dukungan Pembelajaran**")
    lm_mentoring  = st.text_area("1. Pembimbingan/mentoring yang diterima (frekuensi, bentuk):",
                                  height=80, key="lm_mentoring")
    lm_pelatihan  = st.text_area("2. Pelatihan/briefing/rapat/kunjungan kerja yang diikuti:",
                                  height=80, key="lm_pelatihan")

    st.markdown("**D3. Kendala & Solusi**")
    col_knd_btns = st.columns([1, 1, 6])
    with col_knd_btns[0]:
        if st.button("➕ Tambah Kendala", key="lm_add_knd"):
            st.session_state.lm_kendala_count += 1
            st.rerun()
    with col_knd_btns[1]:
        if st.session_state.lm_kendala_count > 1 and st.button("➖ Hapus", key="lm_rmv_knd"):
            st.session_state.lm_kendala_count -= 1
            st.rerun()

    lm_kendala_rows = []
    hdr_k = st.columns([0.5, 3, 2.5, 3, 2])
    hdr_k[0].markdown("**No**"); hdr_k[1].markdown("**Kendala**")
    hdr_k[2].markdown("**Dampak**"); hdr_k[3].markdown("**Solusi/Mitigasi**"); hdr_k[4].markdown("**Status**")
    for i in range(st.session_state.lm_kendala_count):
        ck = st.columns([0.5, 3, 2.5, 3, 2])
        ck[0].markdown(f"**{i+1}**")
        kd = ck[1].text_input(" ", key=f"lm_knd_knd_{i}", label_visibility="collapsed", placeholder="Kendala")
        dp = ck[2].text_input(" ", key=f"lm_knd_dmp_{i}", label_visibility="collapsed", placeholder="Dampak")
        sl = ck[3].text_input(" ", key=f"lm_knd_sol_{i}", label_visibility="collapsed", placeholder="Solusi")
        st_ = ck[4].selectbox(" ", ["Selesai", "Dalam Proses", "Belum Selesai"], key=f"lm_knd_st_{i}", label_visibility="collapsed")
        lm_kendala_rows.append({"no": i+1, "kendala": kd, "dampak": dp, "solusi": sl, "status": st_})

    # ── E. Hasil Kinerja Akhir ───────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### E. Hasil Kinerja Akhir Magang")
    st.markdown("**E1. Rekap Capaian Akhir vs Rencana** *(selaraskan dengan IKI/SKI di C1)*")
    lm_e1_rows = []
    hdr_e1 = st.columns([0.5, 3, 2, 2, 1.5, 3, 2])
    hdr_e1[0].markdown("**No**"); hdr_e1[1].markdown("**Sasaran/IKI**")
    hdr_e1[2].markdown("**Target**"); hdr_e1[3].markdown("**Realisasi**")
    hdr_e1[4].markdown("**% Capaian**"); hdr_e1[5].markdown("**Bukti Output**"); hdr_e1[6].markdown("**Catatan Mentor**")
    for i in range(st.session_state.lm_iki_count):
        default_sas = st.session_state.get(f"lm_iki_sasaran_{i}", "")
        default_tgt = st.session_state.get(f"lm_iki_target_{i}", "")
        ce = st.columns([0.5, 3, 2, 2, 1.5, 3, 2])
        ce[0].markdown(f"**{i+1}**")
        es = ce[1].text_input(" ", key=f"lm_e1_sas_{i}", label_visibility="collapsed", value=default_sas, placeholder="Sasaran/IKI")
        et = ce[2].text_input(" ", key=f"lm_e1_tgt_{i}", label_visibility="collapsed", value=default_tgt, placeholder="Target")
        er = ce[3].text_input(" ", key=f"lm_e1_real_{i}", label_visibility="collapsed", placeholder="Realisasi")
        ep = ce[4].text_input(" ", key=f"lm_e1_pct_{i}", label_visibility="collapsed", placeholder="%")
        eb = ce[5].text_input(" ", key=f"lm_e1_bukti_{i}", label_visibility="collapsed", placeholder="Link/nama file")
        ec = ce[6].text_input(" ", key=f"lm_e1_cat_{i}", label_visibility="collapsed", placeholder="Catatan")
        lm_e1_rows.append({"no": i+1, "sasaran": es, "target": et, "realisasi": er, "persen": ep, "bukti": eb, "catatan": ec})

    st.markdown("**E2. Output Utama yang Dihasilkan** *(tuliskan 3–5 output paling penting)*")
    col_out_btns = st.columns([1, 1, 6])
    with col_out_btns[0]:
        if st.button("➕ Tambah Output", key="lm_add_out"):
            st.session_state.lm_output_count = min(st.session_state.lm_output_count + 1, 5)
            st.rerun()
    with col_out_btns[1]:
        if st.session_state.lm_output_count > 1 and st.button("➖ Hapus", key="lm_rmv_out"):
            st.session_state.lm_output_count -= 1
            st.rerun()
    lm_output_rows = []
    for i in range(st.session_state.lm_output_count):
        ov = st.text_input(f"Output {i+1}:", key=f"lm_output_{i}",
                            placeholder=f"Output {i+1}: … (fungsi/manfaat)")
        lm_output_rows.append(ov)

    # ── F. Pembelajaran ──────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### F. Pembelajaran, Manfaat, dan Rekomendasi")
    lm_f1 = st.text_area("1. Pembelajaran teknis yang diperoleh:", height=80, key="lm_f1")
    lm_f2 = st.text_area("2. Pembelajaran soft skills/etos kerja:", height=80, key="lm_f2")
    lm_f3 = st.text_area("3. Manfaat magang bagi unit kerja:", height=80, key="lm_f3")
    lm_f4 = st.text_area("4. Rekomendasi perbaikan penyelenggaraan magang (batch berikutnya):", height=80, key="lm_f4")

    # ── G. Penutup ───────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### G. Penutup")
    lm_penutup = st.text_area("Paragraf penutup (isi sendiri atau biarkan kosong untuk teks default):",
                               height=100, key="lm_penutup",
                               placeholder="Biarkan kosong untuk menggunakan teks penutup default sesuai format.")

    # ── H. Lampiran ──────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### H. Lampiran")
    lm_lampiran_wajib   = st.text_area("Bukti output utama (wajib) — link drive / nama file / deskripsi:",
                                        height=80, key="lm_lampiran_wajib")
    lm_lampiran_opsional = st.text_area("Logbook ringkas mingguan/bulanan (opsional):",
                                         height=80, key="lm_lampiran_opsional")

    # ── Tanda Tangan ────────────────────────────────────────────────────────
    st.markdown("---")
    st.markdown("#### Tanda Tangan")
    col_ttd1, col_ttd2 = st.columns(2)
    with col_ttd1:
        lm_mentor_jabatan = st.text_input("Jabatan Mentor:", key="lm_mentor_jabatan", placeholder="Kepala Seksi …")
        lm_mentor_nama_ttd = st.text_input("Nama Mentor (tanda tangan):", key="lm_mentor_nama_ttd", value=st.session_state.get("lm_mentor", ""))
    with col_ttd2:
        lm_kota_ttd  = st.text_input("Kota:", key="lm_kota_ttd", placeholder="Jakarta")
        lm_tgl_ttd   = st.text_input("Tanggal TTD (mis. 20 Mei 2026):", key="lm_tgl_ttd")

    st.markdown("---")
    col_gen1, col_gen2, col_gen3 = st.columns([1, 2, 1])
    with col_gen2:
        lm_gen_btn = st.button("📄 Generate Laporan Magang (.docx)", key="lm_gen_btn", type="primary", use_container_width=True)
    with col_gen3:
        if st.button("🔄 Reset", key="lm_reset_btn"):
            for k in list(st.session_state.keys()):
                if k.startswith("lm_"):
                    del st.session_state[k]
            st.rerun()

    if lm_gen_btn:
        if not lm_nama.strip():
            st.error("❌ Nama peserta wajib diisi.")
        else:
            try:
                with st.spinner("Membuat dokumen laporan magang…"):
                    from docx import Document as LMDoc
                    from docx.shared import Pt as LMPt, Cm as LMCm, RGBColor as LMRgb
                    from docx.enum.text import WD_ALIGN_PARAGRAPH as LM_ALIGN
                    from docx.enum.table import WD_TABLE_ALIGNMENT as LM_TBL_ALIGN, WD_ALIGN_VERTICAL as LM_VA
                    from docx.oxml.ns import qn as lm_qn
                    from docx.oxml import OxmlElement as lm_oxe

                    doc = LMDoc()
                    sec = doc.sections[0]
                    sec.page_width  = LMCm(21)
                    sec.page_height = LMCm(29.7)
                    sec.left_margin = sec.right_margin = LMCm(3)
                    sec.top_margin  = sec.bottom_margin = LMCm(2.5)

                    N = 'Times New Roman'
                    SZ = 12

                    def _lm_p(text="", bold=False, italic=False, sz=SZ,
                              align=LM_ALIGN.JUSTIFY, sb=0, sa=6, li=0, fi=0, font=N):
                        p = doc.add_paragraph()
                        p.alignment = align
                        p.paragraph_format.space_before = LMPt(sb)
                        p.paragraph_format.space_after  = LMPt(sa)
                        if li: p.paragraph_format.left_indent       = LMCm(li)
                        if fi: p.paragraph_format.first_line_indent = LMCm(fi)
                        if text:
                            r = p.add_run(text)
                            r.bold = bold; r.italic = italic
                            r.font.name = font; r.font.size = LMPt(sz)
                        return p

                    def _lm_r(para, text, bold=False, italic=False, sz=SZ, font=N):
                        r = para.add_run(text)
                        r.bold = bold; r.italic = italic
                        r.font.name = font; r.font.size = LMPt(sz)
                        return r

                    def _lm_section(label, sa=4):
                        p = doc.add_paragraph()
                        p.alignment = LM_ALIGN.LEFT
                        p.paragraph_format.space_before = LMPt(10)
                        p.paragraph_format.space_after  = LMPt(sa)
                        r = p.add_run(label)
                        r.bold = True; r.font.name = N; r.font.size = LMPt(SZ)
                        return p

                    def _lm_field(label, value, indent=0.75):
                        p = doc.add_paragraph()
                        p.alignment = LM_ALIGN.LEFT
                        p.paragraph_format.space_before = LMPt(0)
                        p.paragraph_format.space_after  = LMPt(3)
                        p.paragraph_format.left_indent  = LMCm(indent)
                        _lm_r(p, label); _lm_r(p, f"\t: {value if value else '………………………………………'}")
                        return p

                    def _shade_row(row, hex_color="D9E1F2"):
                        for cell in row.cells:
                            tc = cell._tc
                            tcPr = tc.get_or_add_tcPr()
                            shd = lm_oxe("w:shd")
                            shd.set(lm_qn("w:val"), "clear")
                            shd.set(lm_qn("w:color"), "auto")
                            shd.set(lm_qn("w:fill"), hex_color)
                            tcPr.append(shd)

                    def _set_cell(cell, text, bold=False, align=LM_ALIGN.CENTER, sz=11):
                        cell.text = ""
                        p = cell.paragraphs[0]
                        p.alignment = align
                        p.paragraph_format.space_before = LMPt(2)
                        p.paragraph_format.space_after  = LMPt(2)
                        r = p.add_run(text)
                        r.bold = bold; r.font.name = N; r.font.size = LMPt(sz)

                    def _set_col_width(table, col_idx, width_cm):
                        for row in table.rows:
                            row.cells[col_idx].width = LMCm(width_cm)

                    def _make_table_borders(table):
                        tbl = table._tbl
                        tblPr = tbl.tblPr if tbl.tblPr is not None else lm_oxe("w:tblPr")
                        tblBorders = lm_oxe("w:tblBorders")
                        for side in ("top","left","bottom","right","insideH","insideV"):
                            b = lm_oxe(f"w:{side}")
                            b.set(lm_qn("w:val"),   "single")
                            b.set(lm_qn("w:sz"),    "4")
                            b.set(lm_qn("w:space"), "0")
                            b.set(lm_qn("w:color"), "000000")
                            tblBorders.append(b)
                        tblPr.append(tblBorders)

                    # ─── HEADER ────────────────────────────────────────────
                    _lm_p("BADAN PENGAWAS OBAT DAN MAKANAN", bold=True, sz=12,
                          align=LM_ALIGN.CENTER, sb=0, sa=2)
                    _lm_p("PUSAT PENGEMBANGAN SUMBER DAYA MANUSIA", bold=True, sz=12,
                          align=LM_ALIGN.CENTER, sb=0, sa=2)
                    _lm_p("PENGAWASAN OBAT DAN MAKANAN", bold=True, sz=12,
                          align=LM_ALIGN.CENTER, sb=0, sa=8)

                    p_judul = doc.add_paragraph()
                    p_judul.alignment = LM_ALIGN.CENTER
                    p_judul.paragraph_format.space_before = LMPt(0)
                    p_judul.paragraph_format.space_after  = LMPt(14)
                    r_judul = p_judul.add_run("FORMAT LAPORAN AKHIR PESERTA MAGANG BPOM")
                    r_judul.bold = True; r_judul.underline = True
                    r_judul.font.name = N; r_judul.font.size = LMPt(13)

                    # ─── A. Identitas ──────────────────────────────────────
                    _lm_section("A. Identitas Peserta")
                    fields_a = [
                        ("1. Nama",                     lm_nama),
                        ("2. Perguruan Tinggi / Prodi",  lm_pt_prodi),
                        ("3. Tahun Lulus",               lm_tahun_lulus),
                        ("4. Unit Kerja Penempatan",     lm_unit_kerja),
                        ("5. Jabatan/Posisi Magang",     lm_jabatan),
                        ("6. Nama Mentor",               lm_mentor),
                    ]
                    for lbl, val in fields_a:
                        _lm_field(lbl, val)

                    # ─── B. Durasi & Ringkasan ─────────────────────────────
                    _lm_section("B. Durasi dan Ringkasan Magang")
                    _lm_field(f"1. Durasi Magang",
                              f"{int(lm_durasi_bulan)} bulan ({lm_tgl_mulai} s.d. {lm_tgl_selesai})")
                    _lm_p("2. Ringkasan Pelaksanaan Magang:", bold=False, sb=4, sa=2, li=0.75)
                    _lm_p(lm_ringkasan if lm_ringkasan.strip() else "…",
                          sa=6, li=1.25, align=LM_ALIGN.JUSTIFY)

                    # ─── C1. IKI/SKI ──────────────────────────────────────
                    _lm_section("C. Perencanaan Kinerja Awal Magang (Input Utama)")
                    _lm_p("C1. Rencana Kinerja Magang (IKI/SKI)", bold=True, sb=0, sa=4)
                    _lm_p("Tabel 1. Rencana Kinerja Awal Magang (IKI/SKI)",
                          italic=True, align=LM_ALIGN.CENTER, sb=0, sa=4)

                    tbl1 = doc.add_table(rows=1 + len(lm_iki_rows), cols=5)
                    _make_table_borders(tbl1)
                    tbl1.alignment = LM_TBL_ALIGN.CENTER
                    hdr1 = tbl1.rows[0]
                    _shade_row(hdr1)
                    for ci, hd in enumerate(["No", "Sasaran/IKI (6 Bulan)", "Target Kuantitatif",
                                              "Indikator Mutu (jika ada)", "Output yang Direncanakan"]):
                        _set_cell(hdr1.cells[ci], hd, bold=True)
                    for ri, row in enumerate(lm_iki_rows):
                        tr = tbl1.rows[ri + 1]
                        _set_cell(tr.cells[0], str(row["no"]))
                        _set_cell(tr.cells[1], row["sasaran"], align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[2], row["target"],  align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[3], row["mutu"],    align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[4], row["output"],  align=LM_ALIGN.LEFT, sz=11)

                    doc.add_paragraph()

                    # ─── C2. Target Bulanan ────────────────────────────────
                    _lm_p("C2. Target Bulanan (Ringkas)", bold=True, sb=6, sa=4)
                    _lm_p("Tabel 2. Target Kinerja Bulanan",
                          italic=True, align=LM_ALIGN.CENTER, sb=0, sa=4)

                    tbl2 = doc.add_table(rows=1 + len(lm_bulanan_rows), cols=5)
                    _make_table_borders(tbl2)
                    tbl2.alignment = LM_TBL_ALIGN.CENTER
                    hdr2 = tbl2.rows[0]
                    _shade_row(hdr2)
                    for ci, hd in enumerate(["No", "Bulan Ke", "Target Utama Bulan Ini (turunan IKI/SKI)",
                                              "Output Bulanan", "Catatan"]):
                        _set_cell(hdr2.cells[ci], hd, bold=True)
                    for ri, row in enumerate(lm_bulanan_rows):
                        tr = tbl2.rows[ri + 1]
                        _set_cell(tr.cells[0], str(row["bulan"]))
                        _set_cell(tr.cells[1], str(row["bulan"]))
                        _set_cell(tr.cells[2], row["target"],  align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[3], row["output"],  align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[4], row["catatan"], align=LM_ALIGN.LEFT, sz=11)

                    doc.add_paragraph()

                    # ─── D. Gambaran Proses ────────────────────────────────
                    _lm_section("D. Gambaran Proses Pelaksanaan Magang")
                    _lm_p("D1. Aktivitas Utama per Bulan (narasi singkat)", bold=True, sb=0, sa=4)
                    for i, ak in enumerate(lm_aktivitas):
                        p_ak = doc.add_paragraph()
                        p_ak.alignment = LM_ALIGN.JUSTIFY
                        p_ak.paragraph_format.space_before = LMPt(0)
                        p_ak.paragraph_format.space_after  = LMPt(4)
                        p_ak.paragraph_format.left_indent  = LMCm(0.75)
                        _lm_r(p_ak, f"{i+1}. Bulan {i+1}: ", bold=True)
                        _lm_r(p_ak, ak if ak.strip() else "…")

                    _lm_p("D2. Dukungan Pembelajaran", bold=True, sb=8, sa=4)
                    p_d2a = doc.add_paragraph()
                    p_d2a.alignment = LM_ALIGN.JUSTIFY
                    p_d2a.paragraph_format.left_indent = LMCm(0.75)
                    p_d2a.paragraph_format.space_after = LMPt(4)
                    _lm_r(p_d2a, "1. Pembimbingan/mentoring: ", bold=True)
                    _lm_r(p_d2a, lm_mentoring if lm_mentoring.strip() else "…")

                    p_d2b = doc.add_paragraph()
                    p_d2b.alignment = LM_ALIGN.JUSTIFY
                    p_d2b.paragraph_format.left_indent = LMCm(0.75)
                    p_d2b.paragraph_format.space_after = LMPt(6)
                    _lm_r(p_d2b, "2. Pelatihan/briefing/rapat/kunjungan kerja: ", bold=True)
                    _lm_r(p_d2b, lm_pelatihan if lm_pelatihan.strip() else "…")

                    _lm_p("D3. Kendala & Solusi", bold=True, sb=4, sa=4)
                    _lm_p("Tabel 3. Kendala dan Mitigasi",
                          italic=True, align=LM_ALIGN.CENTER, sb=0, sa=4)

                    tbl3 = doc.add_table(rows=1 + len(lm_kendala_rows), cols=5)
                    _make_table_borders(tbl3)
                    tbl3.alignment = LM_TBL_ALIGN.CENTER
                    hdr3 = tbl3.rows[0]
                    _shade_row(hdr3)
                    for ci, hd in enumerate(["No", "Kendala", "Dampak", "Solusi/Mitigasi", "Status"]):
                        _set_cell(hdr3.cells[ci], hd, bold=True)
                    for ri, row in enumerate(lm_kendala_rows):
                        tr = tbl3.rows[ri + 1]
                        _set_cell(tr.cells[0], str(row["no"]))
                        _set_cell(tr.cells[1], row["kendala"], align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[2], row["dampak"],  align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[3], row["solusi"],  align=LM_ALIGN.LEFT, sz=11)
                        _set_cell(tr.cells[4], row["status"],  sz=11)

                    doc.add_paragraph()

                    # ─── E. Hasil Kinerja ──────────────────────────────────
                    _lm_section("E. Hasil Kinerja Akhir Magang (Output Akhir)")
                    _lm_p("E1. Rekap Capaian Akhir vs Rencana (Akuntabilitas)", bold=True, sb=0, sa=4)
                    _lm_p("Tabel 4. Matriks Akuntabilitas: Rencana vs Realisasi",
                          italic=True, align=LM_ALIGN.CENTER, sb=0, sa=4)

                    tbl4 = doc.add_table(rows=1 + len(lm_e1_rows), cols=7)
                    _make_table_borders(tbl4)
                    tbl4.alignment = LM_TBL_ALIGN.CENTER
                    hdr4 = tbl4.rows[0]
                    _shade_row(hdr4)
                    for ci, hd in enumerate(["No", "Sasaran/IKI\n(Rencana Awal)", "Target\n(Rencana)",
                                              "Realisasi\n(Capaian)", "% Capaian",
                                              "Bukti Output\n(tautan/nama file)", "Catatan\nMentor"]):
                        _set_cell(hdr4.cells[ci], hd, bold=True, sz=10)
                    for ri, row in enumerate(lm_e1_rows):
                        tr = tbl4.rows[ri + 1]
                        _set_cell(tr.cells[0], str(row["no"]), sz=11)
                        _set_cell(tr.cells[1], row["sasaran"],  align=LM_ALIGN.LEFT, sz=10)
                        _set_cell(tr.cells[2], row["target"],   align=LM_ALIGN.LEFT, sz=10)
                        _set_cell(tr.cells[3], row["realisasi"],align=LM_ALIGN.LEFT, sz=10)
                        _set_cell(tr.cells[4], row["persen"],   sz=10)
                        _set_cell(tr.cells[5], row["bukti"],    align=LM_ALIGN.LEFT, sz=10)
                        _set_cell(tr.cells[6], row["catatan"],  align=LM_ALIGN.LEFT, sz=10)

                    doc.add_paragraph()

                    _lm_p("E2. Output Utama yang Dihasilkan (deskripsi singkat)", bold=True, sb=6, sa=4)
                    for i, ov in enumerate(lm_output_rows):
                        p_ov = doc.add_paragraph()
                        p_ov.alignment = LM_ALIGN.JUSTIFY
                        p_ov.paragraph_format.left_indent = LMCm(0.75)
                        p_ov.paragraph_format.space_after = LMPt(3)
                        _lm_r(p_ov, f"{i+1}. ", bold=True)
                        _lm_r(p_ov, ov if ov.strip() else "…")

                    # ─── F. Pembelajaran ───────────────────────────────────
                    _lm_section("F. Pembelajaran, Manfaat, dan Rekomendasi")
                    for i, (lbl, val) in enumerate([
                        ("1. Pembelajaran teknis yang diperoleh", lm_f1),
                        ("2. Pembelajaran soft skills/etos kerja", lm_f2),
                        ("3. Manfaat magang bagi unit kerja",       lm_f3),
                        ("4. Rekomendasi perbaikan penyelenggaraan magang (batch berikutnya)", lm_f4),
                    ], 1):
                        p_f = doc.add_paragraph()
                        p_f.alignment = LM_ALIGN.JUSTIFY
                        p_f.paragraph_format.left_indent = LMCm(0.75)
                        p_f.paragraph_format.space_after = LMPt(4)
                        _lm_r(p_f, f"{lbl}: ", bold=True)
                        _lm_r(p_f, val if val.strip() else "…")

                    # ─── G. Penutup ────────────────────────────────────────
                    _lm_section("G. Penutup")
                    penutup_default = (
                        f"Demikian laporan akhir pelaksanaan magang ini disusun sebagai bentuk "
                        f"pertanggungjawaban atas pelaksanaan Program Pemagangan Lulusan Perguruan Tinggi "
                        f"di {lm_unit_kerja if lm_unit_kerja.strip() else 'unit kerja terkait'}. "
                        f"Seluruh kegiatan magang telah dilaksanakan sesuai durasi yang ditetapkan, yaitu "
                        f"selama {int(lm_durasi_bulan)} bulan ({lm_tgl_mulai} s.d. {lm_tgl_selesai}). "
                        f"Output akhir yang dihasilkan telah mengacu pada rencana kinerja awal (IKI/SKI) "
                        f"sebagaimana tertuang dalam laporan ini. Peserta menyampaikan terima kasih yang "
                        f"sebesar-besarnya kepada {lm_mentor if lm_mentor.strip() else 'mentor'} selaku "
                        f"mentor magang serta seluruh pimpinan dan rekan-rekan di {lm_unit_kerja if lm_unit_kerja.strip() else 'unit kerja'} "
                        f"atas bimbingan, dukungan, dan kesempatan yang diberikan."
                    )
                    _lm_p(lm_penutup if lm_penutup.strip() else penutup_default,
                          align=LM_ALIGN.JUSTIFY, li=0.75)

                    # ─── H. Lampiran ───────────────────────────────────────
                    _lm_section("H. Lampiran")
                    p_hw = doc.add_paragraph()
                    p_hw.paragraph_format.left_indent = LMCm(0.75)
                    p_hw.paragraph_format.space_after = LMPt(3)
                    _lm_r(p_hw, "(Wajib) ", bold=True)
                    _lm_r(p_hw, lm_lampiran_wajib if lm_lampiran_wajib.strip() else "Bukti output utama (tautan drive/nama file/screenshot)")
                    if lm_lampiran_opsional.strip():
                        p_ho = doc.add_paragraph()
                        p_ho.paragraph_format.left_indent = LMCm(0.75)
                        p_ho.paragraph_format.space_after = LMPt(3)
                        _lm_r(p_ho, "(Opsional) ", bold=True)
                        _lm_r(p_ho, lm_lampiran_opsional)

                    # ─── Tanda Tangan ──────────────────────────────────────
                    doc.add_paragraph()
                    doc.add_paragraph()

                    tbl_ttd = doc.add_table(rows=1, cols=2)
                    tbl_ttd.alignment = LM_TBL_ALIGN.CENTER
                    from docx.oxml import OxmlElement as _nob
                    from docx.oxml.ns import qn as _nqn
                    tblPr = tbl_ttd._tbl.tblPr
                    tblBorders2 = _nob("w:tblBorders")
                    for side in ("top","left","bottom","right","insideH","insideV"):
                        b = _nob(f"w:{side}"); b.set(_nqn("w:val"),"none"); tblBorders2.append(b)
                    tblPr.append(tblBorders2)

                    cell_l = tbl_ttd.cell(0, 0)
                    cell_r = tbl_ttd.cell(0, 1)

                    def _ttd_block(cell, top_line, jabatan_or_kota, name, is_right=False):
                        cell.text = ""
                        al = LM_ALIGN.CENTER
                        for txt, bld, sz_ in [
                            (top_line, False, SZ),
                            (jabatan_or_kota, False, SZ),
                            ("", False, SZ),
                            ("", False, SZ),
                            ("", False, SZ),
                            (name, True, SZ),
                        ]:
                            p = cell.add_paragraph() if cell.paragraphs[0].text or len(cell.paragraphs)>1 else cell.paragraphs[0]
                            if cell.paragraphs[-1].text and len(cell.paragraphs)>1:
                                p = cell.add_paragraph()
                            p.alignment = al
                            p.paragraph_format.space_before = LMPt(0)
                            p.paragraph_format.space_after  = LMPt(2)
                            r = p.add_run(txt)
                            r.bold = bld; r.font.name = N; r.font.size = LMPt(sz_)

                    mentor_jabatan_str = lm_mentor_jabatan if lm_mentor_jabatan.strip() else "(Jabatan Mentor)"
                    mentor_nama_str    = lm_mentor_nama_ttd if lm_mentor_nama_ttd.strip() else lm_mentor if lm_mentor.strip() else "(Nama Mentor)"
                    peserta_nama_str   = lm_nama if lm_nama.strip() else "(Nama Peserta)"
                    kota_tgl_str       = f"{lm_kota_ttd}, {lm_tgl_ttd}" if lm_kota_ttd.strip() else f"………………, {lm_tgl_ttd}"

                    def _fill_ttd(cell, header, sub, name):
                        for p in cell.paragraphs:
                            p.clear()
                        def _add(text, bold=False):
                            p = cell.add_paragraph()
                            p.alignment = LM_ALIGN.CENTER
                            p.paragraph_format.space_before = LMPt(0)
                            p.paragraph_format.space_after  = LMPt(2)
                            r = p.add_run(text)
                            r.bold = bold; r.font.name = N; r.font.size = LMPt(SZ)
                        _add(header)
                        _add(sub)
                        _add(""); _add(""); _add("")
                        _add(f"({name})", bold=True)

                    _fill_ttd(cell_l, "Persetujuan Mentor Magang,", mentor_jabatan_str, mentor_nama_str)
                    _fill_ttd(cell_r, kota_tgl_str, "Peserta Magang", peserta_nama_str)

                    # ─── Save ──────────────────────────────────────────────
                    out_lm = io.BytesIO()
                    doc.save(out_lm)
                    out_lm.seek(0)

                    import re as _re_lm
                    nama_clean = _re_lm.sub(r'[^\w]', '_', lm_nama.strip())[:30]
                    st.session_state.lm_docx = out_lm.getvalue()
                    st.session_state.lm_filename = f"LaporanMagang_{nama_clean}.docx"
                    st.rerun()

            except Exception as e:
                st.error(f"❌ Gagal membuat dokumen: {str(e)}")
                import traceback
                st.code(traceback.format_exc())

    if st.session_state.lm_docx is not None:
        st.success("✅ Laporan berhasil dibuat!")
        st.download_button(
            label="📥 Download Laporan Magang (.docx)",
            data=st.session_state.lm_docx,
            file_name=st.session_state.lm_filename or "LaporanMagang.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

# ─────────────────────────────────────────────────────────
# TAB PDF EDITOR
# ─────────────────────────────────────────────────────────
with tab_pdf:
    st.markdown("### 📄 Editor PDF")
    st.markdown("Berbagai fitur edit PDF langsung di browser — tidak perlu install software apapun.")

    import io as _pdf_io
    from pypdf import PdfWriter as _PdfWriter, PdfReader as _PdfReader
    from reportlab.pdfgen import canvas as _rl_canvas
    from reportlab.lib.pagesizes import A4 as _rl_A4
    from reportlab.lib.colors import Color as _rl_Color

    pdf_fitur = st.selectbox(
        "Pilih fitur:",
        [
            "📎 Gabung PDF (Merge)",
            "✂️ Pisah / Ambil Halaman Tertentu",
            "🗑️ Hapus Halaman",
            "🔄 Putar Halaman",
            "💧 Tambah Watermark Teks",
            "🔢 Tambah Nomor Halaman",
            "🗜️ Kompres / Optimalkan PDF",
        ],
        key="pdf_fitur_select",
        label_visibility="visible"
    )

    st.markdown("---")

    # ── 1. GABUNG PDF ──────────────────────────────────────
    if pdf_fitur == "📎 Gabung PDF (Merge)":
        st.markdown("#### 📎 Gabung Beberapa PDF Menjadi Satu")
        st.markdown("Upload dua atau lebih file PDF. Urutan penggabungan sesuai urutan upload.")
        pdf_merge_files = st.file_uploader(
            "Upload file PDF (bisa lebih dari satu):",
            type=["pdf"],
            accept_multiple_files=True,
            key="pdf_merge_files"
        )
        if pdf_merge_files:
            st.info(f"📂 {len(pdf_merge_files)} file siap digabung: {', '.join(f.name for f in pdf_merge_files)}")
            pdf_merge_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="Gabungan_PDF", key="pdf_merge_outname")
            if st.button("📎 Gabung Sekarang", key="btn_pdf_merge", type="primary"):
                try:
                    _writer = _PdfWriter()
                    _total_pages = 0
                    for _f in pdf_merge_files:
                        _reader = _PdfReader(_pdf_io.BytesIO(_f.read()))
                        for _page in _reader.pages:
                            _writer.add_page(_page)
                            _total_pages += 1
                    _out_buf = _pdf_io.BytesIO()
                    _writer.write(_out_buf)
                    _out_buf.seek(0)
                    st.success(f"✅ Berhasil menggabungkan {len(pdf_merge_files)} file → {_total_pages} halaman total.")
                    st.download_button(
                        label="📥 Download PDF Gabungan",
                        data=_out_buf.getvalue(),
                        file_name=f"{pdf_merge_outname.strip() or 'Gabungan_PDF'}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 2. PISAH / AMBIL HALAMAN ───────────────────────────
    elif pdf_fitur == "✂️ Pisah / Ambil Halaman Tertentu":
        st.markdown("#### ✂️ Pisah / Ambil Halaman Tertentu")
        st.markdown("Upload PDF, lalu tentukan halaman mana saja yang ingin diambil.")
        pdf_split_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_split_file")
        if pdf_split_file:
            _r = _PdfReader(_pdf_io.BytesIO(pdf_split_file.read()))
            _n = len(_r.pages)
            st.info(f"📄 File ini memiliki **{_n} halaman**.")
            pdf_split_range = st.text_input(
                "Halaman yang ingin diambil (contoh: 1-3, 5, 7-10):",
                value=f"1-{_n}",
                key="pdf_split_range",
                help="Pisahkan dengan koma. Rentang menggunakan tanda -. Nomor halaman dimulai dari 1."
            )
            pdf_split_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="Halaman_Dipilih", key="pdf_split_outname")
            if st.button("✂️ Ambil Halaman", key="btn_pdf_split", type="primary"):
                try:
                    def _parse_ranges(s, max_page):
                        pages = set()
                        for part in s.split(","):
                            part = part.strip()
                            if "-" in part:
                                a, b = part.split("-", 1)
                                pages.update(range(int(a.strip()), int(b.strip()) + 1))
                            elif part:
                                pages.add(int(part))
                        return sorted(p for p in pages if 1 <= p <= max_page)

                    _pages_sel = _parse_ranges(pdf_split_range, _n)
                    if not _pages_sel:
                        st.error("❌ Tidak ada halaman valid yang dipilih.")
                    else:
                        pdf_split_file.seek(0)
                        _r2 = _PdfReader(_pdf_io.BytesIO(pdf_split_file.read()))
                        _w = _PdfWriter()
                        for _pg in _pages_sel:
                            _w.add_page(_r2.pages[_pg - 1])
                        _out = _pdf_io.BytesIO()
                        _w.write(_out)
                        _out.seek(0)
                        st.success(f"✅ Berhasil mengambil {len(_pages_sel)} halaman: {_pages_sel}")
                        st.download_button(
                            label="📥 Download Halaman Terpilih",
                            data=_out.getvalue(),
                            file_name=f"{pdf_split_outname.strip() or 'Halaman_Dipilih'}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 3. HAPUS HALAMAN ───────────────────────────────────
    elif pdf_fitur == "🗑️ Hapus Halaman":
        st.markdown("#### 🗑️ Hapus Halaman Tertentu dari PDF")
        pdf_del_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_del_file")
        if pdf_del_file:
            _r = _PdfReader(_pdf_io.BytesIO(pdf_del_file.read()))
            _n = len(_r.pages)
            st.info(f"📄 File ini memiliki **{_n} halaman**.")
            pdf_del_range = st.text_input(
                "Halaman yang ingin DIHAPUS (contoh: 2, 4-6, 10):",
                key="pdf_del_range",
                help="Nomor halaman dimulai dari 1."
            )
            pdf_del_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="PDF_Dihapus", key="pdf_del_outname")
            if st.button("🗑️ Hapus Halaman", key="btn_pdf_del", type="primary"):
                try:
                    def _parse_ranges_del(s, max_page):
                        pages = set()
                        for part in s.split(","):
                            part = part.strip()
                            if "-" in part:
                                a, b = part.split("-", 1)
                                pages.update(range(int(a.strip()), int(b.strip()) + 1))
                            elif part:
                                pages.add(int(part))
                        return set(p for p in pages if 1 <= p <= max_page)

                    _del_set = _parse_ranges_del(pdf_del_range, _n)
                    if not _del_set:
                        st.error("❌ Tidak ada halaman valid yang ditentukan.")
                    else:
                        pdf_del_file.seek(0)
                        _r2 = _PdfReader(_pdf_io.BytesIO(pdf_del_file.read()))
                        _w = _PdfWriter()
                        _kept = 0
                        for _i, _pg in enumerate(_r2.pages, 1):
                            if _i not in _del_set:
                                _w.add_page(_pg)
                                _kept += 1
                        if _kept == 0:
                            st.error("❌ Semua halaman dihapus — tidak ada yang tersisa.")
                        else:
                            _out = _pdf_io.BytesIO()
                            _w.write(_out)
                            _out.seek(0)
                            st.success(f"✅ Halaman {sorted(_del_set)} dihapus. Tersisa {_kept} halaman.")
                            st.download_button(
                                label="📥 Download PDF Hasil",
                                data=_out.getvalue(),
                                file_name=f"{pdf_del_outname.strip() or 'PDF_Dihapus'}.pdf",
                                mime="application/pdf",
                                use_container_width=True
                            )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 4. PUTAR HALAMAN ───────────────────────────────────
    elif pdf_fitur == "🔄 Putar Halaman":
        st.markdown("#### 🔄 Putar Halaman PDF")
        pdf_rot_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_rot_file")
        if pdf_rot_file:
            _r = _PdfReader(_pdf_io.BytesIO(pdf_rot_file.read()))
            _n = len(_r.pages)
            st.info(f"📄 File ini memiliki **{_n} halaman**.")
            col_rot1, col_rot2 = st.columns(2)
            with col_rot1:
                pdf_rot_range = st.text_input(
                    "Halaman yang ingin diputar (kosongkan = semua):",
                    value="",
                    key="pdf_rot_range",
                    help="Contoh: 1-3, 5. Kosongkan untuk putar semua halaman."
                )
            with col_rot2:
                pdf_rot_deg = st.selectbox(
                    "Derajat putaran:",
                    [90, 180, 270],
                    key="pdf_rot_deg",
                    help="90° = putar kanan, 270° = putar kiri, 180° = terbalik"
                )
            pdf_rot_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="PDF_Diputar", key="pdf_rot_outname")
            if st.button("🔄 Putar Sekarang", key="btn_pdf_rot", type="primary"):
                try:
                    pdf_rot_file.seek(0)
                    _r2 = _PdfReader(_pdf_io.BytesIO(pdf_rot_file.read()))

                    def _parse_ranges_rot(s, max_page):
                        if not s.strip():
                            return set(range(1, max_page + 1))
                        pages = set()
                        for part in s.split(","):
                            part = part.strip()
                            if "-" in part:
                                a, b = part.split("-", 1)
                                pages.update(range(int(a.strip()), int(b.strip()) + 1))
                            elif part:
                                pages.add(int(part))
                        return set(p for p in pages if 1 <= p <= max_page)

                    _rot_set = _parse_ranges_rot(pdf_rot_range, _n)
                    _w = _PdfWriter()
                    for _i, _pg in enumerate(_r2.pages, 1):
                        if _i in _rot_set:
                            _pg.rotate(pdf_rot_deg)
                        _w.add_page(_pg)
                    _out = _pdf_io.BytesIO()
                    _w.write(_out)
                    _out.seek(0)
                    _label = "semua halaman" if not pdf_rot_range.strip() else f"halaman {sorted(_rot_set)}"
                    st.success(f"✅ Berhasil memutar {_label} sebesar {pdf_rot_deg}°.")
                    st.download_button(
                        label="📥 Download PDF Hasil",
                        data=_out.getvalue(),
                        file_name=f"{pdf_rot_outname.strip() or 'PDF_Diputar'}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 5. WATERMARK TEKS ──────────────────────────────────
    elif pdf_fitur == "💧 Tambah Watermark Teks":
        st.markdown("#### 💧 Tambah Watermark Teks ke PDF")
        pdf_wm_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_wm_file")
        if pdf_wm_file:
            _r = _PdfReader(_pdf_io.BytesIO(pdf_wm_file.read()))
            _n = len(_r.pages)
            st.info(f"📄 File ini memiliki **{_n} halaman**.")
            col_wm1, col_wm2 = st.columns(2)
            with col_wm1:
                pdf_wm_text = st.text_input("Teks watermark:", value="RAHASIA", key="pdf_wm_text")
                pdf_wm_size = st.slider("Ukuran font:", 20, 120, 60, key="pdf_wm_size")
                pdf_wm_opacity = st.slider("Transparansi (0=bening, 1=solid):", 0.0, 1.0, 0.15, step=0.05, key="pdf_wm_opacity")
            with col_wm2:
                pdf_wm_color = st.selectbox("Warna teks:", ["Abu-abu", "Merah", "Biru", "Hitam"], key="pdf_wm_color")
                pdf_wm_angle = st.slider("Sudut kemiringan (derajat):", 0, 90, 45, key="pdf_wm_angle")
                pdf_wm_pages = st.text_input("Halaman (kosongkan = semua):", value="", key="pdf_wm_pages")
            pdf_wm_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="PDF_Watermark", key="pdf_wm_outname")
            if st.button("💧 Tambah Watermark", key="btn_pdf_wm", type="primary"):
                try:
                    _color_map = {
                        "Abu-abu": (0.5, 0.5, 0.5),
                        "Merah": (0.8, 0.0, 0.0),
                        "Biru": (0.0, 0.0, 0.8),
                        "Hitam": (0.0, 0.0, 0.0),
                    }
                    _cr, _cg, _cb = _color_map[pdf_wm_color]

                    def _make_watermark(text, size, angle, opacity, r, g, b, page_width, page_height):
                        _buf = _pdf_io.BytesIO()
                        _c = _rl_canvas.Canvas(_buf, pagesize=(page_width, page_height))
                        _c.setFillColor(_rl_Color(r, g, b, alpha=opacity))
                        _c.setFont("Helvetica-Bold", size)
                        _c.saveState()
                        _c.translate(page_width / 2, page_height / 2)
                        _c.rotate(angle)
                        _c.drawCentredString(0, 0, text)
                        _c.restoreState()
                        _c.save()
                        _buf.seek(0)
                        return _PdfReader(_buf).pages[0]

                    def _parse_ranges_wm(s, max_page):
                        if not s.strip():
                            return set(range(1, max_page + 1))
                        pages = set()
                        for part in s.split(","):
                            part = part.strip()
                            if "-" in part:
                                a, b = part.split("-", 1)
                                pages.update(range(int(a.strip()), int(b.strip()) + 1))
                            elif part:
                                pages.add(int(part))
                        return set(p for p in pages if 1 <= p <= max_page)

                    pdf_wm_file.seek(0)
                    _r2 = _PdfReader(_pdf_io.BytesIO(pdf_wm_file.read()))
                    _wm_pages_set = _parse_ranges_wm(pdf_wm_pages, _n)
                    _w = _PdfWriter()
                    for _i, _pg in enumerate(_r2.pages, 1):
                        if _i in _wm_pages_set:
                            _pw = float(_pg.mediabox.width)
                            _ph = float(_pg.mediabox.height)
                            _wm_page = _make_watermark(
                                pdf_wm_text, pdf_wm_size, pdf_wm_angle,
                                pdf_wm_opacity, _cr, _cg, _cb, _pw, _ph
                            )
                            _pg.merge_page(_wm_page)
                        _w.add_page(_pg)
                    _out = _pdf_io.BytesIO()
                    _w.write(_out)
                    _out.seek(0)
                    st.success(f"✅ Watermark '{pdf_wm_text}' berhasil ditambahkan ke {len(_wm_pages_set)} halaman.")
                    st.download_button(
                        label="📥 Download PDF dengan Watermark",
                        data=_out.getvalue(),
                        file_name=f"{pdf_wm_outname.strip() or 'PDF_Watermark'}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 6. NOMOR HALAMAN ───────────────────────────────────
    elif pdf_fitur == "🔢 Tambah Nomor Halaman":
        st.markdown("#### 🔢 Tambah Nomor Halaman ke PDF")
        pdf_pn_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_pn_file")
        if pdf_pn_file:
            _r = _PdfReader(_pdf_io.BytesIO(pdf_pn_file.read()))
            _n = len(_r.pages)
            st.info(f"📄 File ini memiliki **{_n} halaman**.")
            col_pn1, col_pn2 = st.columns(2)
            with col_pn1:
                pdf_pn_pos = st.selectbox("Posisi nomor halaman:", ["Bawah Tengah", "Bawah Kanan", "Bawah Kiri", "Atas Tengah", "Atas Kanan", "Atas Kiri"], key="pdf_pn_pos")
                pdf_pn_size = st.slider("Ukuran font:", 6, 24, 10, key="pdf_pn_size")
            with col_pn2:
                pdf_pn_format = st.text_input("Format nomor (gunakan {n} = nomor halaman, {total} = total):", value="Halaman {n} dari {total}", key="pdf_pn_format")
                pdf_pn_start = st.number_input("Mulai penomoran dari angka:", min_value=1, value=1, key="pdf_pn_start")
            pdf_pn_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="PDF_Bernomor", key="pdf_pn_outname")
            if st.button("🔢 Tambah Nomor Halaman", key="btn_pdf_pn", type="primary"):
                try:
                    _pos_map = {
                        "Bawah Tengah": ("center", "bottom"),
                        "Bawah Kanan": ("right", "bottom"),
                        "Bawah Kiri": ("left", "bottom"),
                        "Atas Tengah": ("center", "top"),
                        "Atas Kanan": ("right", "top"),
                        "Atas Kiri": ("left", "top"),
                    }
                    _halign, _valign = _pos_map[pdf_pn_pos]
                    _margin = 20

                    def _make_page_number(num_text, size, page_width, page_height, halign, valign, margin):
                        _buf = _pdf_io.BytesIO()
                        _c = _rl_canvas.Canvas(_buf, pagesize=(page_width, page_height))
                        _c.setFillColor(_rl_Color(0, 0, 0, alpha=0.7))
                        _c.setFont("Helvetica", size)
                        _tw = _c.stringWidth(num_text, "Helvetica", size)
                        if halign == "center":
                            _x = page_width / 2
                        elif halign == "right":
                            _x = page_width - margin - _tw / 2
                        else:
                            _x = margin + _tw / 2
                        _y = margin if valign == "bottom" else page_height - margin - size
                        _c.drawCentredString(_x, _y, num_text)
                        _c.save()
                        _buf.seek(0)
                        return _PdfReader(_buf).pages[0]

                    pdf_pn_file.seek(0)
                    _r2 = _PdfReader(_pdf_io.BytesIO(pdf_pn_file.read()))
                    _w = _PdfWriter()
                    for _i, _pg in enumerate(_r2.pages, 1):
                        _num = _i + int(pdf_pn_start) - 1
                        _label = pdf_pn_format.replace("{n}", str(_num)).replace("{total}", str(_n))
                        _pw = float(_pg.mediabox.width)
                        _ph = float(_pg.mediabox.height)
                        _pn_page = _make_page_number(_label, pdf_pn_size, _pw, _ph, _halign, _valign, _margin)
                        _pg.merge_page(_pn_page)
                        _w.add_page(_pg)
                    _out = _pdf_io.BytesIO()
                    _w.write(_out)
                    _out.seek(0)
                    st.success(f"✅ Nomor halaman berhasil ditambahkan ke {_n} halaman.")
                    st.download_button(
                        label="📥 Download PDF dengan Nomor Halaman",
                        data=_out.getvalue(),
                        file_name=f"{pdf_pn_outname.strip() or 'PDF_Bernomor'}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

    # ── 7. KOMPRES / OPTIMALKAN ────────────────────────────
    elif pdf_fitur == "🗜️ Kompres / Optimalkan PDF":
        st.markdown("#### 🗜️ Kompres / Optimalkan Ukuran PDF")
        st.markdown("Mengurangi ukuran file PDF dengan menghapus data yang tidak diperlukan dan mengoptimalkan struktur file.")
        pdf_cmp_file = st.file_uploader("Upload file PDF:", type=["pdf"], key="pdf_cmp_file")
        if pdf_cmp_file:
            _orig_size = len(pdf_cmp_file.read())
            pdf_cmp_file.seek(0)
            st.info(f"📄 Ukuran asli: **{_orig_size / 1024:.1f} KB** ({_orig_size / 1024 / 1024:.2f} MB)")
            pdf_cmp_outname = st.text_input("Nama file hasil (tanpa .pdf):", value="PDF_Kompres", key="pdf_cmp_outname")
            if st.button("🗜️ Kompres Sekarang", key="btn_pdf_cmp", type="primary"):
                try:
                    pdf_cmp_file.seek(0)
                    _r2 = _PdfReader(_pdf_io.BytesIO(pdf_cmp_file.read()))
                    _w = _PdfWriter()
                    for _pg in _r2.pages:
                        _pg.compress_content_streams()
                        _w.add_page(_pg)
                    _w.compress_identical_objects(remove_identicals=True, remove_orphans=True)
                    _out = _pdf_io.BytesIO()
                    _w.write(_out)
                    _out.seek(0)
                    _new_size = len(_out.getvalue())
                    _saved = _orig_size - _new_size
                    _pct = (_saved / _orig_size * 100) if _orig_size > 0 else 0
                    if _saved > 0:
                        st.success(f"✅ Ukuran berkurang: {_orig_size/1024:.1f} KB → {_new_size/1024:.1f} KB (hemat **{_pct:.1f}%**)")
                    else:
                        st.info(f"ℹ️ Ukuran: {_orig_size/1024:.1f} KB → {_new_size/1024:.1f} KB (PDF sudah teroptimalkan)")
                    _out.seek(0)
                    st.download_button(
                        label="📥 Download PDF Hasil Kompresi",
                        data=_out.getvalue(),
                        file_name=f"{pdf_cmp_outname.strip() or 'PDF_Kompres'}.pdf",
                        mime="application/pdf",
                        use_container_width=True
                    )
                except Exception as _e:
                    st.error(f"❌ Gagal: {str(_e)}")

st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #64748b; padding: 1rem;">
    <p>📊 Aplikasi Perbandingan Data Realisasi Impor</p>
    <p style="font-size: 0.8rem;">Dibuat dengan ❤️ menggunakan Streamlit</p>
</div>
""", unsafe_allow_html=True)
